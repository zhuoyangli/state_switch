
################################################################################################################
################################################################################################################
################ \ ### / ################################################## | ##################################
################# \ # / ###### | ########################################## | ########## | #####################
################## \ / ################ _____  ###### |______  ############ | ##################################
################## / \ ####### | ##### / ### | ###### | #### | ############ | ########## | #####################
################# / # \ ###### | #### | #### | ###### | #### | ############ | ########## | #####################
################ / ### \ ##### | #### | #### | ###### | #### | ############ | ########## | #####################
############### / ##### \ #### | ##### \_____|\__ ### | #### | ############ |_______ ### | #####################
################################################################################################################
################################################################################################################
############### This library of helper functions are written by Xian Li at JHU 09/2020-09/2024 #################
################################################################################################################

#Catalog

    # to_tensor(x)
    # to_ndarray(x)
    # to_list(x)
    # convert_list_array(inp)
    # split_int(input_int)
    # ls2str(inls,join_by=',')
    # intls2strls(int_ls)
    # int2str_underscore(input_int,loc=-1)
    # intls2strls_underscore(input_strls,sep =',')
    # str_to_num(in_str,num_type='intfloat',keep_as_str='_')
    # array_str_to_num(array,to_int=False,to_list=False)
    # strdict_to_dict(strdict)
    # num_to_alphabet(numlist,case='upper')
    # concat_mat(ndarray1,ndarray2,dim=0,mute=False)
    # maxpool(ndarray)
    # maxpool_inds(local_conn)
    # avg(mat,ignore_nan=False,dim=0,tolist=True,mute=False)
    # avg_df(multi_dfs,z=True,to_df=True)
    # std(vector)
    # se(vector)
    # se_ndarray(list_of_lists)
    # summ(mat,sum_dim=0,tolist=True,mute=False)
    # dot(ndarray,ndarray2=False)
    # accum(list,tolist=True)
    # tri(num)
    # binarize(num,thresh= 0)
    # binarize_mat(mat,thresh= 0)
    # list_nonzero(inlist)
    # adjust_baseline(vector,baseline='default')
    # zscore(ndarray,dim=0)
    # zscore_df(df,to_df=True,axis=1)
    # zscore_multi_dfs(multi_dfs,to_df=True,axis=1)
    # shape(ndarray)
    # skip_nan_array(array)
    # skip_nan_ndarray(ndarray)
    # flatten(mat)
    # transpose_mat(mat):

    # pca_tsne(vectors,n_comp=2)
    # fisherZ(r)
    # spearman_corr(vec1,vec2)
    # cronbach_alpha(df)
    # t_test(list1,list2=False,null=0,test_type='ind',var_eq=False)
    # f_test_one_way(inlists)
    # trunc(timeseries,time,window)
    # end_sentence(textls)
    # clean_quotes(input_str,replace_item='\'',replaced_by='‘’',use=1)
    # clean_text(text,clean_up='Default')
    # clean_list(strlist,clean='Default')
    # clean_consec3(inlist,start_ind=2)
    # clean_add_missing(inlist)
    # clean_psiturk_route(inlist)
    # split_num(num,thres=5)
    # split_num_ls(num_ls,thres=5)
    # split_num_ls_remain(rcl,thres=6)
    # pad_int(in_int,dig=3,pad='0')
    # insert_str(file_list,infront_of_where,str_to_insert)
    # replace_str(file_list,str_to_delete, str_to_replce='')
    # replace_item(inlist,to_be_replaced,new_item)
    # rename_files(in_path,file_list,rename_list)
    # file_size(file,unit='kb')
    # file_exist(file)
    # shutil_cp(from_file,to_file,check_exist=False)
    # os_sys_cp(from_file,to_file,check_exist=False)
    # os_buf_cp(from_file,to_file,check_exist=True,buffer=1,mute=False)

    # get_legacy_val(inlist)
    # get_intersect(list1,list2)
    # get_intersect_all(lists)
    # get_interval(inlist)
    # get_ls_interval(input_lists, start='default', end='default')
    # get_union(list1,list2,uniq=False)
    # get_subtract(list1,list2)
    # get_unique_repeat(inlist)
    # get_rep_thr(inlist,minmax=(None,None))
    # get_first_elem_ind(inlist,rng=[1,30],repeat_ind=False)
    # get_indices(search_list,search_item,find=False)
    # get_inds_lsofls(list_of_lists,search_item)
    # nth_percentile(vector, n=50)
    # inds_above_thresh(vector, thresh)
    # inds_ls_above_thresh(ndarray,nth_prcnt=50)
    # translate_pairs(tobe_translated_list,decipher_pairing)
    # get_sequence(seq1,seq2,offset=1)
    # get_min_max(y_axis,roundoff=True)
    # get_range(start,end,step)
    # get_range_float(rng,points=100)

    # distance_mat(ndarray)
    # get_sim_ratio(array1,array2,skip_nan_min=5)
    # get_pairsim_ratio(ndarray1,ndarray2,thresh_n=5,fill='nan')
    # get_corr(array1,array2,skip_nan_min=30)
    # get_paircorr(ndarray1,ndarray2,fill='nan')
    # cosine_sim(ndarray1,ndarray2)
    # get_cosine_sim(ndarray1,ndarray2)
    # get_pairwise_cosinesim(ndarray1,ndarray2)
    # pairwise_col_correlation(matrix1, matrix2)
    # pairwise_row_correlation(ndarray1, ndarray2)
    # get_pairwise_pearsonr(ndarray)
    # get_pairwise_corr(ndarray1,ndarray2)
    # paircorr_list(ndarray1,ndarray2)
    # paircorr_avg(vecs1,vecs2)
    # compute_pairs(list1, list2, compute_func, lists_name=None, default_params=None)
    # get_pairs(list_of_lists,lists_name='list',connect='_',direction=1)
    # get_triu_mat(mat)
    # get_tril_mat(mat)
    # get_triu_coords(list_of_lists,lists_name='list',diag_offset=1)
    # get_triu_val(mat,diag_offset=1)
    # get_tril_coords(list_of_lists,lists_name='list',diag_offset=1)
    # get_tril_val(mat,diag_offset=1)
    # get_diag_val(mat,diag_offset=0)
    # get_offdiag_val(mat,diag_offset=1)
    # avg_below_diag(matrix, num_cells=5)
    # get_non_nans(array)
    # get_nonnan_vals(array)
    # zero_matlow(inmat,k=1)
    # zero_matup(inmat,k=1)
    # set_mat_diag(inmat,k=0,set_to=0)
    # keep_mat_diag(inmat,k=0,others_set_to=0)
    # check_list_thr(inlist, thr=0,check='neg')
    # check_corr(check_lists,lists_name,make_plot=False,title=False,corr_path=False,mute=False)
    # find_peak_latency(timeseries,time,window='all',peak='high')
    # find_halfmax_point(lagcorr_vec,half_max=0.5,interpolate=False)
    # mat_halfmax(mat,half_max=0.5,interpolate=True,rnd=4,window=10,plot=True,to_inds='None',plot_title='Auto-correlation_half-maximum',outpath=False)

    # apply_inds(indlist,inlist)
    # apply_inds_2d(coordlist,inmat)
    # assign_val(inds,inlist,val=0)
    # translate_items_inls(inlist,search='',set_search_vals=0, set_other_vals=1)
    # get_vals(ls1,ls2,forls2=True)
    # get_keydata(partial_key,data):
    # get_cp_items(ls1, ls1_items, ls2)
    # fastfetch_keydata(data_storage,target_key)
    # fastfetch_coldata(file,col_num,target_key,target_dtype)
    # assemble_dict(lists,keynames,stack=True,stack_keyname=['val','cond'],to_df=True)
    # dict_to_lists(data_as_dict)
    # dict_to_df(data_as_dict,colname_by_key=False)
    # keep_rows(keep_start,keep_len,ind_start=1)
    # exclude_rows(all_start,all_len,keep_start,keep_len,ind_start=1)
    # list_exclude_inds(inlist,pop_inds)
    # list_select_inds(inlist,keep_inds)
    # list_rm_mid_items(inlist, n_remove):
    # lists_to_dict(list_of_lists,header='row1')
    # list_add_vals(ls,vals)
    # lists_add_vals(lists,vals)
    # dict_sort_inds(data,sort_by_key,target_key_seq=None)
    # dict_exclude_inds(data,exclude_inds)
    # dict_select_inds(data,select_key,val_match,find=False)
    # dict_exclude_keys(data,exclude_keys='')
    # dict_include_keys(data,keep_keys)
    # dict_exclude_rows(data,by_key,by_search,exclude=1)
    # dict_collapse(data_dict,uni_per_key=False)
    # group_cut(to_be_cut_list,cut_by_list,cut_item)
    # group_by_list_renum(group_by_list)
    # group_by_both_list(group_by_list1,group_by_list2)
    # group_by(group_by_list,to_be_grouped_list)
    # group_lists(list_of_lists,group_by_list_ind=0)
    # group_dict_shrink(data_dict,sep=' ')
    # group_dict(data_dict,group_by_key,exclude_keys=False)
    # group_means(grouped_data,n_groups,keys_not_cal,keys_exclude=[],rnd=2)
    # split_list(inlist)
    # sort_nlists(list_of_lists)
    # merge_items_in_list(list_of_lists,sep=' ')
    # merge_tuples_in_list(list_of_tuples)
    # merge_lists_in_list(list_of_lists)
    # merge_ioflist_in_data(data)
    # merge_lsoflist_in_data(data)
    # flip_lsoflist(list_of_lists)
    # shared_overall(vec1,vec2)
    # shared_history(vec1,vec2,curr_item=False)

    # data_info(in_data)
    # datakey_samelen(data)
    # data_subset(data_dict,keep_keys)
    # data_storage(file_path,file,col_num,target_dtype)
    # data_by_key(data,clear_empty=True,clear = '')
    # dict_by_key(data,keynames='default')
    # grab_all(path,filetype,exclude=False,mute=False)
    # grab_fname(path,search_name,exclude=False,mute=False)
    # grab_base(path,file_search='*',exclude=False,mute=False)
    # read_table(infile,sheet=1)
    # read_rows(table,col_num,key=True,skip_row1=True,clear_empty=True,clear='')
    # read_data(infile,col_num,sheet=1,key=True,skip_row1=True,clear_empty=True,clear='',clean_keylist=False)
    # read_data_num(infile,col_num,num_type='float')
    # num_data_clean(data,keylist='all',num_type='float',exclude='',get_mean=False,pop_empty=False,keep_as_str=False)
    # clean_data(data,numkey=False,num_type=False,keep_as_str=False,textkey=False)
    # read_csv(vectorfile)
    # read_csv_data(rawfile,asnum=True,mute=True,header=False,todict=False)
    # read_text(filename,sep=False)
    # read_docx(filename,clean_up='Default')
    # init_mat(shape)
    # init_list(length,label='lab')
    # init_dict(keylist,repeat=False)
    # combo_ls(lists_of_strings,join_by='_')
    # dict_add_keys(dic,keylist,repeat=False) 
    # dict_add_vals(dic,keys,vals)
    # dict_add_keys(dic,keylist,repeat=False)
    # dict_add_keys_n_vals(dic, key_list,val_list,append=True,verbose=False)
    # reverse_dict(data,by_key)
    # mkdir(inpath)
    # write_dict_samelen(dic,outpath,header='default',dict_ind='default',convert=False,set_format=False)
    # write_dict_anylen(dic,outpath,header='default',allstr=False)
    # write_array(array,outpath,vertical=True)
    # write_ndarray(ndarray,outpath,vertical=True)
    # write_csv(ndarray,outfile,transpose=True)
    # write_csv_ndarray(ndarray,outname):
    # write_highdim_to_csv(ndarray,out,delimiter=",")
    # load_highdim_from_csv(file,shape,delimiter=",")
    # write_txt(instr,outfile)
    # write_docx(content,outpath)
    # write_json(data_dict,outname='./out.json')

    # thresh_mat()
    # rotate_mat(matrix, n=1)
    # zero_pad_mat_diag(original_matrix,new_size=4,direction=1)
    # nan_pad_mat_bone(original_matrix,new_size=4,direction=1,mid_window=15,time_window=20)
    # mat_diag(in_mat,val=1)
    # mat_bound(in_mat,submat,val=-1)
    # mat_info(in_mat,mute=False)
    # plot_figsize(wd,ht)
    # sns_style(style='default') 
    # plot_xylim(x_rng=False,y_rng=False)
    # plot_save(outpath,ftype='.png')
    # plot_bar_se(df,x_var,y_var,kind='bar',alpha=.7,outpath=False,title=False,add=False)
    # plot_bar(df,x_axis,y_axis,split_by,dodge=False,outpath=False,title=False,add=False)
    # plot_bar_stack(df,x_axis,split_by,stats,bg_style='default',color='default',outpath=False,title=False,add=False)
    # plot_bar_sorted(mean_ls, se_ls, cond_ls, pval_ls,thresh=.05,figsize=(50,12),xlab='Conditions',ylab='Mean Value',cb_lab='Magnitude of Means',colormap='viridis', title='Bar Plot of Sorted Means with Standard Errors',fontsize=20,outpath=False,dpi=300,fig_format='.png',add=False):
    # plot_label(x_coords,y_coords,labels='default',font_size=15,xloc=-1,yloc=0.1,add=False)
    # plot_line(y,x='Default',y_lab='None',x_lab='None',legend='Default',title='None',start=0,add=False)
    # plot_anyline(x,y, figsize=(12, 6),xlab='x-axis: list-A',ylab='y-axis: list-B',title='Line Plot',add=False,outpath=False,fig_format='.png'):
    # plot_anylines(x, y_lists, legend_labels, legend_pos_fixed='upper right',figsize=(12, 6), xlab='x-axis: list-A', ylab='y-axis: list-B', title='Line Plot', add=False, outpath=False, fig_format='.png', pad=0.2)
    # plot_vline(x_intercept,y_range,add=False)
    # plot_hline(y_intercept,x_range,add=False)
    # plot_shade(rng,direction='vertical',add=False,outpath=False,title='plot',alpha=.15,lw=0,color='grey',x='default')
    # plot_shades(rng_ls,direction='vertical',add=False)
    # plot_line_pairdot(array1,array2,xlab=['array1','array2'],title=['',20],xylim=[(-0.5,1.5),(-0.15,0.6)],xylab=['',''],xylab_size=(30,30),xytick_size=(20,20),dot_size=70,dot_c='b',fig_hw=(4,8),fig_size=80,lwd=0.3,outpath=False)
    # plot_dist(array,outpath=False,mean=False,median=False,base=False,rug=False,y_rng=(0,1))
    # plot_hist(array,outpath=False,add=False)
    # plot_hist_with_value(value, value_distribution, bins=50,outpath=False,fig_format='.png',dpi=300,add=False)
    # plot_2dmat(mat,xylab=False,title='ISC_Time-Time-Corr',fontsize=20,figsize=(12,10),outpath=False,save_matrix=False)
    # plot_matrix(mat,title=False,outpath=False,labname=False,diagmask=False,map_color="YlGnBu")
    # plot_dots(x,y,labels=False,fontsize=15,marker='o',markerfacecolor='blue',markersize=12,linestyle='none', linecolor='green',add=False,outpath=False,title='scatter')
    # plot_dots_3d(x,y,z,figsize=(10,8),labels=False,markersize=200,markercolor='Green',cmap='Greens',edgecolor='black', linestyle='none',linecolor='grey',outpath=False,add=False,title='scatter3d')
    # plot_scatter(array1,array2,make_plot=True,outpath=False,title=False,xlab=False,ylab=False,pointlab=False,mute=False)
    # plot_swarm(data_as_dict,y,x,outpath=False,color='black',size=7,alpha=1,edgecolor='black')
    # plot_violin(data_as_dict,y,x,outpath=False,color=False,split=False,inner_stick=False,bandwidth=False,palet='muted')
    # plot_autocorr_halfmax(lag_corr,halfmax_point,fontsize=20,outpath=False,style='seaborn-whitegrid',title='Auto-correlation_half-maximum',fig_format='.png')

    # calc_onset_from_dur(inlist,start_time=0)
    # aud_dur(aud_name,unit='s')
    # aud_sound_time(audfile,decimal=0.001,thr=0.1)
    # aud_start_end(audfile,decimal=0.001,thr=0.1)
    # aud_dur_ms(audfile,decimal=0.001,thr=0.1,unit='ms')
    # aud_env(audfile,tr,zscore=False,outcsv=False)
    # aud_rms(audpath,audname,unit='s',plot=True)
    # aud_onset_offset(audpath,audname,unit='s',thr=0.01)
    # downsample_rms(df_aud_rms,plot=True)
    # create_bp(aud_ls,onset_ls,offset_ls,order='default')
    # export_bp(bp_df,outfile='none',comb_gap='gap.wav',export=True)
    # export_audio(bp_df,outfile,unit='ms',fmat='wav')
    # scramble_audio(bp, aud_stories, output_name)
    # calc_gap_length(aud, mindist=None, TRdur=1500)
    # gap_dict(row, mingap, TRdur)
    # insert_gaps(bp, mingap=1000, TRdur=1500, outname=None)
    # fix_col_order(df)
    # scram_list(inlist)
    # scram_worker(item)
    # scramble_lists_in_parallel(inlist, n_permutations)
    # count_combo(unique_items, repeated_items)
    # select_from_repeated_items(unique_items, repeated_items, repeated_couple)
    # compute_trial_select_from_repeated_items(args)
    # select_and_compute(unique_items, repeated_items, repeated_couple, data, n=1000, batch_size=100)
    # pval_from_dist(value, value_distribution)

    # Choose Your Own Adventure (CYOA)
        # get subjects' information for rating files 
            # quiry_yroute(yoked_to='all',yoked_sum='/Users/xianl/Desktop/CYOA/FixR_data/yoked/try/2_data/1_summary.xlsx')
            # quiry_pid(debugid,pidpath='/Users/xianl/Desktop/CYOA/PassR_data/collect_data/data/first28/questiondata.xlsx')
            # cyoa_get_eventinfo(maptable,eventinfo)
            # cyoa_get_merge(mapfile,quiry_events,col_num=[1,2,10],merge='Y',pop_choice=False,choice_marker='_')
            # cyoa_get_choice_vec(story_route)
            # cyoa_get_choice(story_route,choice_marker='_')
            # cyoa_fake_choice(mapfile,col_num)
            # cyoa_get_choiceinfo(mapfile,eventinfo,choices,col_num=1)
            # cyoa_subj_paths(mapfile,sumfile,subj_ind,sum_cols=[1,2,6],keynames=['scene_label','event_label','scene_text'])
            # cyoa_subj_pass(mapfile,sumfile,subj_ind,sum_cols=[2,5,6],keynames=['scene_label','event_label','scene_text'])
            # cyoa_subj_yoked(mapfile,sumfile,subj_ind,sum_cols=[2,7,9,10],keynames=['scene_label','event_label','scene_text','actual_choice','want_not'])
            # cyoa_get_survey_heroku_qdata(ques,resp,search_q='uncertainty_')
            # cyoa_cut_events(data)
        # semantic and causal centrality:
            # get_embeds(all_lines)            
            # get_embeds_glove(glove_text="glove.6B.300d.txt")
            # find_closest_embeddings(embed_dict,word,top=10)
            # theme_sim(rate_word,theme_word,embed_dict,measure='max')
            # semantic_max_match(reference,test,title=False,outpath=False)
            # semantic_cent(all_lines,title=False,outpath=False)
            # modify_conn_weight(choice_events,caus,efct,rew_caus_efct,w1_w2)
            # causal_cent(table,col1,col2,row1,pair,title=False,outpath=False,mute=True,re_weigh=False,rew_events=False,rew_caus_efct=(1,1))
            # contextual_cent(table,col1,col2,row1,pair,title=False,outpath=False,mute=True,re_weigh=False,rew_events=False,rew_curr_cntx=(1,1),decimal=6)
            # cent_bymat(mat,decimal=6)
            # isc_cent(allsub_cents,method='pairwise',mute=True,skip_nan_min=30)
            # isc_cent_cross_group(group1_cents,group2_cents,methods='all',legend=['g1','g2'])
            # align_rcl_to_ev(rcl_ev,rcl_embed,n_ev)
        # draw sample:         
            # sample_one_val_per_key(data_dict,nsample=1)
            # compute_corr_per_sample(sample_dict,cols_to_corr=[1,2])
        # recall temporal:
            # recall_temp_mat(rcl,tot)
            # process_recall_seg(rcl_seg)
            # recall_seg_split_num(rcl,thres=6)
            # recall_temp_vio(rcl)
        # GPT4 API calling
            # mesg(query)
            # ask_model(query,model='gpt-4')
            # correct_quote_in_key(inlist,key_ls=['event_A_number','event_B_number','rating','reasoning'])
            # correct_quote_in_keydata(inlist)
            # process_resp(resp,outpath_f=False)
        # recall by local eventnet conn
            # clean_dot(ls)
            # pred_acc(overlap) 
        # neighbor encoding effect
            # neighbor_pred(rcl,ignore_nan=True,thres=0)
            # neighbor_forward_pred(rcl,ignore_nan=True,thres=0)
            # neighbor_clusters(rcl)

    # Models of Mind and Brain:
        #optimize & matrix transform:
            # rand_vec(vec_size,rng=[0,1],integer=True)
            # plot_points(points,labels='dot')
            # plot_distance_mat(points,labels='dot')
            # rand_points(npoints,ndims,labels='dot',plot=False)
            # get_distance(points,plot=False)
            # shepard_stress(locations_input, targetdistances, npoints, ndims, plot=False)
            # find_locations(targetdistances,npoints,ndims,plot=False)
            # points_circle(npoints=16,radius=1,phaseshift=0)
            # plot_transform(orig_points,trans_points,npoints)
            # mat_transform(points,scale=1,trans_mat='none',rot='none')
            # nlinear_trans_demo(metric,x='default',overlay=True)
        # hopfield net 
            # toLong(A)
            # toSquare(longvector)
            # heb_getWeights(x)
            # hop_update(w,x)
            # hop_trainPatterns(M)
            # get_img_data(img,convert='1/-1',plot=False)
            # plot_states(states,num_states,num_neurons)
            # plot_hop_tc(dynamics,states,num_neurons,num_states,timesteps)
            # plot_hop_match_tc(dynamics,states,state_labs=False)
            # plot_hop_startend(dynamics,num_neurons,timesteps)
            # hopfield_net(input,states,num_neurons,num_states,timesteps,plot_startend=True,plot_timecourse=True,state_labs=False,plot_wmat=True)

    # Naturalistic fMRI data processing:
        #fmri general operations
            # get_nifti_data(file)
            # get_nifti_info(file)
            # nifti_select_time(brain_file,time_lists)
            # get_mask(maskn='50',plot_mask=True)
            # plot_nifti(nii_file,mode='stats',outpath=None)
            # plot_anat(nii_f,minmax=(0,1000),outpath=None)
            # nifti_split_hemi(mask_f,x_mid)
            # nifti_math(nii_f1,nii_f2,math='sub',outpath='default')
            # export_nifti(mat,outpath)
            # wipe_mat(mat,retain,wipe_out=0,binary=False,dtype='int',outpath=False)
            # make_mask(atlas_f,retain,outpath='default',resample=False,interp='nearest')
            # apply_mask(data_file,mask_file,keep=1,pop=0)
            # lagcorr(swav,a1tc,unit=30,yrng=(-25,25),peak=-3)
            # auto_correlation(tc,unit=30)
            # auto_correlation_halfmax(tc,unit=30,half_max=0.5,interpolate=False,show_fig=True,outpath=False,title='Auto-correlation_half-maximum')
            # extract_mvp(data_file,mask_file,trs,mute=True)
            # extract_group_mvp(data_file,mask_file,trs,mute=True)
            # extract_voi(data_file,maskn=200,outfile=False,to_csv=True)
            # load_voi(voi_file,ftype='.csv')
            # load_group_voi(voi_folder,ftype='.csv')
        #fmri GLM related:
            # coord_mni2vox(x,y,z,to_int=False)
            # coord_vox2mni(x,y,z,to_int=False)
            # resample_map(source_img,targ_img,outpath=False,interp='nearest')
            # xyz_label_name(cort_num, subc_num, cort_labtxt,subc_labtxt)
            # xyz_label_num(xyz, atlas_cort_img, atlas_subc_img,xyz_to_vox=True)
            # xyz_label_info(xyz, atlas_cort_img, atlas_subc_img,cort_labtxt,subc_labtxt, xyz_to_vox=True)
            # add_conditions(design_mat,factor_list='default',show=False,outpath=False,add_cons=False)
            # add_contrasts(design_mat,conditions,factor_list='default',show=False,outpath=False,fig_format='.png')
            # plot_conds(conditions,cond_list,design_mat,outpath=False,fig_format='.png')
            # plot_design_mat(design_mat,outpath=False)
            # model_glm(brain_img,events,tr,title=False,hrf='spm',std=False,hp=.01,fwhm=None,roi_mask=None,regressors=None,drift='cosine',outpath=False)
            # thresh_map(stat_map,clust_thresh=20,thresh=1.,alpha=.05,mcc='fdr',outpath=False,show=False)
            # show_glm_results(z_map,alpha=.05,mcc='fdr',clust_thresh=20,title=None,outpath=False,fig_format='.png',view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior'],thresh=1.)
            # save_glm_results(conditions,fmri_glm,outpath,show=True,title=None,alpha=.05,mcc='fdr',clust_thresh=20,thresh=1.,view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior'])
        #fmri processing and ISC related
            # post_fmriprep(base_dir,sub_list,tr,fwhm=6,outlier_cutoff=3,hdf5=False,do_roi=False)
            # hdf5_apply_mask(fmriprep_dir,roi_ind,subj='*',name_in_file='*hdf5',maskn = '50')
            # plot_hyper_isc(data,hyperalign=False,transformed_list=False,voxel_ind=0,tr_ind=0,nsubj=5,lab_ftsize=16,title_ftsize=18,wid=15,hgt=5,map_color='RdBu_r')
            # view_surf(isc_r,isc_p,maskn='50',thresh=.005,outpath=False)
            # view_brain(brain_img,surf_type='smooth',view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior'],hemi_list=['left','right'],title=False,outpath=False,thresh=1.,view=True)
            # combine_figures_w_title(fpath, ftype='.png',outpath=False,show=False,columns='default', dpi=300,title=False,select=['lateral','medial']):
            # crop_image2(img, crop_tuple)
            # combine_figures_w_title2(fpath, ftype='.png', outpath=False, show=False, columns='default', dpi=300, title=False, select=['lateral', 'medial'], crop_tuple=(0, 0, 0, 0))
            # combine_figures(fpath, ftype='.png', outpath='combined_image.png', columns=2, select=['lateral', 'medial'])
            # crop_image(img, left, right, top, bottom)
            # combine_and_crop_images(fpath, crop_sides, ftype='.png', outpath='combined_cropped_image.png', columns=2, select=['lateral', 'medial'], title_space=100)
            # plot_srm_isc(srm=False,transformed_list=False,component_ind = 0,nsubj=4,wid=15,hgt=8,lab_ftsize=16,title_ftsize=18,map_color='RdBu_r')
            # srm_align(all_data,n_features=100,component_ind=1)
            # hyper_align(all_data,voxel_ind=0,tr_ind=0,nsubj=5)
            # funcalign_new_data(new_data,model,tr_ind=0,mthd='procrustes')
            # funcalign_cross_validate(model,test_data,mthd='procrustes',voxel_ind=0,component_ind = 0,tr_ind=0)
            # hdf5_extract_voi(fmriprep_dir,roi_ind='all',subj='*',name_in_file='*hdf5',maskn='50')
            # combine_voi_csv(fmriprep_dir,comb_root,comb,subj='*')
            # get_subject_roi(data, roi)
            # plot_circular_shift(sub_rois, sub, sampling_freq)
            # plot_phase_randomize(sub_rois, sub, sampling_freq)
            # bootstrap_subject_matrix(similarity_matrix, random_state=None)
            # plot_isc_null(stats, method, add_title='none',savefig=False)
            # get_avg_others(ndarray,mute=True)
            # isc_avg(ndarray,fill='nan',mute=True,method='pearson')
            # isc_test(sub_timeseries,roi_ind='all',maskn='50',rand_method='circle_shift',plot_mask=False,plot_whole_brain=False,plot_null=False, permute = 5000, meassure='median', return_bt=True)
            # ttc(mvp_trs1,mvp_trs2='None')
            # time_time_mat(sim_mat,time_window,title,outpath=False,save_matrix=False,map_color='CMRmap')
            # plot_network_degree(mat,savefig=False)
            # get_isfc(sub_timeseries,outpath=False,plot_allsub=False,savefig=False,map_color='RdBu_r')
            # fc_test(isfc_allsub,fdr_thr=.0000001,plot_mat=False, plot_net=False,savefig=False,plot_brain=False,maskn='50')
        #fmri ISC
            # apply_phase_shift_and_correlate(target_series, series)
            # single_subj_isc_at_roi(data, target_subject_key, roi_index, num_permutations)
            # single_subj_isc(data, target_subject_key, roi_indices, num_permutations=100)



import random
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from scipy.stats import pearsonr


def to_tensor(x):
    #x can be any shape of ndarray or list of items/lists
    tensor = torch.from_numpy(np.array(x))
    return tensor

def to_ndarray(x):
    #x can be any shape of tensor or list of items/lists
    ndarray = [np.array(x) if type(x)==list else x.numpy()][0]
    return ndarray

def to_list(x):
    #x can be any shape of tensor or ndarray
    return x.tolist()

def convert_list_array(inp):
    #convert list to array or the reverse
    #input can be list or nd.array
    import numpy as np
    out = 'Error'
    if type(inp) == list:
        out = np.asarray(inp)
    elif type(inp) == np.ndarray:
        out = inp.tolist()
    return out 

def split_int(input_int):
    #this is to split an integer by its digits, e.g. a = split_int(3124) = [3,1,2,4]
    import regex as re
    a = input_int
    out = re.sub(r'(\d{1})', '\\1,', str(a)[::-1])[::-1]
    out = out.split(',')
    out = clean_list(out)
    return out

def ls2str(inls,join_by=','):
    if type(inls[0])!=str:
        inls = [str(i) for i in inls]
    out = join_by.join(inls)
    return out

def intls2strls(int_ls):
    return [str(i) for i in int_ls]
    
def int2str_underscore(input_int,loc=-1):
    #this is to convert 323_2 to '323_2' with inclusion of the underscore 
    #assuming underscore is in front of the last number of this integer, e.g. 323_2
    ls = split_int(input_int)
    part1 = ''.join(ls[0:loc])
    part2 = ls[loc]
    out_str = part1+'_'+part2
    return out_str

def intls2strls_underscore(input_strls,sep =','):
    #e.g. input = '[1,2,3,4,4_1,5,6]', outputs = ['1','2','3','4','4_1','5','6']
    ls = clean_text(a,clean_up=['[',']'])
    out = ls.split(sep)
    return out

def str_to_num(in_str,num_type='intfloat',keep_as_str='_'):
    #if keep_as_str = something, and this something is found in the in_str, keep it as str
    #otherwise, convert the str to a number as specified
    #keep_as_str can be a collection of things e.g. '_','@',etc., and should be in the form '_@' when entered
    if keep_as_str!=False:
        temp = get_subtract(in_str,keep_as_str)
    if keep_as_str!=False and len(in_str)!=len(temp):
        out=in_str
    elif num_type =='float':
        out =float(in_str)
    elif num_type =='int':
        out = int(in_str)
    elif num_type =='intfloat':
        out = int(float(in_str))
    return out

def array_str_to_num(array,to_int=False,to_list=False):
    #convert an array or list of str to a float array 
    import numpy as np
    if type(array)==list or type(array)==np.array:
        out = [float(i) for i in array]
        if to_int==True:
            out = [int(i) for i in out]
    else:
        out = np.array(array, dtype=np.float64, order='C')
        if to_list!=False:
            out = convert_list_array(out)
        if to_int!=False:
            out = [int(i) for i in out]
    return out

def strdict_to_dict(strdict):
    import json
    out = json.loads(strdict)
    return out

def num_to_alphabet(numlist,case='upper'):
    import string
    if case=='upper':
        alphabet_string = string.ascii_uppercase
    else:
        alphabet_string = string.ascii_lowercase
    alphabet_list = list(alphabet_string)
    out = [alphabet_list[i-1] for i in numlist]
    return out

def concat_mat(array1,array2,dim=0,mute=False):
    #concatenate two ndarrays (of the same dimension) along one of their existing dimensions
    import numpy as np
    out = np.concatenate((array1,array2),axis=dim)
    if mute!=True:
        print(np.shape(out))
    return out

def maxpool(ndarray):
    import numpy as np
    pool=[]
    for row in ndarray:
        m = np.max(row)
        pool.append(m)
    return pool

def maxpool_inds(local_conn):
    #get the indices of the max in the local_conn when max !=0
    if max(local_conn)==0:
        return []
    else:
        return get_indices(local_conn,max(local_conn))

def avg(mat,ignore_nan=False,dim=0,tolist=True,mute=False):
    #mat can be list, list of lists, nd.array(s)
    import numpy as np
    if ignore_nan==False:
        avg = np.mean(mat,dim)
    else:
        avg = np.nanmean(mat,dim)
    if mute==False:
        print('Current shape is:')
        print(np.shape(mat))
        print('Averaged shape:')
        print(np.shape(avg))
    if tolist==True:
        return avg.tolist()
    else:
        return avg

def avg_df(multi_dfs,z=True,to_df=True):
    #this is to average across multiple dfs (zscore each column before average)
    #e.g. avg_others = avg_df(sub_timeseries)
        #sub_timeseries is a dict,where each key/subj leads to a df of nTRs x nROI
    from scipy.stats import zscore
    import pandas as pd
    
    alldf = []
    if type(multi_dfs)==dict:
        for key in multi_dfs.keys():
            if z==True:
                zdf = zscore(multi_dfs[key],axis=1) #zscore across each col
                alldf.append(zdf)
            else:
                alldf.append(multi_dfs[key])
        out = avg(alldf)  
        sample_df = multi_dfs[key]
    elif type(multi_dfs)==list:
        for df in multi_dfs:
            if zscore==True:
                zdf = zscore(df,axis=1)
                alldf.append(zdf)
            else:
                alldf.append(df)
        out = avg(alldf)   
        sample_df = df
    #if output to ndarray or df
    if to_df==True:
        head = [i for i in sample_df.keys()]
        inds = list(sample_df.index.values) 
        out = pd.DataFrame(out, columns=head)
    return out

def std(vector):
    import numpy as np
    std = np.std(vector)
    return std

def se(vector):
    import numpy as np
    se = np.std(vector)/np.sqrt(len(vector))
    return se

def se_ndarray(list_of_lists):
    #get the upper and lower bound for each sub-list
    avg_all,upper,lower=[],[],[]
    for sample in list_of_lists:
        avg_all.append(avg(sample,mute=True))
        upper.append(avg(sample,mute=True)+se(sample))
        lower.append(avg(sample,mute=True)-se(sample))
    return avg_all,upper,lower

def summ(mat,dim=0,tolist=True,mute=False):
    #mat can be list, list of lists, nd.array(s)
    import numpy as np
    summed = np.sum(mat,dim)
    if mute==False:
        print('Current shape is:')
        print(np.shape(mat))
        print('Summed shape:')
        print(np.shape(summed))
    if tolist==True:
        return summed.tolist()
    else:
        return summed

def dot(ndarray,ndarray2=False):
    if ndarray2==False:
        ndarray2=ndarray
    return np.dot(ndarray,ndarray2)

def accum(list,tolist=True):
    #input = [1,2,3]; output = [1,3,6]
    import numpy as np
    out = np.cumsum(list)
    if tolist==True:
        out = convert_list_array(out)
    return out

def tri(num):
    # num can be either list/ndarray or a single float/int
    #this is to convert a number to 1/0/-1
    if type(num)==list or 'numpy.ndarray' in str(type(num)):
        bin_list = []
        for i in num:
            if i>0:
                bin_list.append(1)
            if i==0:
                bin_list.append(0)
            if i<0:
                bin_list.append(-1)
        return bin_list
    else:
        i = num
        if i>0:
            i = 1
        if i==0:
            i = 0
        if i<0:
            i = -1
        return i

def binarize(num,thresh= 0):
    # num can be either list/ndarray or a single float/int
    #this is to convert a number to 1/0
    if type(num)==list or 'numpy.' in str(type(num)):
        bin_list = []
        for i in num:
            if i>thresh:
                bin_list.append(1)
            if i<=thresh:
                bin_list.append(0)
        return bin_list
    else:
        i = num
        if i>thresh:
            i = 1
        if i<=thresh:
            i = 0
        return i

def binarize_mat(mat,thresh= 0):
    mat[mat>thresh]=1
    mat[mat<thresh]=0
    mat[mat==thresh]=0
    return mat

def list_nonzero(inlist):
    out = [i for i in inlist if i !=0]
    return out

def adjust_baseline(vector,baseline='default'):
    #this is to adjust every value in vector to the baseline (mean, std)
        #'default': take the average of the first half of the vector as the baseline
        #[0,1,2,3,4]: if input is a list, take the average of these indices of the vector as the baseline
        # 1.23: if input is a single number, take this number as the baseline
    import numpy as np
    if baseline=='default':
        basevec = vector[0:(int(len(vector)/2))]
        std = np.std(basevec)
    elif type(baseline)==list or type(baseline)==np.ndarray: #inds of baseline reference trs
        basevec = apply_inds(baseline,vector)
        std = np.std(basevec)
    elif type(baseline)==int or type(baseline)==float: 
        basevec = [baseline]
        std = np.std(vector)
    base = np.mean(basevec)    
    outvec = [(i-base)/std for i in vector]
    return outvec,base

def zscore(ndarray,dim=0):
    from scipy import stats as st
    out = st.zscore(ndarray,axis=dim)
    return out

def zscore_df(df,to_df=True,axis=1):
    from scipy.stats import zscore
    import pandas as pd
    out = zscore(df,axis=axis)
    #if output to ndarray or df
    if to_df==True:
        head = [i for i in df.keys()]
        out = pd.DataFrame(out, columns=head)
    return out

def zscore_multi_dfs(multi_dfs,to_df=True,axis=1):
    alldf= init_dict([key for key in multi_dfs.keys()])
    if type(multi_dfs)==dict:
        for key in multi_dfs.keys():
            zdf = zscore_df(multi_dfs[key],axis=axis)
            alldf[key]=zdf
    return alldf

def shape(ndarray):
    import numpy as np
    return np.shape(ndarray)

def skip_nan_array(array):
    out = [i for i in array if str(i)!='nan']
    return out

def skip_nan_ndarray(ndarray):
    out = []
    for ls in ndarray:
        out.append(skip_nan_array(ls))
    return out

def flatten(mat):
    ndarray=mat
    ndarray = np.array(ndarray)
    flattened_matrix = ndarray.reshape(-1, ndarray.shape[-2] * ndarray.shape[-1])
    return flattened_matrix

def transpose_mat(mat):
    return np.transpose(mat)

def pca_tsne(vectors,n_comp=2):
    #e.g. vectors = [glove_embed[word] for word in words]
    #n_comp is the number of principal components wanted
    #function outputs the [component1_vec,component2_vec,...]
    from sklearn.manifold import TSNE
    tsne = TSNE(n_components=n_comp, random_state=0)
    comps = tsne.fit_transform(vectors)
    return comps

def fisherZ(r):
    import numpy as np
    z = .5*(np.log(1+r)-np.log(1-r))
    return z

def spearman_corr(vec1,vec2):
    import scipy.stats
    r,p = scipy.stats.spearmanr(vec1,vec2)
    return r,p

def cronbach_alpha(df):
    import pandas as pd
    import numpy as np
    # 1. Transform the df into a correlation matrix
    df_corr = df.corr()
    # 2.1 Calculate N
    # The number of variables equals the number of columns in the df
    N = df.shape[1]
    # 2.2 Calculate R
    # For this, we'll loop through the columns and append every
    # relevant correlation to an array calles "r_s". Then, we'll
    # calculate the mean of "r_s"
    rs = np.array([])
    for i, col in enumerate(df_corr.columns):
        sum_ = df_corr[col][i+1:].values
        rs = np.append(sum_, rs)
    mean_r = np.mean(rs)
   # 3. Use the formula to calculate Cronbach's Alpha 
    cronbach_alpha = (N * mean_r) / (1 + (N - 1) * mean_r)
    return cronbach_alpha

def t_test(list1,list2=False,null=0,test_type='ind',var_eq=False):
    from scipy import stats as st
    if test_type=='ind' and list2!=False:
        out = st.ttest_ind(list1,list2,equal_var=var_eq)
    elif test_type=='paired' and list2!=False:
        out = st.ttest_rel(list1,list2)
    elif test_type =='1samp':
         out = st.ttest_1samp(list1,0)
    return out

def f_test_one_way(inlists):
    #inlists is a list of arrays, with each array being a list of sample means
    #e.g. inlists = [free_recall_means, yoke_recall_mean, pasv_recall_means]
    # free_recall_means = [sub1_recall_score, sub2_recall_score, ...]
    from scipy.stats import f_oneway
    if len(inlists)==3:
        out = f_oneway(inlists[0],inlists[1],inlists[2])
    elif len(inlists)==4:
        out = f_oneway(inlists[0],inlists[1],inlists[2],inlists[3])
    elif len(inlists)==5:
        out = f_oneway(inlists[0],inlists[1],inlists[2],inlists[3],inlists[4])
    elif len(inlists)==6:
        out = f_oneway(inlists[0],inlists[1],inlists[2],inlists[3],inlists[4],inlists[5])
    return out

def trunc(timeseries,time,window):
    #window: the time window to search in, e.g. time points of (100,130)
    ind1 = get_indices(time,window[0])[0]
    ind2 = get_indices(time,window[1])[0]
    return timeseries[ind1:ind2+1], time[ind1:ind2+1]

def end_sentence(textls):
    #for those sentences that did not end with any punctuation, add a '.' to the end and return the text list
    a = clean_list(textls)
    for i in range(len(a)):
        item = a[i]
        if a[i][-1]!='.' and a[i][-1]!='?':
            a[i]=item+'.'
    return a

def clean_quotes(input_str,replace_item='\'',replaced_by='‘’',use=1):
    #this is to replace the quote mark of ' or " or '\'' as ’
    if type(input_str)==list:
        out = []
        for item in input_str:
            out.append(item.replace(replace_item,replaced_by[use]))
    else:
        out= input_str.replace(replace_item,replaced_by[use])
        return out

def clean_text(text,clean_up='Default'):
    if clean_up =='Default':
        clean_up = ['\n', '*', '\t', '~','@','\xa0',]
    for item in clean_up:
        if item in text:
            text = text.replace(item,'')
    while text[-1] ==' ' and len(text)!=0:
        text = text[0:-1]
    while text[0] ==' ' and len(text)!=0:
        text = text[1:]
    return text

def clean_list(strlist,clean='Default'):
    if clean == 'Default':
        to_rm = ['',' ','"','“','”']
    else:
        to_rm = clean
    for i in to_rm:
        while i in strlist:
            strlist.remove(i)
    clean_list = []
    for item in strlist:
        temp = clean_text(item)
        if len(temp)!=0:
            clean_list.append(temp)
    return clean_list

def clean_textls(inlist,clean_item='Default',clean_ls='Default'):
    outlist1 = clean_list(inlist,clean_ls)
    outlist = [clean_text(i,clean_up=clean_item) for i in outlist1]
    return outlist

def clean_consec3(inlist,start_ind=2):
    #psiturk recording sometimes repeated records the scene label right before (rarely after) a choice point 
    #route is individual subject's route list
    #start_ind = 2: start to check from the 3rd element
    #check if there is repeat within consecutive 3 elements
    repeat = []
    route = inlist
    #get the indices for repeated elements within 3 elements
    for i in range(2,len(route)):
        current_elem = route[i]
        prior_elem = route[i-1]
        prior_2elem = route[i-2]
        if current_elem == prior_elem:
            repeat.append(i)
        if current_elem == prior_2elem:
            if i not in repeat:
                repeat.append(i)
    #pop those repeated elements
    for i in sorted(repeat,reverse=True):
        route.pop(i)
    return route,repeat

def clean_add_missing(inlist):
    #psiturk recording sometimes miss records the scene label right before a choice point
    route = inlist
    insert_where = []
    insert_what = []
    for i in range(len(route)):
        scene = route[i]
        prior_scene = route[i-1]
        if scene.find('_')!=-1: #choice point                
            if prior_scene != scene.split('_')[0]:
                insert_what.append(scene.split('_')[0])
                insert_where.append(i)

    count=len(insert_what)-1
    for i in sorted(insert_where,reverse=True):    
        route.insert(i,insert_what[count])
        count-=1  
    return route

def clean_psiturk_route(inlist):
    route = inlist
    route = clean_list(route) #clean-up 1: delete empty elements
    route,repeat = clean_consec3(route) #clean-up 2: delete repeated item within neighboring 3 elem
    route = clean_add_missing(route) #clean-up3: make up for missing scene label right before a choice point
    return route

def split_num(num,thres=5):
    #this is to split 95101 into ['95','101']
    num = str(num)
    if len(num)>=thres and num[0]=='1':
        num1=num[0:3]
        rest=num[3:]
    elif len(num)>=thres and int(num[0])>1:
        num1=num[0:2]
        rest=num[2:]
    return [num1,rest]
    
def split_num_ls(num_ls,thres=4):
    #this is to split [12,34,95101102103,114,115116] into [12,34,95,101,102,103,114,115,116]
    num_ls = [str(i) for i in num_ls]
    out = []
    for i in num_ls:
        cur_num =i
        while len(cur_num)>=thres:
            num1,cur_num = split_num(cur_num,thres=thres)
            out.append(num1)
        out.append(cur_num)
    return out

def split_num_ls_remain(rcl,thres=6):
    #this is the same function with the one named recall_seg_split_num
    #rcl: this is subjects' recalled event number in temporal sequence, e.g. rcl = ['1.0', '2.0', '1,2,4', '18,19,21,22', '10,14', '14.0', '18,19,20,21,22,30,31', '38,39,40', '41.0', '45,46,47,48', '49.0', '52.0', '57,58', '54,55,56,57', '66.0', '66.0', '72.0', '85,86', '82.0', '59,60', '61,64', '61,64', '87,88', '96.0', '97,98,99', '100,101', '98,99,100', '109.0', '119,120,121', '129,130']
    #this function detects potential sticky wrong number e.g. '116117.0' in the list, and cut it into '116,117'
    ind = 0
    rcl_out = rcl.copy()
    for i in rcl_out:
        if i.find(',')==-1 and len(i)>=thres:
            i2 = i[:-2] #drop the .0 at the end of the number
            cutted = split_num(i2,thres=thres-2)
            cutted_store = []
            while len(cutted[-1])>=thres-2:
                cutted_store.append(cutted[:-1])
                cutted = split_num(cutted[-1],thres=thres-2)
            cutted_store.append(cutted)
            new = ','.join(merge_lists_in_list(cutted_store))
            rcl_out[ind]=new
        ind+=1
    return rcl_out

def pad_int(in_int,dig=3,pad='0'):
    #pad integer with 0 before them to reach targeted digit
    #e.g. in_int = 1, out= '001'
    out = []
    if type(in_int)==list:
        for num in in_int:
            num = str(num)
            pad_dig = dig-len(num)
            out.append('0'*pad_dig+num)
    else:
        num = str(in_int)
        pad_dig = dig-len(num)
        out = '0'*pad_dig+num
    return out

def insert_str(file_list,infront_of_where,str_to_insert):
    #example use: insert_str(['a-sub1.xlsx','a-sub2.xlsx'],'-sub','_inserted')
    #this will return: ['a_inserted-sub1.xlsx','a_inserted-sub2.xlsx']
    inserted_list=[]
    for file in file_list:    
        ind = file.find(infront_of_where)
        name = file[0:ind]+str_to_insert+file[ind:]
        inserted_list.append(name)
    return inserted_list

def replace_str(file_list,str_to_delete, str_to_replace=''):
    #example use: delete_str(['a_inserted-sub1.xlsx','a_inserted-sub2.xlsx'],'_inserted')
    #this will return: ['a-sub1.xlsx','a-sub2.xlsx']
    replaced_list=[]
    for file in file_list:
        if str_to_delete.find('*')!=-1:
            ind = file.find(str_to_delete.split('*')[0])
            file = file[0:ind]+str_to_replace
        else:
            file = file.replace(str_to_delete,str_to_replace)
        replaced_list.append(file)
    return replaced_list

def replace_item(inlist,to_be_replaced,new_item):
    #replace all the element in the list to_be_replaced with a new item
    #inlist can be a list of list, where each list may have the item to be replaced
    if type(inlist)==list:
        outls = [i for i in inlist]
        for i,ls in enumerate(inlist):
            for j,item in enumerate(ls):
                if item==to_be_replaced:
                    outls[i][j]=new_item
    else:
        for i,item in enumerate(inlist):
            if item==to_be_replaced:
                inlist[i]=new_item
        outls = inlist
    return outls

def rename_files(in_path,file_list,rename_list):
    import os
    for ind in range(0,len(file_list)):
        file = in_path+file_list[ind]
        new = in_path+rename_list[ind]
        os.rename(file,new)
    return "Files renamed at the same directory."

def file_size(file,unit='kb'):
    import os
    size = os.path.getsize(file)/1000
    if unit =='mb':
        size= size/1000
    elif unit=='gb':
        size= size/1000000
    return size

def file_exist(file):
    import os
    if os.path.exists(file):
        return True
    else:
        return False

def os_buf_mv(from_file,to_file,check_exist=True,mute=False,wait=True,buffer=1):
    import os
    import time
    if not os.path.exists(from_file):
        print('Original file does not exist! Skipped.')
    elif check_exist==True and os.path.exists(to_file) and file_size(from_file)==file_size(to_file):
        print('Target file already exists! Skipped.')
    else:
        if mute==False:
            print(from_file, to_file)
        to_dir = '/'.join(to_file.split('/')[:-1])+'/'
        if not os.path.exists(to_dir):
            mkdir(to_dir)
        os.popen('mv '+from_file+ ' ' +to_file,'r',buffer)
        if wait==True:
            while file_exist(to_file)!=True:
                time.sleep(buffer)
            while file_size(to_file)!=file_size(from_file):
                time.sleep(buffer)
        print('done')

def shutil_cp(from_file,to_file,check_exist=False):
    import shutil
    if check_exist==False:
        shutil.copyfile(from_file, to_file)
        print('done')
    else:
        if file_exist(to_file):
            if file_size(from_file)==file_size(to_file):
                print('Target file already exist! Skipped.')
        else:
            shutil.copyfile(from_file, to_file)
            print('done')

def os_sys_cp(from_file,to_file,check_exist=False):
    import os
    if check_exist==False:
        os.system("cp "+rawf+" "+targf)
        print('done')
    else:
        if file_exist(to_file):
            if file_size(from_file)==file_size(to_file):
                print('Target file already exist! Skipped.')
        else:
            os.system("cp "+rawf+" "+targf)
            print('done')   

def os_buf_cp(from_file,to_file,check_exist=True,mute=False,wait=True,buffer=1):
    import os
    import time
    if not os.path.exists(from_file):
        print('Original file does not exist! Skipped.')
    elif check_exist==True and os.path.exists(to_file) and file_size(from_file)==file_size(to_file):
        print('Target file already exists! Skipped.')
    else:
        if mute==False:
            print(from_file, to_file)
        to_dir = '/'.join(to_file.split('/')[:-1])
        if not os.path.exists(to_dir):
            mkdir(to_dir)
        os.popen('cp '+from_file+ ' ' +to_file,'r',buffer)
        if wait==True:
            while file_exist(to_file)!=True:
                time.sleep(buffer)
            while file_size(to_file)!=file_size(from_file):
                time.sleep(buffer)
        print('done')


def get_legacy_val(inlist,heirs=['',' ']):
    #e.g. inlist = ['A','','','',' ','','B','','','C','','','']
    #this function will repeat the previous none-empty item until a new value shows up
    outls = []
    for i in inlist:
        if i not in heirs:
            temp = i
        outls.append(temp)
    return outls

def get_intersect(list1,list2):
    shared = []
    for i in list1:
        if i in list2:
            shared.append(i)
    return shared

def get_intersect_all(lists):
    process = get_intersect(lists[0],lists[1])
    for i in range(2,len(lists)):
        process = get_intersect(process,lists[i])
    return process

def get_interval(inlist):
    #this is to get the integer intervals (between the neighboring two numbers) from a list (e.g. get TRs for a sentence segment with an onset=1.3TR, offset=4TR: should return [2,3]TRs)
    #e.g. list = [1.3, 4, 7.1, 10], this function takes the ceiling of the onset, and floor of the offset (if float) or the offset-1 (if int), and returns the interval in between each pair of onset and offset
    #so out = [[2,3], [4,5,6], [8,9], [10]] the last element is going to be itself.
    import numpy as np
    out = []
    a = inlist
    for i in range(len(a)):
        #get the range of integers for each interval
        if i!=len(a)-1:
            start = int(np.ceil(a[i]))
            if a[i+1].is_integer():
                end = int(a[i+1]-1)
            else:
                end = int(np.floor(a[i+1]))
            rng = [i for i in range(start,end+1)] #include the end TR
            if rng ==[]:
                rng = [int(float(np.ceil(a[i])))]
        else: #if the last element: is the start TR itself
            start = int(np.ceil(a[i]))
            rng = [start]
        #determine if there's negative numbers, which was error 
        if check_list_thr(rng)=='yes':
            out.append('nan')
        else:
            out.append(rng)
    return out

def get_ls_interval(input_lists, start='default', end='default',skip=(0,0)):
    '''gets the interval lists with the specified start and end
        example: 
            input_lists = [[5, 6, 7], [13, 14, 15]]
            start = 1
            end = 16
            output = [[1, 2, 3, 4], [8, 9, 10, 11, 12], [16]]
            '''
    # Flatten the input lists to find min and max values if start or end is 'default'
    all_values = sorted([item for sublist in input_lists for item in sublist])
    # Set start and end if they are 'default'
    if start == 'default':
        start = min(all_values)
    if end == 'default':
        end = max(all_values)
    # Initialize output list and previous end
    output = []
    previous_end = start - 1
    # Loop through each sublist
    for sublist in input_lists:
        # Find the start and end of this sublist
        sublist_start = min(sublist)
        sublist_end = max(sublist)
        # If there is a gap between the previous end and this sublist's start, add the padding
        if previous_end + 1 < sublist_start:
            output.append(list(range(previous_end + 1+skip[0], sublist_start-skip[-1])))
        # Update the previous end to this sublist's end
        previous_end = sublist_end
    # Check for padding after the last sublist until 'end'
    if previous_end + 1 <= end:
        output.append(list(range(previous_end + 1+skip[0], end + 1-skip[-1])))
    return output

def get_union(list1,list2,uniq=False):
    ls = merge_lists_in_list([list1,list2])
    if uniq!=False:
        ls,rep = get_unique_repeat(ls)
    return sorted(ls)

def get_subtract(list1,list2):
    #this is to get the unique items only from list1, where list2 serves as the reference list
    #e.g. list1 = [1,2,3], list2=[2,3,4], this should return 1
    out = []
    for i in list1:
        if i not in list2:
             out.append(i)
    return out

def get_unique_repeat(inlist,mute=True):
    #this is to get the list of elements that are unique, and the list of items that has been repeated in the inlist
    unique = []
    repeat = []
    for i in range(len(inlist)):
        if inlist[i] not in unique:
            unique.append(inlist[i])
        else:
            repeat.append(inlist[i])
            if mute!=True:
                print('Repeated element is: '+str(inlist[i]))
    return unique,repeat

def get_rep_thr(inlist,minmax=(None,None)):
    #get repeated items in the inlist that were repeated over a certain threshold
    #e.g. inlist = [1,1, 2,2,2,2, 3,3, 6, 7,7,7], min=4, max=None, should return [2]; if min=3, should return [2,7]
    dic = init_dict([str(i) for i in inlist])
    for i in inlist:
        dic[str(i)].append(i)
    qualified_list = []
    number_of_repeats = []
    for key in dic:
        nrep = len(dic[key])
        if minmax[0]!=None and minmax[1]==None and nrep>=minmax[0]:
            qualified_list.append(key)
            number_of_repeats.append(nrep)
        elif minmax[0]==None and minmax[1]!=None and nrep<=minmax[1]:
            qualified_list.append(key)
            number_of_repeats.append(nrep)
        elif minmax[0]!=None and minmax[1]!=None and nrep>=minmax[0] and nrep<=minmax[1]:
            qualified_list.append(key)
            number_of_repeats.append(nrep)
    return qualified_list,number_of_repeats

def get_first_elem_ind(inlist,rng=[1,30],repeat_ind=False):
    #this is to get the first element and its index in each element-strings in a list, i.e. ['1,2,34','4,1,51'] 
    #example: get_first_elem_ind(inlist,rng=[1,100],repeat_ind=False)
            #inlist = ['93.0', '93.0', '93.0', '92, 93', '93, 15, 94, 28', '15, 93', '93, 9', '28, 96']
            #output: elem = [93,92,15,9,28], ind = [0,3,4,6,7]
    #example: get_first_elem_ind(inlist,rng=[1,100],repeat_ind=True)
            #inlist = ['93.0', '93.0', '93.0', '92, 93', '93, 15, 94, 28', '15, 93', '93, 9', '28, 96']
            #output: elem = [93,92,15,94,28,9,96], ind = [0,3,4,4,4,6,7]
    out_elems=[]
    out_inds=[]    
    count_ind = 0
    for i in inlist:
        elems= array_str_to_num(clean_list(i.split(',')),True,True)
        if repeat_ind==False: #the same ind do not count for multiple elems
            outlen = len(out_elems)
            for j in elems:
                if j not in out_elems and len(out_elems)==outlen and j in range(rng[0],rng[1]+1):
                    out_elems.append(j)
                    out_inds.append(count_ind)
            count_ind+=1
        else: #count all unmentioned elems in the same ind sequentially
            for j in elems:
                if j not in out_elems and j in range(rng[0],rng[1]+1):
                    out_elems.append(j)
                    out_inds.append(count_ind)
            count_ind+=1
    return out_elems,out_inds

def get_indices(search_list,search_item,find=False):
    #this is to get indices of the target item(s) from a list
    #example use: ind_list = get_indices(search_list = data['wanted_not'], search_item = ['1','0'])
    #if find=True, will return the inds of items in the list that contains the search item.
    counter=0
    ind_list = []
    for item in search_list:
        if type(search_item)==list:
            if find ==False:
                if item in search_item:
                    ind_list.append(counter)
            else:
                for search in search_item: #e.g. search_item = ['_','!']
                    if search in str(item): #if '_' is in the current item in the search list
                        ind_list.append(counter)
        else:
            if find ==False:
                if item == search_item:
                    ind_list.append(counter)
            else:
                if search_item in str(item): #if '_' is in the current item in the search list
                    ind_list.append(counter)
        counter+=1
    return ind_list

def get_inds_lsofls(list_of_lists,search_item):
    # Finding indices of sublists that contain the search_item
    indices = [i for i, sublist in enumerate(list_of_lists) if search_item in sublist]
    return indices

def nth_percentile(vector, n=50):
    #e.g. n=50 is getting the median of the vector
    if not 0 <= n <= 100:
        raise ValueError("n should be between 0 and 100")
    percentile = np.percentile(vector, n)
    return percentile

def inds_above_thresh(vector, thresh):
    #get the inds of the elements in the vector whose value is above the threshold
    inds = [index for index, value in enumerate(vector) if value > thresh]
    return inds

def inds_ls_above_thresh(ndarray,nth_prcnt=50):
    #for each array in the ndarray, get the inds of the elements in the vector whose value is above the threshold
    inds = [inds_above_thresh(vec,thresh=nth_percentile(vec,n=nth_prcnt)) for vec in ndarray]
    return inds

def translate_pairs(tobe_translated_list,decipher_pairing):
    #example: to-be-translated e.g. [27,34,29]; decipher_pairing e.g. [27,28,29,30,31,32,33,34] -> [1,2,3,4,5,6,7,8]; so returns [1,8,3]
    #to-be_translated is a list of values to be deciphered
    #decipher_pairing is a list of two lists of equal length (1-to-1 projection), the frist ls is the observed number, the second is the actual value we want
    obs,seq= decipher_pairing
    inds = [get_indices(obs,i)[0] for i in tobe_translated_list]
    out = apply_inds(inds,seq)
    return out

def get_sequence(seq1,seq2,offset=1):
    #this function is to get the sequence of items from seq2 with reference to seq1
        #seq1 is the first sequence of items that serve as a reference
        #seq2 is the same list of items but in a different sequence
        #offset is if the sequence of the reference should start from 0 or 1
    if sorted(seq1)==sorted(seq2):
        seq =[]
        for i in seq2:
            ind = get_indices(seq1,i)[0]
            seq.append(ind+offset)
        return seq
    else:
        warning='Error: seq1 and seq2 should contain the exact same elements'
        print(warning)

def get_min_max(y_axis,roundoff=True):
    import numpy as np
    if type(y_axis[0])==float or type(y_axis[0])==int or type(y_axis[0])==np.float64 or type(y_axis[0])==np.float32:
        high = max(y_axis)
        low = min(y_axis)
    elif type(y_axis[0])==list:
        high = max([max(lst) for lst in y_con])
        low = min([min(lst) for lst in y_con])
    if roundoff ==True:
        high= np.ceil(high)
        low= np.floor(low)
    return high,low

def get_range(start,end,step):
    #this returns a list of integers at the specified start, end, step-size
    import numpy as np
    out = np.arange(start, end+1, step).tolist()
    return out

def get_range_float(rng,points=100):
    #rng is the boundaries: e.g. (0,10)
    #points is the density of the line
    #this function returns a list of floats e.g. [0, 0.1, 0.2,...,9.8, 9.9, 10]
    import numpy as np
    ind1 = int(float(np.ceil(rng[0]*points)))
    ind2 = int(float(np.floor(rng[1]*points)))
    step = int((ind1-ind2)/points)
    float_rng = []
    for i in range(ind2,ind1+1,step):
        float_rng.append(i/points)
    return float_rng




def distance_mat(ndarray):
    from scipy.spatial.distance import pdist, squareform
    vector = np.array(ndarray)
    if len(shape(vector))==1:
        vector = vector[:, np.newaxis]
    # Compute the pairwise Euclidean distance
    pairwise_dist = pdist(vector)
    # Convert the pairwise distance to a square matrix form
    dm = squareform(pairwise_dist)
    return dm

def get_sim_ratio(array1,array2,skip_nan_min=5,num_count=False):
    #skip_nan_min is to limit the length of array after excluding NA to be no less than 5
    import numpy as np
    array1 = np.array(array1) 
    array2 = np.array(array2)     
    if len(array1)!=len(array2):
        print('The two arrays need to be of same length')
    else:
        ls1 = [str(i) for i in array1]
        ls2 = [str(i) for i in array2]
        inds1 = get_indices(ls1,'nan')
        inds2 = get_indices(ls2,'nan')
        inds = get_union(inds1,inds2,uniq=True)
        if len(array1)-len(inds)>=skip_nan_min:
            array1 = list_exclude_inds(array1,inds)
            array2 = list_exclude_inds(array2,inds)
        #count same elements
        count= 0
        for i in range(len(array1)):
            if array1[i]==array2[i]:
                count+=1
        ratio= count/len(array1)
        if num_count==False:
            return ratio
        else:
            return count

def get_pairsim_ratio(ndarray1,ndarray2,thresh_n=5,fill='nan',num_count=False):
    import numpy as np
    ratio_mat = init_mat((len(ndarray1),len(ndarray2)),fill)
    a1 = 0
    for i in ndarray1:
        a2 = 0
        for j in ndarray2:       
            if 'nan' in i or 'nan' in j or np.nan in i or np.nan in j:
                ratio = np.nan
            else:
                ratio = get_sim_ratio(i,j,skip_nan_min=thresh_n,num_count=num_count)
            ratio_mat[a1][a2]= ratio
            a2+=1
        a1+=1
    return ratio_mat

def get_corr(array1,array2,skip_nan_min=30,na_replace=False,decimal=3):
    #skip_nan_min is setting the lower boundary for the array length after excluding the nan in both arrays
    #skip_nan_min is to limit the length of array after excluding NA to be no less than 30
    import scipy.stats
    import numpy as np
    array1 = np.array(array1) 
    array2 = np.array(array2) 
    if len(array1)!=len(array2):
        print('The two arrays need to be of same length')
    else:
        ls1 = [str(i) for i in array1]
        ls2 = [str(i) for i in array2]
        inds1 = get_indices(ls1,'nan')
        inds2 = get_indices(ls2,'nan')
        inds = get_union(inds1,inds2,uniq=True)
        if len(array1)-len(inds)>=skip_nan_min:
            array1 = list_exclude_inds(array1,inds)
            array2 = list_exclude_inds(array2,inds)
            pearson_r= scipy.stats.pearsonr(array1,array2)
            corr=np.round(pearson_r[0],decimal)
            pval=np.round(pearson_r[1],decimal)
        else:
            if na_replace==False:
                print('The two arrays are shorter than the minimum threshold after excluding the NaNs in them!')
            else:
                corr = na_replace[0]
                pval = na_replace[1]
        return corr,pval

def get_paircorr(ndarray1,ndarray2,fill='nan',thresh_n=30):
    #outputs rmat and pmat based on pairwise pearson correlation
    import numpy as np
    rmat = init_mat((len(ndarray1),len(ndarray2)),fill)
    pmat = init_mat((len(ndarray1),len(ndarray2)),fill)
    a1 = 0
    for i in ndarray1:
        a2 = 0
        for j in ndarray2:       
            if 'nan' in i or 'nan' in j or np.nan in i or np.nan in j:
                rval = np.nan
                pval = np.nan
            else:
                rval,pval = get_corr(i,j,skip_nan_min=thresh_n)
            rmat[a1][a2]= rval
            pmat[a1][a2]= pval
            a2+=1
        a1+=1
    return rmat,pmat

def cosine_sim(ndarray1,ndarray2):
    if len(shape(ndarray1))==1 and len(shape(ndarray2))==1:
        out = get_cosine_sim(ndarray1,ndarray2)
    else:
        out = get_pairwise_cosinesim(ndarray1,ndarray2)
    return out

def get_cosine_sim(array1,array2):
    import numpy as np
    from numpy.linalg import norm
    A,B = array1,array2
    # compute cosine similarity
    cosine = np.dot(A,B)/(norm(A)*norm(B))
    return cosine 

def get_pairwise_cosinesim(ndarray1,ndarray2):
    #this is the same function from the next one, but with a better name
    #outputs pairwise cosine similarity (correlation) matrix
    #example: ndarray1 = [ls1,ls2,ls3],ndarray2 = [ls4,ls5,ls6], out_mat = [r(ls1,ls4),r(ls1,ls5),r(ls1,ls6); r(ls2,ls4),r(ls2,ls5),r(ls2,ls6); r(ls3,ls4),r(ls3,ls5),r(ls3,ls6)]
    import numpy as np
    import sklearn    
    from sklearn import metrics
    sim_mat = sklearn.metrics.pairwise.cosine_similarity(ndarray1,ndarray2,dense_output=True)
    print(sim_mat.shape)
    return sim_mat

def pairwise_col_correlation(matrix1, matrix2):
    """
    Compute the correlation between corresponding columns of two matrices efficiently.

    Args:
    matrix1 (numpy.ndarray): The first matrix of shape (n, m).
    matrix2 (numpy.ndarray): The second matrix of shape (n, m).

    Returns:
    numpy.ndarray: An array of correlation coefficients of length m.
    """
    # Check that the matrices have the same shape
    if shape(matrix1) != shape(matrix2):
        raise ValueError("Both matrices must have the same shape")
    # Center the matrices (subtract the mean of each column)
    matrix1_centered = matrix1 - np.mean(matrix1, axis=0)
    matrix2_centered = matrix2 - np.mean(matrix2, axis=0)
    # Compute the covariance of corresponding columns
    covariance = np.sum(matrix1_centered * matrix2_centered, axis=0)
    # Compute the standard deviations of each column
    std1 = np.sqrt(np.sum(matrix1_centered**2, axis=0))
    std2 = np.sqrt(np.sum(matrix2_centered**2, axis=0))
    # Compute the correlation for each column
    correlations = covariance / (std1 * std2)
    return correlations

def pairwise_row_correlation(ndarray1, ndarray2):
    '''fast computes the pair-wise pearson correlation between the rows of arr1 and arr2
        ignores nan in the inputs'''
    arr1, arr2 = ndarray1, ndarray2
    # Convert arrays to masked arrays
    arr1_m = np.ma.masked_invalid(arr1)
    arr2_m = np.ma.masked_invalid(arr2)
    # Calculate means of rows for both arrays ignoring NaN
    mean1 = np.ma.mean(arr1_m, axis=1, keepdims=True)
    mean2 = np.ma.mean(arr2_m, axis=1, keepdims=True)
    # Demean the rows
    demeaned1 = arr1_m - mean1
    demeaned2 = arr2_m - mean2
    # Use masked dot product. Filled NaNs with 0 for dot product
    numerators = np.ma.dot(demeaned1.filled(0), demeaned2.filled(0).T)
    # Calculate L2 norm for rows of both arrays ignoring NaN
    norm1 = np.sqrt(np.ma.sum(demeaned1 ** 2, axis=1, keepdims=True))
    norm2 = np.sqrt(np.ma.sum(demeaned2 ** 2, axis=1, keepdims=True))
    # Calculate denominators for each pair
    denominators = np.dot(norm1, norm2.T)
    # Calculate pairwise correlation, taking care of division by zero
    with np.errstate(divide='ignore', invalid='ignore'):
        result = numerators / denominators
    # Handle potential NaNs and infinite values due to division by zero
    result = np.nan_to_num(result, nan=0.0, posinf=0.0, neginf=0.0)
    rmat = result.data
    return rmat

def get_pairwise_pearsonr(ndarray,ndarray2='None'):
    #this returns the pairwise pearson r matrix, ignores nan in the data
    import numpy as np
    if ndarray2=='None':
        ndarray2 = ndarray.copy()
    sim_mat = pairwise_row_correlation(ndarray,ndarray2)
    return sim_mat

def get_pairwise_corr(ndarray1,ndarray2):
    #outputs pairwise cosine similarity (correlation) matrix
    #example: ndarray1 = [ls1,ls2,ls3],ndarray2 = [ls4,ls5,ls6], out_mat = [r(ls1,ls4),r(ls1,ls5),r(ls1,ls6); r(ls2,ls4),r(ls2,ls5),r(ls2,ls6); r(ls3,ls4),r(ls3,ls5),r(ls3,ls6)]
    import numpy as np
    import sklearn    
    from sklearn import metrics
    sim_mat = sklearn.metrics.pairwise.cosine_similarity(ndarray1,ndarray2,dense_output=True)
    print(sim_mat.shape)
    return sim_mat

def paircorr_list(ndarray1,ndarray2):
    #outputs the list of rval, each derived from the correspongding elements from ndarray1 and ndarray2
    #example: ndarray1 = [ls1,ls2,ls3],ndarray2 = [ls4,ls5,ls6], out_ls = [r(ls1,ls4),r(ls2,ls5),r(ls3,ls6)]
    vecs1,vecs2 = ndarray1,ndarray2
    if len(vecs1)!=len(vecs2):
        print('the two vectors array need to be of the same length!')
    else:
        allr = []
        allp = []
        for i in range(len(vecs1)):
            r,p = get_corr(vecs1[i],vecs2[i])
            allr.append(r)
            allp.append(p)
    return allr,allp    

def paircorr_avg(ndarray1,ndarray2,mute=True):
    allr,allp = paircorr_list(ndarray1,ndarray2)
    r = avg(allr,mute=mute)
    p = avg(allp,mute=mute)
    return r,p

def compute_pairs(list1, list2, compute_func, lists_name=None, default_params=None):
    """
    Call a function 'compute_func' on each combination of sub-lists (matrices) from two lists of lists efficiently.
    You can override the default parameters of the function if provided.

    Args:
    list1 (list of numpy.ndarray): The first list containing matrices.
    list2 (list of numpy.ndarray): The second list containing matrices.
    compute_func (function): The function to apply on each pair of matrices.
    lists_name (list of list of str, optional): Lists to generate names for combinations.
    default_params (dict, optional): A dictionary of default parameters for the compute function.

    Returns:
    tuple: A list of results from applying the compute function on each combination of matrices, and a list of combination names.
    """
    from functools import partial
    from itertools import product
    if default_params is not None:
        # Create a new function with default parameters
        compute_func = partial(compute_func, **default_params)
    # Using list comprehension for faster iteration and processing
    results = [compute_func(matrix1, matrix2) for matrix1 in list1 for matrix2 in list2]
    # Generate combination names
    if lists_name:
        names = combo_ls(lists_name)
    else:
        names = combo_ls([[str(i+1) for i in range(len(list1))], [str(i+1) for i in range(len(list2))]])
    return results, names

def get_pairs(list_of_lists,lists_name='list',connect='_',direction=1):
    #direction = 1 or 2
    #returns the name of all possible pairs and index
    if type(lists_name)==str:
        lists_name=[lists_name+str(i+1) for i in range(len(list_of_lists))]
    
    #get all possible pairs of vars
    all_pairs = []
    for x1_ind in range(0,len(list_of_lists)):
        for x2_ind in range(0,len(list_of_lists)):
            if x2_ind==x1_ind:
                continue
            pair1 = (x1_ind,x2_ind)
            pair2 = (x2_ind,x1_ind)
            if direction==1:
                if pair1 not in all_pairs and pair2 not in all_pairs:
                    all_pairs.append(pair1)
            else:
                all_pairs.append(pair1)
    #examine each of these pairs
    pair_num=0
    all_pairs_name =[]
    while pair_num<len(all_pairs):            
        #get the current pair of vars
        x1_ind = all_pairs[pair_num][0]
        x2_ind = all_pairs[pair_num][1]
        x1 = list_of_lists[x1_ind]
        x2 = list_of_lists[x2_ind]
        x1_name = str(lists_name[x1_ind])
        x2_name = str(lists_name[x2_ind])
        this_pair = x1_name+connect+x2_name
        all_pairs_name.append(this_pair)
        pair_num+=1
    return all_pairs_name,all_pairs

def get_triu_mat(mat):
    return np.triu(mat)

def get_tril_mat(mat):
    return np.tril(mat)

def get_triu_coords(list_of_lists,lists_name='list',diag_offset=1):
    #this is to get the non-directional pairs (upper triangle of a symmetric matrix)
    #nondir: non-directional when considering the pairs (i.e. irrespective of sequence)
    import numpy as np
    #check input list of lists type: int or actual list of lists
    if type(list_of_lists)==int:
        length=list_of_lists
    elif type(list_of_lists)==list:
        length = len(list_of_lists)
    else:
        length = len(list_of_lists)
    #name of the lists in sequence
    if type(lists_name)==str:
        name_list = init_list(length=length,label=lists_name)
    else:
        name_list = lists_name
    #get index pairs for upper triangle
    iu=np.triu_indices(n=length,k=diag_offset)
    x,y = iu[0],iu[1]
    pairs_ind=[]
    pairs_name=[]
    for i in range(0,len(x)):
        pair = (x[i],y[i])
        name = name_list[x[i]]+'-'+name_list[y[i]]
        pairs_ind.append(pair)
        pairs_name.append(name)        
    return pairs_ind,pairs_name

def get_triu_val(mat,diag_offset=1):
    import numpy as np
    length = len(mat)
    pairs_ind,pairs_name = get_triu_coords(length,diag_offset=diag_offset)
    upval =[]
    for pair in pairs_ind:
        upval.append(mat[pair[0]][pair[1]])
    return upval

def get_tril_coords(list_of_lists,lists_name='list',diag_offset=1):
    #this is to get the non-directional pairs (lower triangle of a symmetric matrix)
    #nondir: non-directional when considering the pairs (i.e. irrespective of sequence)
    import numpy as np
    #check input list of lists type: int or actual list of lists
    if type(list_of_lists)==int:
        length=list_of_lists
    elif type(list_of_lists)==list:
        length = len(list_of_lists)
    elif type(list_of_lists)==np.ndarray:
        length = len(list_of_lists)
    else:
        print('need to check the data-type of the input <list_of_list>!')
    #name of the lists in sequence
    if type(lists_name)==str:
        name_list = init_list(length=length,label=lists_name)
    else:
        name_list = lists_name
    #get index pairs for upper triangle
    iu=np.tril_indices(n=length,k=-1*diag_offset)
    x,y = iu[0],iu[1]
    pairs_ind=[]
    pairs_name=[]
    for i in range(0,len(x)):
        pair = (x[i],y[i])
        name = name_list[x[i]]+'-'+name_list[y[i]]
        pairs_ind.append(pair)
        pairs_name.append(name)        
    return pairs_ind,pairs_name

def get_tril_val(mat,diag_offset=1):
    import numpy as np
    length = len(mat)
    pairs_ind,pairs_name = get_tril_coords(length,diag_offset=diag_offset)
    lowval =[]
    for pair in pairs_ind:
        lowval.append(mat[pair[0]][pair[1]])
    return lowval

def get_diag_val(mat,diag_offset=0):
    import numpy as np
    out = np.diagonal(mat,offset=diag_offset)
    return out

def get_offdiag_val(mat,diag_offset=1):
    #diag_offset=1 means we exclude the 1 center diagonal; 2 means exclude the 3 center diagonals
    import numpy as np
    triu = get_triu_val(mat,diag_offset=diag_offset)
    tril = get_tril_val(mat,diag_offset=diag_offset)
    out = merge_lists_in_list([triu,tril])
    return out

def avg_below_diag(matrix, num_cells=5):
    """
    Computes the average of the specified number of cells below the diagonal for each column in a square matrix.

    Parameters:
    matrix (2D array-like): A square matrix.
    num_cells (int): The number of cells below the diagonal to consider.

    Returns:
    list of float: A list containing the average of the values below the diagonal for each column.
    """
    n = len(matrix)  # Size of the matrix (assuming it's square)
    averages = []
    for col in range(n):
        # Get the specified number of elements below the diagonal for this column
        values_below_diagonal = []
        for row in range(col + 1, min(col + 1 + num_cells, n)):
            values_below_diagonal.append(matrix[row][col])
        # Compute the average of these values, if any
        if values_below_diagonal:
            avg = sum(values_below_diagonal) / len(values_below_diagonal)
        else:
            avg = None  # If no values are below the diagonal for this column
        averages.append(avg)
    return averages

def get_non_nans(array):
    import numpy as np
    exclude = [np.nan,'nan']
    num_types = [int,float,np.float64,np.int64,np.nan,'nan']
    list_types= [np.ndarray,list]
    out = []
    if type(array[0]) in num_types:
        out = array[~np.isnan(array)]
    elif type(array[0]) in list_types:
        for ls in array:
            ls = np.array(ls)
            non_nan = ls[~np.isnan(ls)]
            uni,rep = get_unique_repeat(non_nan)
            if len(uni)==0:
                continue
            else:
                out.append(ls)
        trans = np.array(out).T
        out = []
        for ls in trans:
            ls = np.array(ls)
            non_nan = ls[~np.isnan(ls)]
            uni,rep = get_unique_repeat(non_nan)
            if len(uni)==0:
                continue
            else:
                out.append(ls)
        out = np.array(out).T
    return out

def get_nonnan_vals(array):
    import numpy as np
    vals = []
    inds = []
    ind = 0
    for i in array:
        if np.isnan(i)==False:
            vals.append(i)
            inds.append(ind)
        ind+=1
    return vals,inds

def zero_matlow(inmat,k=1):
    return np.triu(inmat,k=k) #k=0 will keep the diagonal

def zero_matup(inmat,k=1):
    return np.tril(inmat,k=k) #k=0 will keep the diagonal

def set_mat_diag(inmat,k=0,set_to=0):
    if type(k)==int:
        np.fill_diagonal(inmat[k:], set_to)
        np.fill_diagonal(inmat[:, k:], set_to)
    else: #a list or ndarray of k
        for i in k:
            inmat = set_mat_diag(inmat,i,set_to)
    return inmat

def keep_mat_diag(inmat,k=0,others_set_to=0):
    # Re-creating the original 10x10 matrix with arbitrary values
    original_matrix = inmat
    # Creating a copy of the matrix to preserve the original
    modified_matrix = original_matrix.copy()
    # Setting all elements to zero except the first off-diagonals
    # First, set the entire matrix to zero
    modified_matrix[:, :] = others_set_to
    # Then, restore the first (if k=1) off-diagonals from the original matrix
    if type(k)==int:
        k = [k]
    for i in k:
        np.fill_diagonal(modified_matrix[i:], original_matrix.diagonal(i))
        np.fill_diagonal(modified_matrix[:, i:], original_matrix.diagonal(-1*i))
    return modified_matrix

def check_list_thr(inlist, thr=0,check='neg'):
    #thr=0, check='neg' (smaller than thr): this checks if a list of numbers had an element that is < 0
    if check =='neg':
        out = ['yes' if i>=thr else 'no' for i in inlist]
    elif check=='pos':
        out = ['no' if i>thr else 'yes' for i in inlist]
    if 'no' in out:
        out = 'yes' #'The list has element of checking'
    else:
        out = 'no' #'The list does not have element of checking'
    return out

def check_corr(check_lists,lists_name,make_plot=False,title=False,corr_path=False,mute=False,thresh=30):
    import os
    if len(check_lists)!=len(lists_name):
        print('ERROR: number of lists need to match with number of names')
    else:
        #get all possible pairs of vars
        all_pairs = []
        for x1_ind in range(0,len(check_lists)):
            for x2_ind in range(0,len(check_lists)):
                if x2_ind==x1_ind:
                    continue
                pair1 = (x1_ind,x2_ind)
                pair2 = (x2_ind,x1_ind)
                if pair1 not in all_pairs and pair2 not in all_pairs:
                    all_pairs.append(pair1)
        #examine each of these pairs
        pair_num=0
        corr=[]
        all_pairs_name =[]
        while pair_num<len(all_pairs):
            corr.append([]) 
            #get the current pair of vars
            x1_ind = all_pairs[pair_num][0]
            x2_ind = all_pairs[pair_num][1]
            x1 = check_lists[x1_ind]
            x2 = check_lists[x2_ind]
            x1_name = str(lists_name[x1_ind])
            x2_name = str(lists_name[x2_ind])
            this_pair = x1_name+'_'+x2_name
            all_pairs_name.append(this_pair)
            #create a separate path for the current var pair
            pair_name = x1_name+'_vs_'+x2_name
            if corr_path!=False:
                if os.path.isdir(corr_path) ==False:
                    parentdir =  '/'+'/'.join(corr_path.split('/')[1:-2])
                    if os.path.isdir(parentdir) ==False:
                        os.mkdir(parentdir)
                    os.mkdir(corr_path)
                pair_path = corr_path+pair_name+'/'
            else:
                pair_path=False
            if not os.path.exists(pair_path):
                os.makedirs(pair_path)
            #check out the current var pair correlation:
            if title!=False:
                graph_title = title+'_'+ pair_name
            else:
                graph_title = pair_name
            corr1=plot_scatter(x1,x2,outpath=pair_path,make_plot=make_plot,title=graph_title,xlab=x1_name,ylab=x2_name,pointlab=False,mute=mute,thresh=thresh)
            #store the corr and pval of the current pair:
            corr[pair_num].append(corr1[0])
            corr[pair_num].append(corr1[1])
            #get ready for the next pair
            pair_num+=1
        if mute ==False:
            print(all_pairs_name)
    return corr,all_pairs_name

def find_peak_latency(timeseries,time,window='all',peak='high'):
    #timeseries: the y-axis that corresponds to the value of each time-point
    #time: the time points
    #window: the time window to search in, e.g. (100,130)
    #peak: defaults to 'high', but can also specify 'low'
    import numpy as np
    if len(time)!=len(timeseries):
        print('input time series needs to be the same length as time points')
    else:
        if window !='all':
            series,time = trunc(timeseries,time,window=window)
        #get peak    
        if peak == 'high':
            peak = max(series)
        else:
            peak = min(series)
        #get ind of peak
        ind = get_indices(series,peak)
        #get time of peak from ind
        if len(ind)==1:
            print(time[ind[0]])
            return time[ind[0]],timeseries[ind[0]]
        else: #multiple peaks
            uni,rep = get_unique_repeat([timeseries[i] for i in ind])
            if len(uni)==1:
                return [time[i] for i in ind], uni[0]
            else:
                return [time[i] for i in ind], [timeseries[i] for i in ind] 

def find_halfmax_point(lagcorr_vec,half_max=0.5,interpolate=False,rnd=3):
    """
    Finds the exact point where the correlation values first fall to or below (<=) half of its maximum (0.5) through interpolation.
    
    Parameters:
    correlation_values (list of float): The autocorrelation values at different lags.
    
    Returns:
    float: The exact point between indices where the correlation value falls to 0.5.
           Returns -1 if no such point is found.
    """
    correlation_values = lagcorr_vec
    half_maximum = half_max
    # Iterate over the list of correlation values to find where it falls to half-maximum
    if interpolate==False:
        for i, value in enumerate(correlation_values):
            if value <= half_maximum:
                return round(i,3)
    # Iterate over the list of correlation values to find the points between which to interpolate
    else:
        for i in range(1, len(correlation_values)):
            if correlation_values[i-1] >= 0.5 and correlation_values[i] <= 0.5:
                # Perform linear interpolation
                # y = y0 + (y1 - y0) * ((x - x0) / (x1 - x0))
                # Here, x represents the index, and y represents the correlation value
                # We solve for x when y = 0.5
                x0, y0 = i-1, correlation_values[i-1]
                x1, y1 = i, correlation_values[i]
                exact_point = x0 + (0.5 - y0) * ((x1 - x0) / (y1 - y0))
                return round(exact_point,rnd)

def mat_halfmax(mat,half_max=0.5,interpolate=True,rnd=4,window=10,plot=True,to_inds='None',plot_title='Auto-correlation_half-maximum',outpath=False):
    #for TTC mat, find diagonal width using halfmax TR
    out_vec= [find_halfmax_point(mat[i][i:i+window],half_max=half_max,interpolate=interpolate,rnd=rnd) for i in range(len(mat)-window)]
    ttc_vecs = [mat[i][i:i+window] for i in range(len(mat)-window)]
    if to_inds!='None':
        out_vec = apply_inds(to_inds,out_vec)
        ttc_vecs= apply_inds(to_inds,ttc_vecs)
    if plot==True:
        exclude_inds = get_indices(out_vec,None)
        ttc_vecs = list_exclude_inds(ttc_vecs,exclude_inds)        
        ttc_avg_vec = avg(ttc_vecs,mute=True)
        a = avg(list_exclude_inds(out_vec,exclude_inds),mute=True)
        plot_autocorr_halfmax(ttc_avg_vec,a,title=plot_title,outpath=outpath)
    return out_vec,exclude_inds

def apply_inds(indlist,inlist,ind_pos=0):
    outlist = []
    for ind in indlist:
        if ind_pos == 0:
            outlist.append(inlist[ind])
        if ind_pos == 1:
            outlist.append(inlist[:,ind])
        if ind_pos == 2:
            outlist.append(inlist[:,:,ind])
        if ind_pos == 3:
            outlist.append(inlist[:,:,:,ind])
    return outlist

def apply_inds_2d(coordlist,inmat):
    outmat = init_mat(inmat.shape)
    vals = []
    for coord in coordlist:
        outmat[coord[0]][coord[1]]=inmat[coord[0]][coord[1]]
        vals.append(inmat[coord[0]][coord[1]])
    return outmat,vals

def assign_val(inds,inlist,val=0):
    #re-assigns the value to inlist at inds
    outls = [i for i in inlist]
    if type(val)==list:
        if len(inds)==len(val):
            for i in range(len(inds)):
                outls[inds[i]]=val[i]
        else:
            print('Error: the length of inds need to be the same as val!')        
    else:
        for i in inds:
            outls[i]=val
    return outls

def translate_items_inls(inlist,search='',set_search_vals=0, set_other_vals=1):
    outls = init_mat(shape=(1,len(inlist)),fill=set_other_vals)[0]
    inds = get_indices(inlist,search)
    outls = assign_val(inds,outls,val=set_search_vals)
    return outls

def get_vals(ls1,ls2,forls2=True):
    #list1: a list of items (or lists, if forls2=False, then get values for each element in list1)
    #list2: a list of items or lists
    #aim: for each unique element in list2, get the corresponding values in list1
    if forls2 ==True: #get vals for each element in ls2
        vals = {}
        if type(ls2[0])==list:
            uni,rep = get_unique_repeat(merge_lists_in_list(ls2))
            for i in sorted(uni):
                vals[i]=[]
                count=0
                for ls in ls2:
                    if i in ls:
                        vals[i].append(ls1[count])
                    count+=1
        else:
            uni,rep = get_unique_repeat(ls2)
            for i in sorted(uni):
                vals[i]=[]
                inds = get_indices(ls2,i)
                vals[i]= apply_inds(inds,ls1)
    else:#only if list1 and list2 are both list of lists
        if type(ls2[0])==list and type(ls1[0])==list: #get values for each element in list1
            uni,rep = get_unique_repeat(merge_lists_in_list(ls1))
            vals = {}
            for i in sorted(uni):
                vals[i]=[]
                count=0
                for ls in ls1:
                    if i in ls:
                        vals[i].append(ls2[count])
                    count+=1
                vals[i]=merge_lists_in_list(vals[i])
        else:
            vals = 'List1 and list2 are not both list of lists!'
    return vals

def get_cp_items(ls1, ls1_items, ls2,find=False):
    #get counterpart items of list2 based on the items from list1
    #e.g. list1 = [1,2,8,2,4,1,2,10], ls1_item = [2,10], list2 = [1,2,3,4,5,6,7,8], returns ls2_item = [2,4,7,8] 
    ls2_item = []
    inds = get_indices(ls1,ls1_items,find=find)
    ls2_item = apply_inds(inds,ls2)
    # if len(ls2_item)==1:
    #     ls2_item=ls2_item[0]
    return ls2_item
        
def get_keydata(partial_key,data,partial=True):
    #returns the data of that key and the full name of that key
    #e.g. keyd,key = xf.get_keydata(str(sub+'_'),rcl)
    out_key = []
    for key in data:
        if partial!=True:
            if partial_key==key:
                out_key.append(key)
        else:
            if partial_key in key:
                out_key.append(key)
    if len(out_key)==1:
        return data[out_key[0]],out_key[0]
    else:
        return [data[out_key[i]] for i in range(len(out_key))],out_key

# def fastfetch_keydata(data_storage,target_key):
#     #fast fetch a particular key data from the data storage, e.g. pic1_1
#     #this function may need some revision; i cant trace back to where i used it
#     out =[]
#     if type(file)==list:
#         for f in data_storage:
#             keydata,key = get_keydata(target_key,data_storage[f],partial=False)         
#             out.append(keydata)   
#     else:
#         keydata,key = get_keydata(target_key,data_storage[f],partial=False)
#         out = keydata
#     return out

def fastfetch_coldata(file,col_num,target_key,target_dtype):
    #fast fetch a particular column's data from excel files from the column's title, e.g. pic1_1
    #this function may need some revision; i cant trace back to where i used it
    out =[]
    if type(file)==list:
        for f in file:
            data,row = read_rows(read_table(f),col_num)
            clean = num_data_clean(data,num_type=target_dtype)
            keydata,key = get_keydata(target_key,clean,partial=False)         
            out.append(keydata)   
    else:
        data,row = read_rows(read_table(file),col_num)
        clean = num_data_clean(data,num_type=target_dtype)
        keydata,key = get_keydata(target_key,clean,partial=False)
        out = keydata
    return out

def assemble_dict(lists,keynames,stack=True,stack_keyname=['val','cond'],to_df=True):
    #e.g. df = xf.assemble_dict([free,yoke,pasv],['free','yoke','pasv'],stack_keyname=['recall','cond'])
    if stack==True:
        data = init_dict(stack_keyname)
        data[stack_keyname[0]]=merge_lists_in_list(lists)
        cond=[[keynames[i]]*len(lists[i]) for i in range(len(keynames))]
        data[stack_keyname[1]]=merge_lists_in_list(cond)
    else:
        data = init_dict(keynames)
        for i in range(len(lists)):
            data[keynames[i]]=lists[i]
    if to_df==True:
        data = dict_to_df(data,colname_by_key=True)
    return data

def dict_to_lists(data_as_dict):    
    list_of_keys =[i for i in data_as_dict.keys()]
    list_of_values = [data_as_dict[i] for i in list_of_keys]
    list_lengths = [len(i) for i in list_of_values]
    unique,repeat = get_unique_repeat(list_lengths)
    if len(unique)!=1:
        print('Warning: the length of the lists do not all equal!')
    return list_of_values,list_of_keys

def dict_to_df(data,colname_by_key=False):
    import pandas as pd 
    if colname_by_key==False:
        out = pd.DataFrame.from_dict(data,orient='index')
    else:
        out = pd.DataFrame.from_dict(data)
    return out
 
def keep_rows(keep_start,keep_len,ind_start=1):
    #only for continuous row numbers
    #ind_start =0 means start should -1
    if type(keep_start)==list:
        if type(keep_len)!=list:
            keep_len = [keep_len]*len(keep_start)
        count=0
        keeprows=[]
        for start in keep_start:
            keep = [i+(ind_start-1) for i in range(start,start+keep_len[count])]
            keeprows.append(keep)
            count+=1
        keeprows = merge_lists_in_list(keeprows)
    else: #a single number
        keeprows = [i+(ind_start-1) for i in range(keep_start,keep_start+keep_len)]
    return keeprows

def exclude_rows(all_start,all_len,keep_start,keep_len,ind_start=1):
    allrows = [i+(ind_start-1) for i in range(all_start,all_start+all_len)] 
    keeprows = keep_rows(keep_start,keep_len,ind_start=ind_start)
    exclude = get_subtract(allrows,keeprows)
    return exclude

def list_exclude_inds(inlist,pop_inds):
    import numpy as np
    pop_inds = sorted(pop_inds,reverse=True)
    if type(inlist)==list:    
        for i in pop_inds:
            inlist.pop(i)
    elif type(inlist)==np.ndarray:
        inlist = inlist.tolist()
        for i in pop_inds:
            inlist.pop(i)
        inlist = np.array(inlist)
    return inlist

def list_select_inds(inlist,keep_inds):
    out = []
    for ind in keep_inds:
        out.append(inlist[ind])
    return out

def list_rm_mid_items(inlist, n_remove):
    lst = inlist
    n = len(lst)
    if n <= n_remove:
        return []  # If the removal count exceeds or matches list size, return an empty list.
    mid_point = n // 2
    if n % 2 == 0:
        start = mid_point - (n_remove // 2)
    else:
        start = mid_point - (n_remove // 2)
    end = start + n_remove
    # Ensure start and end indices are within the bounds of the list
    start = max(start, 0)
    end = min(end, n)
    return lst[:start] + lst[end:]

def lists_to_dict(list_of_lists,header='row1'):
    data = list_of_lists
    if header =='row1':
        tidy_data = init_dict(data[0])
    else:
        tidy_data = init_dict([str('col'+str(i)) for i in range(len(data[0]))])
    key_list = [i for i in data[0]]
    for i in range(len(key_list)):
        key_name = key_list[i]
        for j in range(len(data)):
            if j !=0:
                tidy_data[key_name].append(data[j][i])
    return tidy_data

def list_add_vals(ls,vals):
    for i in vals:
        ls.append(i)
    return ls

def lists_add_vals(lists,vals):
    if len(lists)==len(vals):
        for i in range(len(lists)):
            lists[i].append(vals[i])
        return lists
    else:
        print('Error: the two inputs need to have the same length!')


def dict_sort_inds(data, sort_by_key, target_key_seq=None):
    if target_key_seq is None:
        # Sort by the values of sort_by_key
        sorted_indices = sorted(range(len(data[sort_by_key])), key=lambda i: data[sort_by_key][i])
    else:
        # Sort by the sequence of target_key_seq
        sorted_indices = [data[sort_by_key].index(val) for val in target_key_seq]
    sorted_data = {}
    for key in data:
        sorted_data[key] = [data[key][i] for i in sorted_indices]
    return sorted_data

def dict_exclude_inds(data,exclude_inds):
    inds = sorted(exclude_inds,reverse=True)
    out = init_dict(keylist=[key for key in data.keys()])
    for key in data:
        full_len = len(data[key])
        for ind in range(full_len):
            if ind not in exclude_inds:
                out[key].append(data[key][ind])
    return out

def dict_select_inds(data,select_key,val_match,find=False):
    #e.g. data_dict = {key1:[val11,val11,val12,val13];key2:[val21,val22,val23,val24]}
        # select_key = key1; 
        # val_match = val11
        # out = {key1:[val11,val11],key2:[val21,val22]}
    out  = {}
    if type(select_key)==list:
        inds = []
        for i in range(len(select_key)):
            inds.append(get_indices(data[select_key[i]],val_match[i],find=find))
        final_inds = get_intersect_all(inds) 
        for key in data:
            out[key]=apply_inds(final_inds,data[key])
    elif type(select_key)==str:
        inds = get_indices(data[select_key],val_match,find=find)
        for key in data:
            out[key]= apply_inds(inds, data[key])
    return out


def dict_remove_keys(data,exclude_keys=''):
    out = data.copy()   
    if type(exclude_keys)==str:
        out.pop(exclude_keys)
    else: #a list of keys to pop
        for key in exclude_keys:
            out.pop(key)
    return out

def dict_exclude_keys(data,exclude_keys='',general=False,find=False):
    #if general==True: kick out keys/vals of a dict according to exclude keys
        #if not: this is to exclude the keys that had no values in them
    #if find==True, as long as the data key contains the string specified by exclude_keys, it should be excluded
    if find==True:
        keyls = [k for k in data]
        keyinds = get_indices(keyls,exclude_keys,True)
        exclude_keys = apply_inds(keyinds,keyls)
    if general==True:
        if type(exclude_keys)==str:
            data.pop(exclude_keys)
        else: #e.g. list
            for i in exclude_keys:
                data.pop(i)
        return data
    else:
        pop_list = []
        for key in data:
            uni,rep = get_unique_repeat(data[key])
            if len(uni)==1 and uni[0]==exclude_keys:
                pop_list.append(key)
        out_data = data
        for i in pop_list:
            out_data.pop(i)
        return out_data

def dict_include_keys(data,keep_keys,find=False):
    #if find==True, partial match with any of the keep_keys should be kept
    #this is to keep keys in the data
    out={}
    if find==False: #complete match
        for key in data:
            if key in keep_keys:
                out[key]=data[key]
    else: #partial match
        a = [i for i in data.keys()]
        b = keep_keys
        keep_list = [item for substring in b for item in a if substring in item]
        for key in keep_list:
            out[key]=data[key]
    return out

def dict_exclude_rows(data,by_key,by_search,exclude=1):
    #if exclude=1, exclude the rows where the by_key contains by_search
    #if exclude=0, keey the rows where the by_key contains by_search
    ind_list = []
    count=0
    for i in data[by_key]:
        if by_search in str(i):
             ind_list.append(count)
        count+=1    
    if exclude==1:
        exclude_inds = ind_list
    else:
        all_inds = [i for i in range(len(data[by_key]))]
        exclude_inds = get_subtract(all_inds,ind_list)
    out_data = dict_exclude_inds(data,exclude_inds)
    return out_data

def dict_exclude_rows_multi(data,by_key,search_list):
    for i in search_list:
        data = dict_exclude_rows(data,by_key,i)
    return data

def dict_collapse(data_dict,uni_per_key=False,layered=False):
    #this is to collapse dict to list or from dict to dict with one layer down 
        #e.g. return a list of [key_val1,key_val2,...] from data[key]=[val1,val2,...]
    #uni_per_key if True then through away repeated key_val in the outlist
    #layered if True assumes data_dict[key] is also a dict and data_dict[key][key2] is a list or array object
    out=[]
    outdict = {}
    for key in data_dict:
        if type(data_dict[key])==dict:
            for key2 in data_dict[key]:
                outdict[key+'_'+key2]=[]
                for val in data_dict[key][key2]:
                    outdict[key+'_'+key2].append(str(val))
        else: #assume data_dict[key] is a list or array
            for val in data_dict[key]:
                out.append(str(key)+ '_' +str(val))
    if layered ==False: #collapse dict to list
        if uni_per_key !=False:
            uni,rep = get_unique_repeat(out)
            out = uni
        return out
    else: #collapse dict to dict with one layer down 
        return outdict

def group_cut(to_be_cut_list,cut_by_list,cut_item):
    #example use: group_cut(list1,list2,'_')
        #list1=['1','1','2','3','3','3'], list2=['a','b','c','d','d_1','e']
        #cut_item is the item to search for in the cut_by_list as cut-off points
        #returns list1 in: [['1','1','2','3'],['3'],['3']]
    if len(to_be_cut_list) != len(cut_by_list):
        print('Error: the lists need to be of the same length!')
    group_by_list = []
    count=1
    for i in cut_by_list:
        if cut_item in i:
            count+=1
            group_by_list.append(count)
            count+=1
        else:
            group_by_list.append(count)           
    grouped = group_by(group_by_list,to_be_cut_list)
    return grouped,group_by_list

def group_by_list_renum(group_by_list):
    group_by = []
    count=1
    for i in range(0,len(group_by_list)):
        if i !=0:
            prev= group_by_list[i-1]
            cur = group_by_list[i]
            if prev!=cur:
                count+=1
        group_by.append(count)
    return group_by

def group_by_both_list(group_by_list1,group_by_list2):
    #this function requires the two lists be in natural number sequence
    #example use: group_by_both_list(group_by_list,group_by_cut) 
        # group_by_list= ['1','1','2','2','3','3','3'] 
        # group_by_cut = ['1','1','1','1','2','3','3']
        # returns: [1, 1, 2, 2, 3, 4, 4]
    if len(group_by_list1) != len(group_by_list2):
        print('Error: the lists need to be of the same length!')
    else:
        list1 = group_by_list_renum(group_by_list1)
        list2 = group_by_list_renum(group_by_list2)
    count=1
    group_by_both = []
    for i in range(0,len(list1)):
        if i !=0:
            prev_grp= list1[i-1]
            cur_grp = list1[i]
            prev_cut = list2[i-1]
            cur_cut  = list2[i]        
            if cur_grp != prev_grp or prev_cut!=cur_cut:
                count+=1
        group_by_both.append(count)
    return group_by_both

def group_by(group_by_list,to_be_grouped_list):
    #example use-1: group_by(list1,list2) 
        # list1 = ['1','1','2','3','3','3','2']; 
        # list2=['1','0','0','1','1','0','1']
        # returns list2 in: [['1','0'],['0','1'],['1','1','0']]
    #example use-2: group_by(list1,list2,group_cut=[list3,'_'])
        # list3 = ['a','b','c','d','d_1','e','f']
        # returns list2 in: [['1','0'],['0','1'],['1'],['1','0']]
    if len(group_by_list) != len(to_be_grouped_list):
        print('Error: the lists need to be of the same length!')
    else:
        grouped = []
        groups,repeat= get_unique_repeat(group_by_list)
        for grp in groups:
            inds = get_indices(group_by_list,grp)
            cur_grp = apply_inds(inds,to_be_grouped_list)
            grouped.append(cur_grp)
        return grouped
    
def group_lists(list_of_lists,group_by_list_ind=0):
    #example use: 
        # group_lists(alldat)
        # list1 = ['1','1','2','3','3','3','2']
        # list2=['1','0','0','1','1','0','1']
        # list3=['right','wrong','wrong','right','right','wrong','right']
        # alldat = [list1,list2,list3]
        # returns: 
            # [[['1', '1'], ['2', '2'], ['3', '3', '3']],
            #  [['1', '0'], ['0', '1'], ['1', '1', '0']],
            #  [['right', 'wrong'], ['wrong', 'right'], ['right', 'right', 'wrong']]]
    group_by_list = list_of_lists[group_by_list_ind]
    out_list_of_lists = []
    for cur_list in list_of_lists:
        out_list_of_lists.append(group_by(group_by_list,cur_list))
    return out_list_of_lists

def group_dict_shrink(data_dict,sep=' '):
    #e.g. data_dict = groupped_data
    group = data_dict
    for key in group:
        for i in range(len(group[key])):
            ls = group[key][i]
            uni,rep= get_unique_repeat(ls)
            if len(uni)==1:
                group[key][i]=ls[0]
            else:
                group[key][i]=sep.join(group[key][i])
    return group

def group_dict(data_dict,group_by_key,exclude_keys=False,merge=False,sep=' '):
    #group all the lists in the dictionary by one of the key-corresponded list
    #returns a dictionary with keys each correpond to a list of lists (i.e. list being grouped)
    #exclude_keys: e.g. ['key1','key2'] to be excluded from the group_by_key, so corresponding data will also be excluded from the returned data
    list_of_lists,names_of_lists = dict_to_lists(data_dict)
    ind = get_indices(names_of_lists,group_by_key)
    grouped_lists = group_lists(list_of_lists,ind[0])
    #put back in dict format:
    grouped_data = {}
    count=0
    for name in names_of_lists:
        if name not in grouped_data:
            grouped_data[name]=grouped_lists[count]
            count+=1
    if exclude_keys !=False:
        exclude_ind =[]
        for ind in range(0,len(grouped_data[group_by_key])):
            for item in exclude_keys:
                if item in grouped_data[group_by_key][ind]:
                    exclude_ind.append(ind)
        grouped_data = dict_exclude_inds(grouped_data,exclude_ind)
    if merge!=False:
        grouped_data=group_dict_shrink(grouped_data,sep=sep)
    n_groups = len(grouped_data[group_by_key])
    return grouped_data,n_groups

def group_means(grouped_data,n_groups,keys_not_cal,keys_exclude=[],rnd=2):
    #this is to calculate the mean of each grouped list from grouped_data, returns a dictionary of group means by key  
    #keys_not_cal are the key names that we do not calculate the means for; these are usually keys with vals of str
    #e.g. keys_not_cal = ['event','section'], keys_exclude=[]
    #e.g. rnd = 4 is to keep 4 decimals for calculating the mean
    mean_by_group = init_dict([i for i in grouped_data.keys() if i not in keys_exclude])
    for i in range(n_groups):
        for key in mean_by_group:
            if key not in keys_exclude:
                if key in keys_not_cal:
                    cur_event = grouped_data[key][i][0]
                    mean_by_group[key].append(cur_event)
                else:
                    cur_mean = round(avg(grouped_data[key][i],mute=True),rnd)
                    mean_by_group[key].append(cur_mean)
    return mean_by_group

def split_list(inlist):
    #this is to split a list into 2 of approximately equal size
    #if inlist is a list of lists, this splits all the lists, each into 2 lists
    if type(inlist[0])==list:
        split_list=[]
        for ls in inlist:
            length = len(ls)
            middle_index = (length//2)+1
            half1 = ls[:middle_index]
            half2 = ls[middle_index:]
            split_ls = [half1,half2]
            split_list.append(split_ls)
    else:
        length = len(inlist)
        middle_index = (length//2)+1
        half1 = inlist[:middle_index]
        half2 = inlist[middle_index:]
        split_list = [half1,half2]
    return split_list

def sort_nlists(list_of_lists):
    #sort by the first list in the list_of_lists
    out = list_of_lists
    import numpy as np
    list1 = list_of_lists[0]
    idx = np.argsort(list1)
    for i in range(len(out)):
        out[i]=apply_inds(idx,list_of_lists[i])     
    return out

def get_top_n(list_of_lists, n, minmax='min'):
    #get the n smallest values from the first list in the list_of_lists, retaining info from other lists 
    use = sort_nlists(list_of_lists)
    out = use
    for i in range(len(use)):
        if minmax=='min':
            out[i]=use[i][:n]
        elif minmax=='max':
            out[i]=use[i][-n:]
    return out

def merge_items_in_list(list_of_lists,sep=' '):
    #example use: merge_items_in_list([['1','2'],[2,4],['lisa','donald','ana']])
        #returns: ['1 2','2 4','lisa donald ana'] 
    full_list=[]
    simp_list = []
    for ls in list_of_lists:
        #convert list of int/float to list of str:
        ls = [str(i) for i in ls]
        uni_ls,repeat = get_unique_repeat(ls)
        if sep==' ':
            new_item = ' '.join(ls)
            simp_item = ' '.join(uni_ls)
        elif sep ==',':
            new_item = ','.join(ls)
            simp_item = ','.join(uni_ls)
        elif sep =='.':
            new_item = '.'.join(ls)
            simp_item = '.'.join(uni_ls)
        full_list.append(new_item)
        simp_list.append(simp_item)
    return simp_list,full_list

def merge_tuples_in_list(list_of_tuples):
    flat_list = [item for tup in list_of_tuples for item in tup]
    return flat_list

def merge_lists_in_list(list_of_lists):
    #example use: merge_items_in_list([[1,2],[2,4],['lisa','donald','ana']])
        #returns: [1,2,2,4,'lisa','donald','ana'] 
    new_list = []
    for ls in list_of_lists:
        if type(ls)==list:
            new_list = new_list+ls
        else:
            new_list.append(ls)
    return new_list
    
def merge_ioflist_in_data(data,sep=' '):
    #this is best used following grouped_data where we have each key-corresponded list are of grouped elements (i.e. data[key]=list of lists, e.g. [[1],[0,1]])
    #returns the data with data[key]=list (each item is the merged items in the previously grouped list)
    full_data = {}
    simp_data = {}
    for key in data:
        cur_list = data[key]
        simp_list,full_list = merge_items_in_list(cur_list,sep=sep)
        simp_data[key]=simp_list
        full_data[key]=full_list
    return simp_data,full_data

def merge_lsoflist_in_data(data):
    new_data = {}
    for key in data:
        cur_list = data[key]
        new_list = merge_lists_in_list(cur_list)
        new_data[key]=new_list
    return new_data

def flip_lsoflist(list_of_lists):
    #this assumes that all of the sub-lists are of equal length
    a = list_of_lists
    new_a = []
    for j in range(len(a[0])):
        b =[]
        for i in range(len(a)):
            b.append(a[i][j])
        new_a.append(b)
    return new_a

def shared_overall(vec1,vec2):
    #example: vec1=[15,9,28], vec2=[3,15,9], out=len([15,9])/len([3,15,9,28])=.5
    shared = len(get_intersect(vec1,vec2))
    total = len(get_unique_repeat(merge_lists_in_list([vec1,vec2]))[0])
    if total==0:
        pp=0
    else:
        pp = shared/total
    return pp

def shared_history(vec1,vec2,curr_item=False):
    #up until each shared element, what's the proportion of overlap (shared unique #of elems/total unique #of elems) between them? 
    #returns a list of history proportion overlay. 
    shared = get_intersect(vec1,vec2)
    history = []
    for i in shared:
        ind1 = vec1.index(i)
        ind2 = vec2.index(i)
        if curr_item==True:
            ind1= ind1+1
            ind2= ind2+1
        subset1 = vec1[0:ind1]
        subset2 = vec2[0:ind2]
        pp = shared_overall(subset1,subset2)
        history.append(pp)
    return history




def data_info(in_data):
    #this is to browse the data strcuture of the input data (mostly for dictionaries)
    info =init_dict(['dtype','dstructure'])
    dt = type(in_data)
    info['dtype']=dt
    if type(in_data)==np.array or  type(in_data)==list:
        ds = shape(in_data)
        info['dstructure']=ds
    elif type(in_data)==dict:
        ds=init_dict(['nkeys','keyls','keyvalue'])
        ds['keyls']=[i for i in in_data.keys()]
        ds['nkeys']=len(ds['keyls'])
        ds['keyvalue']=[data_info(in_data[k]) for k in ds['keyls']]
        info['dstructure']=ds
    else:
        info['dstructure']=shape(in_data)
    return info

def datakey_samelen(data):
    #this is to check if all data[key] are of the same length, assuming every data[key] is a list of values
    uni,rep = get_unique_repeat([len(data[key]) for key in data])
    if len(uni)==1:
        mesg = "same length confirmed: " + str(uni[0])
        return mesg
    else:
        return "data[key] of different length!"
        
def data_subset(data_dict,keep_keys):
    #this is to select a subset of the dictionary, keeping data of the specified keys
    subset={}
    for key in keep_keys:
        subset[key] = data_dict[key]
    return subset

def data_storage(file_path,file,col_num,target_dtype=False):
    #fast fetch a particular column's data from excel files from the column's title, e.g. pic1_1
    out ={}
    if type(file)==list:
        for f in file:
            data,row = read_rows(read_table(file_path+f),col_num)
            if target_dtype==False:
                clean = data
            else:
                clean = num_data_clean(data,num_type=target_dtype)
            out[f]=clean
    else:
        data,row = read_rows(read_table(file_path+file),col_num)
        if target_dtype==False:
            clean = data
        else:
            clean = num_data_clean(data,num_type=target_dtype)
        out = clean
    return out

def data_by_key(data,clear_empty=True,clear = ''):
    data_org = {}
    for row in data:
        for key in data[row].keys():
            if key not in data_org:
                data_org[key]=[]
            data_org[key].append(data[row][key][0])
    
    if clear_empty ==True:
        key_list = [key for key in data_org.keys()]
        key_num = len(key_list)
        num = len(data_org[key_list[0]])
        ind_list = []
        for ind in range(num):
            check = [data_org[key][ind]==clear for key in data_org.keys()]
            if check == [True]*key_num:
                ind_list.append(ind)
                #print(ind)
        for ind in sorted(ind_list,reverse=True):
            for key in data_org.keys():
                data_org[key].pop(ind)


    return data_org

def dict_by_key(data,keynames='default'):
    keylist = [key for key in data.keys()]
    datlen = len(data[keylist[0]])
    if keynames=='default':
        keynames = ['keys']
        for i in range(0,datlen):
            col_name = 'data_col'+str(i+1)
            keynames.append(col_name)
    data_org={}
    for name in keynames:
        data_org[name]=[]
    for key in data:
        data_org[keynames[0]].append(key)
        for i in range(0,datlen):
            data_org[keynames[i+1]].append(data[key][i])
    return data_org

def grab_all(path,filetype,exclude=False,mute=False):
    #path: location of your data files, e.g. '/Users/xianl/Desktop/data/'
    #filetype: specify the file suffix seeking for, e.g. 'xlsx', 'txt'
    #example use:
        #file_list = grab_all('/Users/xianl/Desktop/data/','.txt')
        #this function returns a list of file names in the path you specified
    import glob
    allfile = []
    for file in glob.glob(path + '*'+filetype):
        file = file.split("/")
        file = file[len(file)-1]
        if exclude!=False:
            if file not in exclude:
                allfile.append(file)  
        else:
            allfile.append(file) 
    if mute !=False:
        print(allfile)
    return sorted(allfile)

def grab_fname(path,search_name,exclude=False,mute=False):
    #path: location of your data files, e.g. '/Users/xianl/Desktop/data/'
    #search_name: specify the thing to search for. e.g. 'sub*' or '*_mvp.xlsx' etc.
    #example use:
        #file_list = grab_fname('/Users/xianl/Desktop/data/','*.txt')
        #this function returns a list of file names in the path you specified
    import glob
    allfile = []
    for file in glob.glob(path +search_name):
        file = file.split("/")
        file = file[len(file)-1]
        if exclude!=False:
            if file not in exclude:
                allfile.append(file)  
        else:
            allfile.append(file) 
    if mute !=False:
        print(allfile)
    return sorted(allfile)

def grab_base(path,file_search='*',exclude=False,mute=False):
    #path: location of your data files, e.g. '/Users/xianl/Desktop/data/'
    #file_search: specify the file name and suffix seeking for, e.g. 'sub*'
    #example use:
        #file_list = grab_all('/Users/xianl/Desktop/data/','.txt')
        #this function returns a list of file names in the path you specified
    import glob
    import os
    allfile = []
    for file in glob.glob(path +file_search):
        base = os.path.basename(file).split('.')[0]
        allfile.append(base)
    if mute !=False:
        print(allfile)
    return sorted(allfile)

def read_table(infile,sheet=1):
    #infile specifies the name of the file and its path
    #example use:
        # table = read_table('/Users/xianl/Desktop/data/sub1.xlsx') 
        # this returns a table that you can process the rows for, e.g.:
            # nrows = table.nrows
            # for row in range(nrows):
            #    row_counter+=1
            #    vec = table.row_values(row)
    import xlrd
    data = xlrd.open_workbook(infile)
    table = data.sheets()[sheet-1]
    return table

def read_rows(table,col_num,key=True,skip_row1=True,clear_empty=True,clear=''):
    #table: read_table('./2_data/1_summary.xlsx')    
    #col_num: either int or list of int for the column number (not index) wanted
    #key: defaults key to be the title row; if no title row, need to input a list of strings for key name
    #skip_row1: if no title row, needs to input the var key, and specify skip_row1=False
    #example use: 
        #all_subs = read_rows(table,8)
        #all_subs = read_rows(table,[1,7,8])
    row_count =0
    nrows = table.nrows
    all_rows = {}
    #read row:
    for row in range(nrows):
        row_count+=1
        v = table.row_values(row)
        #skip first row (if key not specified: grab title)
        if skip_row1==True and row_count == 1:
            if key ==True:
                if type(col_num)==int:
                    key = str(v[col_num-1])
                elif type(col_num)==list:
                    key = []
                    for i in col_num:
                        key.append(str(v[i-1]))
                else:
                    print('ERROR: col_num type needs to be either int or list of int')
            continue
        elif skip_row1==False and row_count==1:
            key=['col_'+str(i) for i in col_num]
        #update progress of row reading:
        current_row = 'row'+str(row_count)
        #print(current_row)
        #check out and save the col value(s)
        if type(col_num)==int:
            out = v[col_num-1]
            all_rows[current_row]={}
            if key not in all_rows[current_row]:
                all_rows[current_row][key]=[]
            all_rows[current_row][key].append(out)
            #print(out) 
        if type(col_num)==list:
            all_rows[current_row]={}
            key_count=0
            for keyname in key:
                key_count+=1
                if keyname not in all_rows[current_row]:
                    all_rows[current_row][keyname]=[]
                item = str(v[col_num[key_count-1]-1])
                all_rows[current_row][keyname].append(item)    
            #print(all_rows[current_row])
    data = data_by_key(all_rows,clear_empty=clear_empty,clear=clear)
    return data, all_rows 

def read_data(infile,col_num,sheet=1,key=True,skip_row1=True,clear_empty=True,clear='',clean_keylist=False):
    data,row = read_rows(read_table(infile,sheet=sheet),col_num=col_num,key=key,skip_row1=skip_row1,clear_empty=clear_empty,clear=clear)
    if clean_keylist!=False:
        data = num_data_clean(data,keylist=clean_keylist)
    return data

def read_data_num(infile,col_num,num_type='float',skip_row1=True):
    #first read in from read_data function, then clean num
    dat = read_data(infile,col_num,skip_row1=skip_row1)
    out = {}
    for key in dat:
        if num_type=='float':
            out[key]=[float(i) for i in dat[key]]
        elif num_type =='int':
            out[key]=[int(i) for i in dat[key]]
        elif num_type =='intfloat':
            out[key]=[int(float(i)) for i in dat[key]]
        elif num_type =='floatint':
            out[key]=[float(int(i)) for i in dat[key]]
    return out

def num_data_clean(data,keylist='all',num_type='float',exclude='',get_mean=False,pop_empty=False,keep_as_str=False):
    #this is to clean the data dict read from the vertically arranged excel file: 
    #dict can be a mix of num & str want-to-be by the column 
    #take items in each column as num, pop the ''    
    #if get_mean !=False, it also outputs the mean per key
    out_data = {} 
    if keylist=='all':
        keylist = [i for i in data]
    else:
        other_keys = [i for i in data if i not in keylist]     
        for i in other_keys:
            out_data[i]=data[i]
    for key in keylist:
        out_data[key] = [str_to_num(i,num_type,keep_as_str) for i in data[key] if i!=exclude]
    if pop_empty !=False:
        pop_key = []
        for key in out_data:
            if out_data[key]==[]:
                pop_key.append(key)
        for key in pop_key:
            out_data.pop(key)    
    if get_mean!=False:
        import numpy as np
        mean = []
        for key in out_data:
            mean.append(np.mean(out_data[key])) 
        return out_data,mean
    else:
        return out_data

def clean_data(data,numkey=False,num_type=False,keep_as_str=False,textkey='default'):
    # this function cleans up the number columns and text columns separately in the specified way
    # example use:
    # clean_data(data_dict,
    #            numkey=['Scene_lab','Section','Storyline','Outcome_1','Outcome_2'],
    #            num_type=['intfloat']*5,
    #            keep_as_str=['_',False,False,'NA,','NA,'], #if met these special cases, keep it as a str   
    #            textkey=['Story','Response_1','Response_2'])
    out = data
    if numkey !=False:
        for i in range(len(numkey)):        
            if keep_as_str==False:
                keep_as_str = [False]*len(numkey)
            out = num_data_clean(out,keylist=[numkey[i]],num_type=num_type[i],keep_as_str=keep_as_str[i])
    if textkey=='default': #default to be all other keys' data
        for key in get_subtract(data.keys(),numkey):
            out[key]=clean_textls(data[key])
    else:
        for j in range(len(textkey)):
            out[textkey[j]] = clean_textls(data[textkey[j]])
    return out

def read_csv(vectorfile,asnum=True,mute=True,skip_row1=False):
    # read in csv file
    # vectorfile: the csv file saved for USE vectors
    # example use:
        # read_vector('./autoScore/sub2_recall_SentenceEmbeddings_6nodes.csv')
        # outputs the ndarray: nodes (number of events) x length of the USE vector (512)
    import csv
    import numpy as np
    with open(vectorfile) as csvfile:
        readCSV = csv.reader(csvfile, delimiter=',')
        vectorlist = []
        for row in readCSV:
            vectorlist.append(row)
    num_events = len(vectorlist)
    vector_len = len(vectorlist[0])
    array = np.array(vectorlist).reshape((num_events,vector_len))
    if skip_row1==True:
        array= array[1:]
    if asnum==True:
        array = array_str_to_num(array)
    if mute!=True:
        print(array.ndim,'\n',array.shape)
    else:
        return(array)

def read_csv_data(rawfile,asnum=True,mute=True,header=False,todict=False):
    # read in csv file which can contain any type of data in any format
        #if header==False: returns a list with each element being a list containing data from a row, e.g. [[row1_i1, row1_i2],[row2_i1],[row3_i1,row3_i2,row3_i3]]
        #if header=='col1': returns a dictionary with key being the first column items
    import csv
    with open(rawfile) as csvfile:
        readCSV = csv.reader(csvfile, delimiter=',')
        rawlist = []
        for row in readCSV:
            rawlist.append(row)            
    outlist = rawlist
    if header=='col1':
        outlist = {}
        for row in rawlist:
            outlist[row[0]]=row[1:]                     
    elif header!=False:
        print('header needs to be either col1 or False')
    out=outlist
    if todict==True:
        out = lists_to_dict(out)
    return(out)

def read_tsv(file):
    import pandas as pd
    df = pd.read_csv(file, sep='\t')
    return df

def read_text(filename,sep=False,header=False,item_sep=','):
    # read in txt file which can contain any type of data in any format
        #if header==False: returns a list with each element being a list containing data from a row, e.g. [[row1_i1, row1_i2],[row2_i1],[row3_i1,row3_i2,row3_i3]]
        #if header=='col1': returns a dictionary with key being the first column items
        #sep='\n' means we are separating by lines in text file
        #item_sep=',' means we are further separating each line by comma, producing a list per line
    file1 = open(filename,"r", errors='ignore')
    text = file1.read()
    if sep !=False:
        text = text.split(sep)
        text = clean_list(text)
    else:
        text = clean_text(text)
    outlist = text
    if header=='col1':
        header = []
        for line in text:
            header.append(line.split(item_sep)[0])
    if type(header)==list:
        if len(header)==len(text):
            outlist = init_dict(header)
            count=0
            for line in text:
                key = header[count]
                content = line.split(key)[-1][1:]
                content_ls = content.split(item_sep)
                outlist[key]=content_ls
                count+=1
        else:
            print('header needs to be the same length as the line number in text file!')    
    elif header!=False:
        print('header needs to be either a list of keys, col1, or False')
    return outlist

def read_docx(filename,sep=False,clean_up='Default'):
    #example use: read_docx('sub2.docx') 
        #this returns a text string that contains all content of the docx file
    import docx
    import re
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    raw_text = '\n'.join(fullText)
    cleaned_text = clean_text(raw_text,clean_up=clean_up)
    if sep !=False:
        cleaned_text_list = clean_list(re.split(r'[.?!]',cleaned_text))
        clean_text_list = []
        for i in cleaned_text_list:
            clean_text_list.append(clean_text(i))
        return raw_text, cleaned_text,clean_text_list
    else:
        return raw_text, cleaned_text

def init_mat(shape,fill=0):
    import numpy as np
    if fill==0:
        mat = np.zeros(shape)
    elif fill==1:
        mat = np.ones(shape)
    elif fill=='nan':
        mat = np.zeros(shape)
        mat[mat==0]=np.nan
    else: #only works when fill is an integer
        mat = np.zeros(shape)
        mat[:]=fill
    return mat

def init_list(length,label='lab'):
    ls = []
    for i in range(1,length+1):
        ls.append(label+str(i))
    return ls

def init_dict(keylist,repeat=False,value='list'):
    dic={}
    uni,rep = get_unique_repeat(keylist)         
    for i in uni:
        if value=='list':
            dic[i]=[]
        elif value=='dict':
            dic[i]={}
    if repeat!=False:   
        uni2,rep2 = get_unique_repeat(rep)
        for i in uni2:
            inds = get_indices(keylist,i)
            for j in range(1,len(inds)):
                key = str(i)+'_'+str(j+1)
                if value=='list':
                    dic[key]=[]
                elif value=='dict':
                    dic[key]={}
    return dic

def init_dict_w_header(keylist,header_tab,title='title',repeat=False,value='list'):
    #default keylist does not contain the title key ('title') nor its corresponding values (title_tab)
    #title_tab is the header
    data = init_dict(merge_lists_in_list([title,keylist]))
    data[title]=header_tab
    return data

def combo_ls(lists_of_strings,join_by='_'):
    #this is to help create the key list when using init_dict
    """
    Takes a list of lists containing strings and returns a list of all possible
    combinations of these strings.

    Args:
    lists_of_strings (list of list of str): List of lists, each containing strings.

    Returns:
    list of str: List of all possible combinations of the strings.
    """
    from itertools import product
    if not lists_of_strings:
        return []
    # Use itertools.product to get all combinations
    combined = product(*lists_of_strings)    
    # Join each combination into a single string
    combined_strings = [join_by.join(combination) for combination in combined]
    return combined_strings

def dict_add_vals(dic,keys,vals,append=True):
    #this is to add vals to a list of specified keys in an existing dictionary
    if len(keys)!=len(vals):
        print('Error: keys and vals need to be lists of the same length for this function!')
        print(keys)
        print(vals)
        print(len(keys),len(vals))
    else:
        for i in range(len(keys)):
            if append==False:
                dic[keys[i]]=vals[i]
            else:
                dic[keys[i]].append(vals[i])
        out = dic
        return out

def dict_add_keys(dic,keylist,repeat=False):
    #this is to add a list of new keys to an existing dictionary
    #if repeat=False, cannot add new key that's already in dict; else will add it as key2
    existing_list = [i for i in dic.keys()]
    keys_added = []
    for key in keylist:
        if repeat==False:
            if key not in existing_list: 
                dic[key]=[]
                existing_list.append(key)
                keys_added.append(key)
        else:
            if key not in existing_list: 
                dic[key]=[]
                existing_list.append(key)
                keys_added.append(key)
            else:
                cur_num = str(len(get_indices(existing_list,key))+1)
                dic[key+cur_num]
                existing_list.append(key)
                keys_added.append(key+cur_num)
    return dic,keys_added

def dict_add_keys_n_vals(dic, key_list,val_list,append=True,verbose=False):
    #this is to add a list of new keys and their corresponding vals to an existing dictionary
    #if append==True, the values will be appended to the key's corresponding list, not directly equal to the key's value
    dic,keys_added = dict_add_keys(dic,key_list)
    if len(keys_added)!=len(key_list): #meaning that some of the keys already exists
        dic = dict_add_vals(dic,key_list,val_list,append=append)
        if verbose!=False:
            print('warning: some keys in the key_list already exist in the dict. those keys will not be re-added to the dict.')
            print('below are the keys in the key_list that was actually added to the dict:')
            print(keys_added)
    else:
        dic = dict_add_vals(dic,keys_added,val_list,append=append)
    return dic

def reverse_dict(data,by_key):
    #this is to reverse the keys and keydata of the dictionary
    lenlist = [len(data[key]) for key in data]
    uni,rep = get_unique_repeat(lenlist)
    if len(uni)!=1:
        print('All lists in the dictionary must be of the same length to reverse the key and data!')
    else:
        keylist = merge_lists_in_list([by_key,data[by_key]])
        newdict = init_dict(keylist)
        newdict[by_key]=[key for key in data if key!=by_key]
        counter=0
        for k in newdict:
            if k!=by_key:
                newdict[k]= [data[key][counter] for key in data if key!=by_key]
                counter+=1
        return newdict

def mkdir(inpath):
    import os
    level_to_make=[]
    level_count=-1
    check_dir =  '/'.join(inpath.split('/')[0:level_count])
    while os.path.isdir(check_dir) ==False:
        level_to_make.append(level_count)
        level_count-=1
        check_dir =  '/'.join(inpath.split('/')[0:level_count])
    for level in sorted(level_to_make):
        make_dir = '/'.join(inpath.split('/')[0:level])
        os.mkdir(make_dir)
    if len(level_to_make)==0:
        print('path already exists.')
    else:
        print('path made.')

def write_dict_samelen(dic,outpath,header='default',dict_ind='default',convert=False,set_format=False):
    #applies to: data stored as dictionary with each of the data[key] of same length
    #example use:
        #write_dict_samelen(simp,outfile): will write columns according to the natural sequence of [key for key in dic.keys()]
        #if reorder columns: write_dict(simp,outfile,header=['event_num','event_label','scene_label','','scene_text'],dict_ind=[0,1,2,4])
        #if rename columns:  write_dict(simp,outfile,header=['event','old_seg','scenes','','text'],dict_ind=[0,1,2,4],convert=['event_num','event_label','scene_label','scene_text'])
        #e.g. header = ['score(0/1)','event','old_seg','scene_lab','scene_text','','event_recalled','recall_in_temporal_order','','cause_event','effect_event','reasoning']
        #e.g. dict_ind = [1,2,3,4]
        #e.g. convert=['event_num','event_label','scene_label','scene_text']
        #e.g. set_format = [[3,5],['na',60],[3,5],[2]] #: [[colnums],[colwidth],['wrap these cols'],['hides these cols']]
        #note: header and convert both corresponds to dict_ind
    import xlsxwriter  
    import numpy as np
    workbook = xlsxwriter.Workbook(outpath)
    worksheet = workbook.add_worksheet() 
    if set_format !=False:
        cols = set_format[0]
        width = set_format[1]
        #setting wrap
        wrap = set_format[2]
        for i in range(0,len(wrap)):
            if wrap[i]!='na':
                format = workbook.add_format({'text_wrap': True})
            else:
                format = workbook.add_format({'text_wrap': False})            
            alph = num_to_alphabet([wrap[i]])
            rng = str(alph[0])+':'+str(alph[0])
            worksheet.set_column(rng, None, format)
        #setting hide
        hide = set_format[3]
        for i in range(0,len(hide)):
            if hide[i]!='na':
                alph = num_to_alphabet([hide[i]])
                rng = str(alph[0])+':'+str(alph[0])
                worksheet.set_column(rng, None, None,{'hidden': True})
        #setting width
        for i in range(0,len(cols)):
            if width[i]!='na':
                ind = cols[i]-1
                worksheet.set_column(ind,ind, width[i]) 
    # Start from the first cell
    row, col = 0, 0
    #specify header
    if header=='default':
        header = [key for key in dic.keys()]
    #specify dict_ind:
    if dict_ind =='default':
        dict_ind = [i for i in range(0,len(header))]
    #writing data into each row
    for i in range(0,len(header)): #write down the headers in first row  
        worksheet.write(row, i, header[i])
    row+=1
    keydatalen = len(dic[[key for key in dic.keys()][0]])
    for i in range(0,keydatalen): #specify each of the following rows
        count_j = 0
        for j in dict_ind:
            if convert ==False:
                content = dic[header[j]][i]
                if type(content)==list or type(content)==np.ndarray or type(content)==np.array:
                    content = ','.join([str(i) for i in content])
                worksheet.write(row, col + j, content) 
            else:
                content =  dic[convert[count_j]][i]
                if type(content)==list or type(content)==np.ndarray or type(content)==np.array:
                    content = ','.join([str(i) for i in content])
                worksheet.write(row, col + j, content) 
                count_j+=1
        row += 1
    workbook.close()     
    return 'Excel file saved'

def write_dict_anylen(dic,outpath,vertical=True, header='default',allstr=False):
    #if vertical == True, key written across first row; otherwise, key written across first col
    import xlsxwriter  
    import math
    mkdir(outpath)
    workbook = xlsxwriter.Workbook(outpath)
    worksheet = workbook.add_worksheet() 
    # Start from the first cell
    row = 0
    col = 0
    #specify header
    if header=='default':
        header = [key for key in dic.keys()]
    if vertical==True:
        #write down the headers in first row  
        for i in range(0,len(header)): 
            worksheet.write(row, i, header[i])
        row+=1
        #write in each column
        for i in range(0,len(header)):
            cur_keydata = dic[header[i]]
            for j in range(0,len(cur_keydata)):
                if allstr==True:
                    result = str(cur_keydata[j])
                elif type(cur_keydata[j])==str:
                    result = cur_keydata[j]
                elif cur_keydata[j]==None:
                    result = 'None'
                elif math.isnan(cur_keydata[j]):
                    result = 'nan'
                elif math.isinf(cur_keydata[j]):
                    result = 'inf'
                else:
                    result = cur_keydata[j]
                worksheet.write(row +j, col, result)
            col+=1
    else:
        #write down the headers in first col
        for i in range(0,len(header)): 
            worksheet.write(i, col, header[i])
        col+=1
        #write in each row
        for i in range(0,len(header)):
            cur_keydata = dic[header[i]]
            for j in range(0,len(cur_keydata)):
                if allstr==True:
                    result = str(cur_keydata[j])
                elif type(cur_keydata[j])==str:
                    result = cur_keydata[j]
                elif math.isnan(cur_keydata[j]):
                    result = 'nan'
                elif math.isinf(cur_keydata[j]):
                    result = 'inf'
                else:
                    result = cur_keydata[j]
                worksheet.write(row, col+j, result)
            row+=1        
    workbook.close()     
    return 'Excel file saved'
    
def write_array(array,outpath,header='default',vertical=True):
    dic = {}
    if header =='default':
        dic['values']=array
    else:
        dic[header]=array
    write_dict_anylen(dic,outpath,vertical=vertical)
    return 'Excel file saved'

def write_ndarray(ndarray,outpath,header='default',vertical=True):
    dic={}
    counter=0
    for array in ndarray:
        counter+=1
        if header=='default':
            dic['values'+str(counter)]=array
        else:
            dic[header[counter-1]]=array
    write_dict_anylen(dic,outpath,vertical=vertical)
    return 'Excel file saved'

def write_csv(ndarray,out,transpose=True):
    '''only works for 2d matrices'''
    # note that the first row is header of index
    import pandas as pd
    if transpose!=True:
        pd.DataFrame(ndarray).to_csv(out, index=False)
    else:
        pd.DataFrame(ndarray.T).to_csv(out, index=False)
    print('csv file saved')

def write_csv_ndarray(ndarray,outname):
    import csv
    data = ndarray
    with open(outname+'.csv', 'w', newline='') as file:
        writer = csv.writer(file)
        for row in data:
            writer.writerow(row)
    print(outname+' saved.')

def write_highdim_to_csv(ndarray,out,delimiter=","):
    # e.g. mat = np.random.rand(20, 17, 41, 41)
    # flattened_matrix = matrix_4d.reshape(-1, 41 * 41)
    flattened_matrix = flatten(np.array(ndarray))
    np.savetxt(out, flattened_matrix, delimiter=delimiter)
    return shape(ndarray)

def load_highdim_from_csv(file,shape,delimiter=","):    
    flattened_matrix = np.genfromtxt(file, delimiter=delimiter)
    reshaped_matrix = flattened_matrix.reshape(shape)
    return reshaped_matrix
    
def write_txt(instr,outfile):
    file = open(outfile,"w")
    if type(instr)==str:
        file.write(instr)
    elif type(instr)==list:
        out = ' '.join(instr)
        file.write(out)
    file.close()
    return "Text file has been saved."

def write_docx(content,outpath):
    import docx
    #create file and write content to file
    doc = docx.Document()
    mesg = 'Docx file saved.'
    if type(content)==str:
        doc.add_paragraph(content)
    elif type(content)==list:
        for p in content:
            doc.add_paragraph(p)
    else:
        mesg= 'Error: content to be outputed can only be str or list!'
    #output to .docx file
    if outpath.split('.')[-1]=='docx':
        doc.save(outpath)
    else:
        doc.save(outpath+'.docx')
    return mesg

def write_json(data_dict,outname='./out.json'):
    import json
    #convert dict to json object with specified parameters
    out_json = json.dumps(data_dict,ensure_ascii=False,indent=2)
    with open(outname,'w') as outfile:
        outfile.write(out_json)
    print('json file saved.')






def rotate_mat(matrix, n=1):
    """Rotates a given square 2D matrix list clockwise n times."""
    N = len(matrix)  # Assume square matrix
    for _ in range(n % 4):  # Optimize to only rotate up to 3 times
        # Create a new matrix with the same dimensions
        new_matrix = [[0] * N for _ in range(N)]
        for i in range(N):
            for j in range(N):
                new_matrix[j][N-1-i] = matrix[i][j]
        matrix = new_matrix  # Update matrix to the newly rotated matrix
    return matrix

def zero_pad_mat_diag(original_matrix,new_size=4,direction=1):
    #this is to zero pad the diagonals of the matrix to make it bigger
    #direction=1 is to pad the diagonal; direction=2 is to pad the reverse diagonal    
    if direction!=1:
        original_matrix = rotate_mat(original_matrix) 
    padded_matrix = np.zeros((new_size, new_size), dtype=float)
    # Populate the new matrix
    cur_size = len(original_matrix)
    dif_size = new_size-cur_size
    for i in range(cur_size):
        for j in range(cur_size):
            if i > j:  # Upper Left Triangle 
                padded_matrix[i + dif_size, j] = original_matrix[i][j]  # Shift one row down
            elif i==j: #shift the Diagonal
                # padded_matrix[i + dif_size, j] = original_matrix[i][j]  # Shift one row down
                padded_matrix[i, j + dif_size] = original_matrix[i][j]  # Shift one column right
            else:  # Bottom Right Triangle 
                padded_matrix[i, j + dif_size] = original_matrix[i][j]  # Shift one column right
    if direction!=1:
        padded_matrix = rotate_mat(padded_matrix,3)
    return padded_matrix

def nan_pad_mat_bone(original_matrix,new_size=4,direction=1,mid_window=15,time_window=20):
    #this is to zero pad the diagonals of the matrix to make it bigger
    #direction=1 is to pad the diagonal; direction=2 is to pad the reverse diagonal    
    if direction!=1:
        original_matrix = rotate_mat(original_matrix) 
    padded_matrix = np.full((new_size, new_size), np.nan)
    # Populate the new matrix
    cur_size = len(original_matrix)
    dif_size = new_size-cur_size
    thresh = time_window+mid_window
    for i in range(cur_size):
        for j in range(cur_size):
            if i >= j:  # Upper Left Triangle 
                if i>=time_window and j>=thresh:
                    padded_matrix[i+dif_size, j+dif_size] = original_matrix[i][j]
                elif i<time_window and j<time_window:
                    padded_matrix[i, j] = original_matrix[i][j]
                else:
                    padded_matrix[i + dif_size, j] = original_matrix[i][j]  # Shift n row down

            else:  # Bottom Right Triangle 
                if i<time_window and j<time_window:
                    padded_matrix[i, j] = original_matrix[i][j]
                elif i>=thresh and j>=thresh:
                    padded_matrix[i+dif_size, j+dif_size] = original_matrix[i][j]
                else:
                    padded_matrix[i, j + dif_size] = original_matrix[i][j]  # Shift n column right
    if direction!=1:
        padded_matrix = rotate_mat(padded_matrix,3)
    return padded_matrix

def mat_diag(in_mat,val=1):
    #defult fill in diagonal with 1
    mat = in_mat
    for i in range(len(mat)):  
        for j in range(len(mat[0])):
            if i==j:
                mat[i][j]=val    
    return mat
        
def mat_bound(in_mat,submat,val=-1):
    #submat = [20,24,21] #length of each submat, assume square, starts from 1, all submat adjacent 
    #val is the number to mark out the boundary
    import numpy as np
    bound = [i-1 for i in np.cumsum(submat)]
    mat = in_mat
    for i in range(len(bound)):
        if i==0:
            b = [0,bound[i]]
        else:
            b = [bound[i-1],bound[i]]
        rng = [i for i in range(b[0],b[1]+1)]
        for i in range(len(mat)):  
            for j in range(len(mat[0])):
                if i in b or j in b:
                    if i in rng and j in rng:
                        mat[i][j]=val
    return mat
            
def mat_info(in_mat,mute=False):
    #return: dim, shape, total #of elem, nonzero #of elem 
    import numpy as np
    info =[len(np.shape(in_mat)),np.shape(in_mat),np.size(in_mat),len(np.nonzero(in_mat)[0])]
    info_title = ['dim','shape','total items','nonzero items']
    if mute==False:
        mesg_dic = {'dim': info[0],'shape': info[1],'total elem#':info[2],'nonzero elem#': info[3]}
        for key, value in mesg_dic.items():
            print(key, ' : ', value)
    return info, info_title


def plot_figsize(wd,ht):
    import matplotlib.pyplot as plt
    plt.figure(figsize=(wd, ht))

def sns_style(style='default'):
    #style can be 'white', 'whitegrid','dark','darkgrid'
    import seaborn as sns
    if style=='default':
        sns.set()
    else:
        sns.set_style(style)

def plot_xylim(x_rng=False,y_rng=False):
    import matplotlib.pyplot as plt
    if x_rng!=False:
        plt.xlim((x_rng[0],x_rng[1]))
    if y_rng!=False:
        plt.ylim((y_rng[0],y_rng[1]))

def plot_save(outpath,ftype='.png'):
    import matplotlib.pyplot as plt
    plt.tight_layout()
    plt.savefig(outpath+ftype,dpi=300)
    print('figure saved.')

def plot_bar_se(df,x_var,y_var,kind='bar',alpha=.7,outpath=False,title=False,add=False):
    #this is to plot the tradditional simple bars with error-bars
    # example use: 
        #df = xf.assemble_dict([free,yoke,pasv],['free','yoke','pasv'],stack_keyname=['recall','cond'])
        #xf.plot_bar_se(df,x="cond", y="recall",outpath=name)   
    import seaborn as sns
    import matplotlib.pyplot as plt
    sns.set_theme(style="whitegrid")
    sns.catplot(data=df, x=x_var, y=y_var,kind=kind,alpha=alpha)
    #show and save figure
    if outpath!=False:
        if title!=False:
            plt.savefig(outpath+title+'.png')  
        else:
            plt.savefig(outpath+'.png')  
        print('Figure saved.')  
    if add==False:
        plt.show()
        print('Figure shown.')

def plot_bar(df,x_axis,y_axis,split_by,dodge=False,outpath=False,title=False,add=False):
    #this plots the bar either on top of each other (overlaid) or side by side (dodge)
    import seaborn as sns
    import matplotlib.pyplot as plt
    sns.barplot(data=df, x=x_axis, y=y_axis,hue=split_by,dodge=dodge)
    #show and save figure
    if outpath!=False:
        if title!=False:
            plt.savefig(outpath+title+'.png')  
        else:
            plt.savefig(outpath)  
        print('Figure saved.')  
    if add==False:
        plt.show()

def plot_bar_stack(df,x_axis,y_axis,split_by,bg_style='default',color='default',outpath=False,title=False,add=False):
    #the dataframe df needs to be have at least 3 columns: 1) the name of the varaible for x-axis, 2) the conditions to split bar by, 3) the count/stats for each variable's each condition
    #the default color palette only has 5 specified scale of blue, so the split_by cannot have more than 5 types
    import seaborn as sns
    import numpy as np
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    sns_style(bg_style)
    #get how many conditions there are to split the bar by
    bar_names,_ = get_unique_repeat([i for i in df[split_by]])
    if color=='default':
        color_palette = ['lightcyan','lightblue','lightskyblue','skyblue','darkblue']
        bar_color=color_palette[-len(bar_names):]
    # plot the bars stacked on top of each other
    total = df.groupby(x_axis)[y_axis].sum().reset_index()
    bar1 = sns.barplot(x=x_axis,  y=y_axis, data=total, color=bar_color[0])
    count=1
    for i in bar_names[1:]:
        opt = df[df[split_by]==i]
        sns.barplot(x=x_axis,  y=y_axis, data=opt, estimator=sum, ci=None,  color=bar_color[count])
        count+=1
    # add legend
    legend = []    
    for i in range(len(bar_names)):
        lgd = mpatches.Patch(color=bar_color[i], label=bar_names[i])
        legend.append(lgd)
    plt.legend(handles=legend)
    #show and save figure
    if outpath!=False:
        if title!=False:
            plt.savefig(outpath+title+'.png')  
        else:
            plt.savefig(outpath)  
        print('Figure saved.')  
    if add==False:
        plt.show()
        print('Figure shown.')

def plot_bar_sorted(mean_ls, se_ls, cond_ls, pval_ls,thresh=.05,figsize=(50,12),xlab='Conditions',ylab='Mean Value',cb_lab='Magnitude of Means',colormap='viridis', title='Bar Plot of Sorted Means with Standard Errors',fontsize=20,outpath=False,dpi=300,fig_format='.png',add=False):
    '''Plot bars with SE, with means sorted from max to min before plotting, marking out significance'''
    means,standard_errors,condition_names,p_values=mean_ls,se_ls,cond_ls,pval_ls
    # Combine the data for sorting
    data = list(zip(means, standard_errors, condition_names,p_values))
    # Sort data by means from largest to smallest
    sorted_data = sorted(data, key=lambda x: x[0], reverse=True)
    # Unpack the sorted data
    sorted_means, sorted_standard_errors, sorted_condition_names, sorted_p_values = zip(*sorted_data)
    # Normalize and map colors
    norm = plt.Normalize(min(sorted_means), max(sorted_means))
    colors = plt.cm.viridis(norm(sorted_means))
    # Create a very wide plot
    fig, ax = plt.subplots(figsize=(figsize[0], figsize[1]))  # Adjusted figure size to accommodate all labels
    bars = ax.bar(range(len(sorted_means)), sorted_means, yerr=sorted_standard_errors, color=colors, capsize=2)
    # Customize axes
    ax.set_xlabel(xlab, fontsize=fontsize, fontweight='bold')
    ax.set_ylabel(ylab, fontsize=fontsize, fontweight='bold')
    ax.set_title(title, fontsize=fontsize, fontweight='bold')
    # Set x-axis labels with all labels displayed
    ax.set_xticks(range(len(sorted_condition_names)))
    ax.set_xticklabels(sorted_condition_names, rotation=90, ha='right', fontsize=fontsize*((figsize[0]/figsize[1])/10))  # Rotate to 90 degrees
    ax.tick_params(axis='y', labelsize=fontsize)  # Increase the font size for y-axis values
    # Annotating bars with asterisks where p-values are < 0.05
    for i, (mean, p_val) in enumerate(zip(sorted_means, sorted_p_values)):
        if p_val < thresh:
            if mean >= 0:
                # Place asterisk above the bar for positive means
                ax.text(i, mean + sorted_standard_errors[i] + 0.1, '*', ha='center', va='bottom', color='red', fontsize=20)
            else:
                # Place asterisk below the bar for negative means
                ax.text(i, mean - sorted_standard_errors[i] - 0.1, '*', ha='center', va='top', color='red', fontsize=20)
    # Set the background color to light grey; use a darker grey for the grid to enhance contrast
    ax.set_facecolor('#f0f0f0')  
    ax.grid(True, linestyle='-', color='#bdbdbd', linewidth=0.5)
    # Adjust the y-axis limits to include space for asterisks
    ax.set_ylim([min(sorted_means) - max(sorted_standard_errors) - 0.5, max(sorted_means) + max(sorted_standard_errors) + 0.5])
    # Remove the axes edge color
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    # Color bar
    sm = plt.cm.ScalarMappable(cmap=colormap, norm=norm)
    sm.set_array([])
    cbar = plt.colorbar(sm, ax=ax)
    cbar.set_label(cb_lab, fontsize=fontsize, fontweight='bold')
    plt.tight_layout()
    #show and save figure
    if outpath!=False:
        plt.savefig(outpath+fig_format,dpi=dpi)  
        print('Figure saved.')  
    if add==False:
        plt.show()
        print('Figure shown.')
        
def is_title_too_long(title, fontsize, figure_width_inches, dpi):
    import math
    # Rough estimation of average character width in points for the given font size
    # This is a very rough estimate; actual values vary by font and character
    avg_char_width_in_points = fontsize # Adjust this based on your font
    # Convert title length to points
    title_width_points = len(title) * avg_char_width_in_points
    # Convert figure width to points (inches to points)
    figure_width_points = figure_width_inches * dpi    
    # Check if the title is too long
    a = title_width_points > figure_width_points
    # Check how many lines are needed
    b = math.ceil(title_width_points/figure_width_points) 
    return a,b

def split_title_into_lines(title, n_lines):
    if n_lines < 2:
        return title  # No need to split if less than 2 lines are requested
    # Replace underscores with spaces for splitting
    words = title.replace('_', ' ').split()
    # Calculate roughly how many words per line
    words_per_line = len(words) // n_lines
    lines = []
    for i in range(n_lines - 1):
        line = ' '.join(words[i*words_per_line:(i+1)*words_per_line])
        lines.append(line)
    # Add the remaining words to the last line
    lines.append(' '.join(words[(n_lines-1)*words_per_line:]))
    # Join the lines with newline characters
    return '\n'.join(lines)

def plot_boxes(cond_ls,val_ls,figsize='default',bg_style='whitegrid',title='boxplots',box_width=0.5,box_outline_width=3,swarm=True,swarm_size=6,fontsize=20,xylab=['Condition','Value'],outpath=False,fig_format='.png',dpi=100,add=False):
    #plot the boxes of values split by condition, with swarm of values in each box
    import seaborn as sns
    import matplotlib.pyplot as plt
    import pandas as pd
    
    figout = outpath+title+fig_format
    sns_style(bg_style)
    conditions,values = cond_ls,val_ls
    # Create a DataFrame
    xlab,ylab = xylab[0],xylab[1]
    data = pd.DataFrame({xlab: conditions, ylab: values})
    # Map conditions to numeric values to control the x-axis limits more easily
    data['ConditionCode'] = data[xlab].astype('category').cat.codes
    # Create the plot
    if figsize=='default':
        fig_width = 2*(max(data['ConditionCode'])+2)
        print(fig_width)
        figsize = (fig_width,8)
    plt.figure(figsize=figsize)
    box_plot = sns.boxplot(x='ConditionCode', y=ylab, data=data, palette="vlag", width=box_width,linewidth=box_outline_width)
    #swarm overlay
    if swarm==True:
        sns.swarmplot(x='ConditionCode', y=ylab, data=data, color='lightblue', edgecolor='black',size=swarm_size, linewidth=1)
    means = data.groupby('ConditionCode')[ylab].mean().reset_index()
    sns.pointplot(x='ConditionCode', y=ylab, data=means, join=False, color='gold', markers='D', scale=1.5, ci=None)
    # Put title to multiple lines if needed:
    a,b= is_title_too_long(title,fontsize,fig_width,dpi)
    if a:
        print(b)
        title = split_title_into_lines(title,b)
    # Set plot title and labels with specified font sizes
    box_plot.set_title(title, fontsize=fontsize*3/2,pad=30)
    box_plot.set_xlabel(xlab,fontsize=fontsize*5/4)
    box_plot.set_ylabel(ylab,fontsize=fontsize*5/4)
    # Adjust tick labels to show condition names instead of codes
    box_plot.set_xticklabels(data[xlab].unique())
    # Rotate and set font size for x-axis labels
    plt.xticks(rotation=45, fontsize=fontsize)
    # Set font size for y-axis labels
    plt.yticks(fontsize=fontsize)
    # Open plot on top and right side, and highlight left and bottom spines
    sns.despine()
    ax = plt.gca()
    ax.spines['left'].set_color('black')
    ax.spines['left'].set_linewidth(2)
    ax.spines['bottom'].set_color('black')
    ax.spines['bottom'].set_linewidth(2)
    # Explicitly set the x-axis limits to include -1
    plt.xlim(-1, max(data['ConditionCode'])+1)
    std = data[ylab].std()
    plt.ylim(min(data[ylab])-std/2, max(data[ylab])+std/2)
    plt.tight_layout(pad=2.0)
    #save and show figure
    if outpath!=False:
        mkdir(outpath)    
        plt.savefig(figout, dpi=dpi, bbox_inches='tight')
    if add==False:
        plt.show()

def plot_label(x_coords,y_coords,labels='default',color='black',font_size=15,xloc=-1,yloc=0.1,add=False,title=False,outpath=False):
    import matplotlib.pyplot as plt
    if labels=='default':
        labels = [str(i+1) for i in range(len(x_coords))]
    elif type(labels)==str:
        labels = [labels+str(i+1) for i in range(len(x_coords))]
    for i in range(len(x_coords)):
        x,y,lab = x_coords[i],y_coords[i],labels[i]
        plt.text(x+xloc,y+yloc,lab,size=font_size,color=color)
    #show and save figure
    if outpath!=False:
        if title!=False:
            plt.savefig(outpath+title+'.png')  
        else:
            plt.savefig(outpath)  
        print('Figure saved.')  
    if add==False:
        plt.show()
    print('label plotted')

def plot_line(y,x='Default',y_lab='None',x_lab='None',legend='Default',title='None',start=1,add=False,outpath=False,ylim=False,dpi=100):
    #y: the y-axis is a list of numbers; or a list of lists of numbers
    #x: the x-axis is a list of numbers; or a list of lists of numbers;
       #Defaults to be natural numbers starting from 0
    #x_lab, y_lab: label for x- and y- axis, defaults to none
    #legend: defaults to line#, otherwise, specify each line's name in order of y-axis
    #title: titile of the plot defaults to none
    #start: x_axis if not specified, defaults to start from 0, but can also start at 1 or other number
    #example:
        #plot_line(y=y_axis,x=[x_axis]*len(y_axis),legend=[i for i in data.keys()[1:]])
    import matplotlib.pyplot as plt
    import numpy as np
    #get x-axis
    x_axis = []
    if x =='Default':
        if type(y[0])==float or type(y[0])==int or type(y[0])==np.float64 or type(y[0])==np.float32:
            x_axis = [i for i in range(0+start,len(y)+start)]
        else: #y is a list of lists 
            for lst in y:
                x_axis.append([i for i in range(0+start,len(lst)+start)])
    else:
        x_axis = x
    #start the plot
    if type(y[0])==float or type(y[0])==int or type(y[0])==np.float64 or type(y[0])==np.float32:
        plt.plot(x_axis,y,label=legend)
        if ylim!=False:
            plt.ylim(ylim[0], ylim[1]) 
        if title !='None':
            plt.title(title,fontsize=18)
        if x_lab !='None':
            plt.xlabel(x_lab,fontsize=18)
        plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.) 
    else:
        for ind in range(len(y)):
            if legend =='Default':
                name = 'line'+str(ind+1)
            else:
                name = legend[ind]
            plt.plot(x_axis[ind],y[ind],label=name)
        if ylim!=False:
            plt.ylim(ylim[0], ylim[1]) 
        plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.) 
        if x_lab !='None':
            plt.xlabel(x_lab, fontsize=18)
        if y_lab != 'None':
            plt.ylabel(y_lab, fontsize=18)
        if title != 'None':
            plt.title(title, fontsize=18)
    if outpath!=False:
        plt.tight_layout()
        plt.savefig(outpath,dpi=dpi)  
        print('Figure saved.')    
    if add==False:
        plt.show()
    print('Figure plotted.')

def plot_anyline(x,y, figsize=(12, 6),xlab='x-axis: list-A',ylab='y-axis: list-B',title='Line Plot',add=False,outpath=False,fig_format='.png',pad=0.2):
    """
    Plots a line plot using list_A for the x-axis labels and list_B for the y-axis values.
    Handles non-numeric and repeated x-axis labels by creating a unique identifier for each data point.
    np.nan values in list_B will be shown as breaks in the line.

    Parameters:
    list_A (list): The list containing x-axis labels which may include non-numeric and repeated elements.
    list_B (list): The list containing y-values corresponding to list_A, can include np.nan.
    figsize (tuple): Figure size for the plot, default is (12, 6).
    """
    list_A,list_B = x,y
    # Creating a DataFrame
    df = pd.DataFrame({'X': list_A, 'Y': list_B})
    # Generate unique identifiers for X to avoid confusion in repeated labels
    df['X_ID'] = range(len(df))
    # Plotting
    plt.figure(figsize=figsize)
    plt.plot(df['X_ID'], df['Y'], marker='o', linestyle='-', drawstyle='default')
    # Set custom x-ticks
    plt.xticks(ticks=df['X_ID'], labels=df['X'])
    # Set labels and titles
    plt.xlabel(xlab)
    plt.ylabel(ylab)
    plt.title(title)
    plt.grid(True)
    # Optional: Rotate the labels if they overlap
    plt.xticks(rotation=45)
    # Adjusting y-axis limits to add some padding
    y_min, y_max = df['Y'].min(), df['Y'].max()
    if pd.notna(y_min) and pd.notna(y_max):  # Ensure there are non-NaN values to calculate min/max
        y_range = y_max - y_min
        padding = pad * y_range  # 10% padding on each end
        plt.ylim(y_min - padding, y_max + padding)
    if outpath!=False:
        plt.savefig(outpath+title+fig_format)  
    if add==False:
        plt.show()

def plot_anylines(x, y_lists, legend_labels, legend_pos_fixed='upper right',figsize=(12, 6), xlab='x-axis: list-A', ylab='y-axis: list-B', title='Line Plot', add=False, outpath=False, fig_format='.png', pad=0.2):
    """
    Plots multiple line plots using a single list for the x-axis and multiple lists for the y-axis, with legends.
    Parameters:
    x (list): The list containing x-axis labels which may include non-numeric and repeated elements.
    y_lists (list of lists): Each sublist contains y-values corresponding to x, can include np.nan.
    legend_labels (list of str): Labels for each line in the legend.
    figsize (tuple): Figure size for the plot, default is (12, 6).
    xlab (str): Label for the x-axis.
    ylab (str): Label for the y-axis.
    title (str): Title of the plot.
    add (bool): If True, the plot will be added to an existing figure, default is False.
    outpath (str): Path to save the figure, if False no file is saved.
    fig_format (str): File format for saving the figure, default is '.png'.
    pad (float): Padding ratio for the y-axis limits, default is 0.2.
    """
    plt.figure(figsize=figsize)
    # Creating a DataFrame from x-axis labels
    df = pd.DataFrame({'X': x})
    # Generate unique identifiers for X to avoid confusion in repeated labels
    df['X_ID'] = range(len(df))
    # Plotting each y list with corresponding legend label
    for y, label in zip(y_lists, legend_labels):
        df['Y'] = y
        plt.plot(df['X_ID'], df['Y'], marker='o', linestyle='-', drawstyle='default', label=label)
    # Set custom x-ticks
    plt.xticks(ticks=df['X_ID'], labels=x)
    # Set labels and titles
    plt.xlabel(xlab)
    plt.ylabel(ylab)
    plt.title(title)
    plt.grid(True)
    # Optional: Rotate the labels if they overlap
    plt.xticks(rotation=45)
    # Adding legend
    if legend_pos_fixed!=False:
        plt.legend(loc=legend_pos_fixed)
    else:
        plt.legend()
    # Adjusting y-axis limits to add some padding
    all_y = np.concatenate(y_lists)  # combine all y lists
    y_min, y_max = np.nanmin(all_y), np.nanmax(all_y)
    if np.isfinite(y_min) and np.isfinite(y_max):  # Ensure there are finite values to calculate min/max
        y_range = y_max - y_min
        padding = pad * y_range  # Padding on each end
        plt.ylim(y_min - padding, y_max + padding)
    if outpath:
        plt.savefig(outpath + title + fig_format)
    if not add:
        plt.show()

def plot_vline(x_intercept,y_range,add=False,outpath=False,title='lineplot_output',lwd=2,color='black',ltype='solid'):
    import matplotlib.pyplot as plt
    y_axis = get_range_float(rng=y_range)
    if type(x_intercept)==list:
        for x in x_intercept:
            x_axis = [x]*len(y_axis)
            plt.plot(x_axis,y_axis,linewidth=lwd,color=color,linestyle=ltype)
    else:
        x_axis = [x_intercept]*len(y_axis)
        plt.plot(x_axis,y_axis,linewidth=lwd,color=color,linestyle=ltype)
    if outpath!=False:
        plt.savefig(outpath+title+'.png') 
    if add==False:
        plt.show()
    print('Figure plotted.')

def plot_hline(y_intercept,x_range,add=False,outpath=False,title='lineplot_output',lwd=2,color='black',ltype='solid'):
    #ltype = 'dashed','dotted','dashdot','solid'
    import matplotlib.pyplot as plt
    x_axis = get_range_float(rng=x_range)
    if type(y_intercept)==list:
        for y in y_intercept:
            y_axis = [y]*len(x_axis)
            plt.plot(x_axis,y_axis,linewidth=lwd,color=color,linestyle=ltype)
    else:
        y_axis = [y_intercept]*len(x_axis)
        plt.plot(x_axis,y_axis,linewidth=lwd,color=color,linestyle=ltype)
    if outpath!=False:
        plt.savefig(outpath+title+'.png')  
    if add==False:
        plt.show()
    print('Figure plotted.')

def plot_shade(rng,direction='vertical',add=False,outpath=False,title='plot',alpha=.15,lw=0,color='grey',x='default'):
    #if direction is vertical, rng is x_intercept range, e.g. rng=(0,10), shading the area between x=0 and x=10
    import matplotlib.pyplot as plt

    if type(rng[0])==list and type(rng[1])==list:
        y1,y2 = rng[0],rng[1]
        if x=='default':
            x = [i for i in range(len(y1))]
        plt.fill_between(x, y1, y2,alpha=alpha,color=color)
    else: 
        if direction=='vertical':
            plt.axvspan(rng[0],rng[1], color=color, alpha=alpha, lw=lw)
        else:
            plt.axhspan(rng[0],rng[1], color=color, alpha=alpha, lw=lw)
    if outpath!=False:
        plt.savefig(outpath+title+'.png')  
    if add==False:
        plt.show()
    print('Figure plotted.')

def plot_shades(rng_ls,direction='vertical',add=False):
    #e.g., rng_ls = [(20,35),(20,40),(42,62),(47,62)]
    for i in range(len(rng_ls)):
        ls = rng_ls[i]
        if i!=len(rng_ls)-1:
            add_not=True
        else:
            add_not=add
        plot_shade(ls,direction=direction,add=add_not)

def plot_line_pairdot(array1,array2,xlab=['array1','array2'],title=['',20],xylim=[(-0.5,1.5),(-0.15,0.6)],xylab=['',''],xylab_size=(30,30),xytick_size=(20,20),dot_size=70,dot_c='b',fig_hw=(4,8),fig_size=80,lwd=0.3,outpath=False,add=False):
    import matplotlib.pyplot as plt
    from matplotlib.pyplot import figure
    import numpy as np
    plt.figure(figsize=fig_hw, dpi=fig_size)
    plt.rc('xtick', labelsize=xytick_size[0]) 
    plt.rc('ytick', labelsize=xytick_size[1]) 
    # plotting the points
    plt.scatter(np.zeros(len(array1)),array1,c=dot_c, s=dot_size)
    plt.scatter(np.ones(len(array2)), array2,c=dot_c, s=dot_size)
    # plotting the lines
    for i in range(len(array1)):
        plt.plot( [0,1], [array1[i], array2[i]],linewidth=lwd)
    plt.xticks([0,1], [xlab[0], xlab[1]])
    plt.xlabel(xylab[0],fontsize=xylab_size[0])
    plt.ylabel(xylab[1],fontsize=xylab_size[1])
    plt.xlim(xylim[0])
    plt.ylim(xylim[1])
    plt.suptitle(title[0],fontsize=title[1])
    if outpath!=False:
        if title[0]!='':
            title= '_'+title[0]
        else:
            title=title[0]
        plt.savefig(outpath+title, dpi=fig_size)  
    if add==False:
        plt.show()

def plot_dist(array,outpath=False,legend=False,mean=False,median=False,base=False,rug=False,y_rng=(0,1),bg_style='default',add=False):
    #example use to plot two overlaid distributions:
        # init_l=[12,13,15,15,16,71,12,4,34,24]
        # init_h=[243,252,136,124,125,224,112,124,124,512]
        # xf.plot_figsize(20,12)
        # xf.plot_dist(init_l,y_rng=(0,0.1),mean=False)
        # xf.plot_dist(init_h,y_rng=(0,0.1),mean=False)
    import seaborn as sns
    import numpy as np
    import matplotlib.pyplot as plt
    if bg_style=='default':
        sns.set()
    else:
        sns_style(bg_style)
    ax = sns.distplot(array,rug=rug,label=legend)
    if mean ==True:
        avg = np.nanmean(array)
        if median ==False and base==False:
            plot_vline(avg,y_rng,add=add)
        else:
            plot_vline(avg,y_rng,add=True)
        print('Mean = '+str(avg))
    if median==True:
        med = np.median(array)
        if base==False:
            plot_vline(med,y_rng,add=add)
        else:
            plot_vline(med,y_rng,add=True)
        print('Median = '+str(med))
    if base!=False:
        if base==True: #assume base = 0
            base = 0
        plot_vline(base,y_rng)
    if legend!=False:
        plt.legend(loc='upper right')
    if outpath!=False:
        mkdir(outpath)
        ax.figure.savefig(outpath+'_dist.png')
    if add==False:
        plt.show()

def plot_hist(array,color=False,title=False,xylab=False,outpath=False,legend=False,bins=15,fontsize=20,bg_style='default',alpha=0.5,add=False):
    import seaborn as sns
    import matplotlib.pyplot as plt
    import numpy as np
    if bg_style=='default':
        sns.set()
    else:
        sns_style(bg_style)
    if color==False:
        color='grey'
    sns.histplot(array, bins=bins, color=color, label=legend, alpha=alpha)
    sns.kdeplot(array, color=color)# Plot KDE (estimate curve) 
    if legend!=False:
        plt.legend(loc='upper right')
    if title!=False:
        plt.title(title,fontsize=fontsize)
    if xylab!=False:
        plt.xlabel(xylab[0],fontsize=fontsize)
        plt.ylabel(xylab[1],fontsize=fontsize)
    if outpath!=False:
        mkdir(outpath)
        plot_save(outpath+'_dist.png')
    if add==False:
        plt.show()

def plot_hist_with_value(value, value_distribution, bins=50,title='default',outpath=False,fig_format='.png',dpi=300,add=False):
    import seaborn as sns
    from scipy.stats import gaussian_kde
    # Calculate the proportion of values in the dist <= value
    proportion = sum(1 for x in value_distribution if x <= value) / len(value_distribution)
    # Create the histogram
    plt.figure(figsize=(10, 6))    
    # Plot the histogram
    counts, bin_edges, patches = plt.hist(value_distribution, bins=bins, color='lightblue', alpha=0.7, label='Histogram')
    # Find the bin that contains the value and shade up to this bin
    for bin_start, bin_end, patch in zip(bin_edges[:-1], bin_edges[1:], patches):
        if bin_start <= value:
            patch.set_facecolor('grey')
        else:
            break
    # Plot the vertical line at the given value
    plt.axvline(value, color='red', linestyle='--', label=f'Value = {value}')    
    # Overlay a KDE curve above the histogram
    kde = gaussian_kde(value_distribution)
    x_values = np.linspace(min(value_distribution), max(value_distribution), 1000)
    kde_values = kde(x_values)
    # Normalize KDE to match the scale of the histogram
    scale_factor = len(value_distribution) * (bin_edges[1] - bin_edges[0])
    kde_values_scaled = kde_values * scale_factor
    plt.plot(x_values, kde_values_scaled, color='blue', label='KDE Curve')
    # Annotate the proportion on the plot
    max_count = counts.max()
    plt.text(value, max_count / 2, f'Proportion = {proportion:.4f}', fontsize=12, color='black', ha='right', va='center', bbox=dict(facecolor='white', alpha=0.5))
    # Set the background color to white and enable the grid
    plt.gca().set_facecolor('white')
    plt.grid(True, which='both', linestyle='--', linewidth=0.5)   
    # Labels and title
    plt.xlabel('Value')
    plt.ylabel('Count per Bin')
    if title=='default':
        plt.title('Distribution of '+str(len(value_distribution))+' trials against actual value',fontsize=20)
    else:
        plt.title(title,fontsize=20)
    plt.legend()
    if outpath!=False:
        mkdir(outpath)
        plt.savefig(outpath+fig_format,dpi=dpi)
    if add==False:
        plt.show()
    p_val =proportion
    return p_val



def plot_2dmat(mat,xylab=False,title='ISC_Time-Time-Corr',fig_format='.png',fontsize=20,figsize=(12,10),outpath=False,save_matrix=False,add=False):
    if len(shape(mat))==1:#1d array needs to be turned into a 2d array
        matrix = np.array(mat).reshape(1, -1)
    else:
        matrix = np.array(mat)
    plot_figsize(figsize[0],figsize[1])
    plt.imshow(matrix, cmap='viridis', aspect='auto')
    plt.colorbar()
    if xylab!=False:
        plt.xlabel(xylab[0],fontsize=fontsize)
        plt.ylabel(xylab[1],fontsize=fontsize)
    plt.title(title,fontsize=fontsize)
    mesg = 'Figure '
    if outpath!=False:
        mkdir(outpath)
        plt.savefig(outpath+title+'_mat'+fig_format)        
        if save_matrix!=False:
            np.savetxt(outpath+title+'_mat.csv', matrix, delimiter=",")
            mesg = mesg+'and matrix '
        mesg = mesg+ 'saved '
    if add==False:
        plt.show()
    mesg = mesg+'shown.'
    print(mesg)

def plot_matrix(mat,title=False,outpath=False,xylab=False,labname=False,diagmask=False,map_color="YlGnBu",minmax=(None,None),fontsize=20,save_matrix=False,add=False):
    #mat: matrix needs to be np.array with 2 dimentions with equal length
    #title: for the plot and the name saved to outpath
    #labname: individual label name, e.g. 'event'
    #diagmask: whether to mask out the upper half of the matrix or not
    #map_color: 
        # "icefire" - (min) blue/black/red (max)
        # "Spectral" - (min) red/yellow/blue (max)
        # "coolwarm" - (min) blue/red (max)
        # 'RdYlBu_r' - (min) blue/yellow/red (max)
        # "YlOrBr" - (min) yellow/red (max)
        # 'CMRmap_r' - (min) yellow/black (max) or use 'CMRmap' for the reversed color-palette
        # 'viridis'
    # example use:
        # matrix.plot(npmat,'causal_matrix','./results/graphs/',True)
        # this will show and save the matrix plot to specified path
    import seaborn as sns
    import matplotlib.pyplot as plt
    import numpy as np
    mesg=''
    if type(mat)!=np.ndarray:
        mat = convert_list_array(mat)
    if mat.ndim !=2:
        mesg='ERROR: input matrix needs to be 2D'
    else:
        if mat.shape[0]!=mat.shape[1]:
            mesg = 'ERROR: input matrix length not equal to width'
        else:
            length = len(mat)                
            mask = np.zeros_like(mat)
            mask[np.triu_indices_from(mask)] = diagmask
            with sns.axes_style("white"):
                if labname!=False and type(labname)==str:
                    #create individual labels 
                    labels = []
                    for i in range(0,length):
                        lab = labname+str(i+1)
                        labels.append(lab)
                    ax = sns.heatmap(mat, mask=mask, square=True, vmin=minmax[0],vmax=minmax[1], cmap=map_color,xticklabels=labels,yticklabels=labels)  
                elif labname!=False and type(labname)==list:
                    ax = sns.heatmap(mat, mask=mask, square=True, vmin=minmax[0],vmax=minmax[1], cmap=map_color,xticklabels=labname,yticklabels=labname)  
                else:
                    ax = sns.heatmap(mat, mask=mask, square=True, vmin=minmax[0],vmax=minmax[1], cmap=map_color,xticklabels=False,yticklabels=False)
                
                if xylab!=False:
                    plt.xlabel(xylab[1],fontsize=fontsize)
                    plt.ylabel(xylab[0],fontsize=fontsize)

                if title!=False:
                    ax.set_title(title,fontsize=fontsize)
                    if outpath!=False:
                        mkdir(outpath)
                        ax.figure.savefig(outpath+title+'_mat.png')
                        if save_matrix!=False:
                            np.savetxt(outpath+title+'_mat.csv', mat, delimiter=",")
                            mesg = 'Figure and mat saved for '+title
                else:
                    if outpath!=False:
                        mkdir(outpath)
                        ax.figure.savefig(outpath+'_mat.png')
                        mesg = 'Figure '
                        if save_matrix!=False:
                            np.savetxt(outpath+'_mat.csv', mat, delimiter=",")
                            mesg = mesg+'and matrix '
                        mesg = mesg+'saved.'
                if add==False:
                    plt.show()
                    mesg='Figure shown. '+mesg 
                else:
                    mesg='plotmat complete.'          
    return mesg

def plot_dots(x,y,figsize=(10,8),labels=False,fontsize=15,marker='o',markerfacecolor='blue',markersize=12,linestyle='none', linecolor='green',add=False,outpath=False,title='scatter'):
    #linestyle (will connect the dots by sequence): [‘solid’ | ‘dashed’, ‘dashdot’, ‘dotted’ | (offset, on-off-dash-seq) | '-' | '--' | '-.' | ':' | 'None' | ' ' | '']
    import matplotlib.pyplot as plt
    from matplotlib.pyplot import plot 
    if figsize!=False:
        plot_figsize(figsize[0],figsize[1])
    plot(x, y, color=linecolor, marker=marker,linestyle=linestyle, markerfacecolor=markerfacecolor, markersize=markersize)
    if labels!=False:
        for lab, x1, y1 in zip(labels,x,y):
            plt.annotate(lab, xy=(x1,y1), xytext=(0,5), textcoords="offset points",size=fontsize)
    if outpath!=False:
        plt.savefig(outpath+str(title)+'.png')  
    if add==False:
        plt.show()


def annotate3D(ax, s, *args, **kwargs):
    '''add anotation text s to to Axes3d ax'''
    from mpl_toolkits.mplot3d.proj3d import proj_transform
    from matplotlib.text import Annotation
    
    class Annotation3D(Annotation):
        def __init__(self, s, xyz, *args, **kwargs):
            Annotation.__init__(self,s, xy=(0,0), *args, **kwargs)
            self._verts3d = xyz        

        def draw(self, renderer):
            xs3d, ys3d, zs3d = self._verts3d
            xs, ys, zs = proj_transform(xs3d, ys3d, zs3d, renderer.M)
            self.xy=(xs,ys)
            Annotation.draw(self, renderer)

    tag = Annotation3D(s, *args, **kwargs)
    ax.add_artist(tag)

    
def plot_dots_3d(x,y,z,figsize=(10,8),labels=False,markersize=200,markercolor='Green',cmap='Greens',edgecolor='black', linestyle='none',linecolor='grey',outpath=False,add=False,title='scatter3d'):
    #markercolor can be a variable, e.g. =x
    #cmap: 'Blues','Greens','Reds',etc. 
    #linestyle (will connect the dots by sequence): [‘solid’ | ‘dashed’, ‘dashdot’, ‘dotted’ | (offset, on-off-dash-seq) | '-' | '--' | '-.' | ':' | 'None' | ' ' | '']
    #e.g. plot_dots_3d(x,y,z,markercolor=x,cmap='Reds',linestyle='dashed')
    from mpl_toolkits import mplot3d
    import matplotlib.pyplot as plt
    if figsize!=False:
        plot_figsize(figsize[0],figsize[1])
    ax = plt.axes(projection='3d')
    if type(markercolor)==str:
        ax.scatter3D(x,y,z, s=markersize,c=markercolor,edgecolor=edgecolor)
    else:
        ax.scatter3D(x,y,z, s=markersize,c=markercolor,cmap=cmap,edgecolor=edgecolor)
    if linestyle!='none':
        ax.plot3D(x,y,z, linecolor,linestyle=linestyle)
    if labels!=False:
        for i,lab in enumerate(labels): 
            annotate3D(ax, s=lab, xyz=(x[i],y[i],z[i]), fontsize=10, xytext=(-5,5),textcoords='offset points', ha='right',va='bottom')  
    if outpath!=False:
        plt.savefig(outpath+str(title)+'.png')  
    if add==False:
        plt.show()
    

def plot_scatter(array1,array2,cal_corr=True,outpath=False,fig_format='.png',make_plot=True,title=False,xlab=False,ylab=False,pointlab=False,mute=False,thresh=30):
    #x-axis: array1 or list1
    #y-axis: array2 or list2
    #outpath: the path to save the scatterplot
    #title: the title of the plot, also the name of the saved plot
    #xlab: label for x-axis
    #ylab: label for y-axis
    #example use:
        # corr_scatter(list1,array2) 
        # this will show the scatter plot in jupyter notebook
        # corr_scatter(list1,array2,outpath='./data/results/graphs/',title='sub1_sct',xlab='array1',ylab='array2',pointlab='event')
        # this will show and save the plot as specified (with tiles for the graph, the x and y-axis, and every dot labeled)
    import matplotlib.pyplot as plt
    import scipy.stats
    import numpy as np
    if cal_corr!=False:
        corr,pval = get_corr(array1,array2,skip_nan_min=thresh)
    if make_plot==True:
        fig, ax = plt.subplots()
        ax.scatter(array1,array2)
        if pointlab!=False:
            pointlab = str(pointlab)
            lab = []
            for i in range(0,len(array1)):
                lab.append(pointlab+str(i))
            for i, txt in enumerate(lab):
                ax.annotate(txt, (array1[i], array2[i]))
        if title!=False and cal_corr!=False:
            plt.title(title+':\n pearson r= '+ str(corr)+', p-value='+str(pval));
        elif cal_corr!=False:
            plt.title('pearson r= '+ str(corr)+', p-value='+str(pval));
        if xlab!=False:
            plt.xlabel(str(xlab));
        if ylab!=False:
            plt.ylabel(ylab);
        z = np.polyfit(array1,array2,1)
        p = np.poly1d(z)
        plt.plot(array1,p(array1),"r-")
        if title!=False and outpath!=False:
            mkdir(outpath)
            plt.savefig(outpath+str(title)+fig_format)  
            mesg = 'Figure saved for '+title
            if mute !=True:
                print(mesg)
        else:
            print('Scatter plot did not save due to no title or outpath available')
    if cal_corr!=False:
        return corr,pval
    else:
        return 'scatter plotted.'

def plot_swarm(data_as_dict,y,x,outpath=False,fig_format='.png',hue=False,size=7,alpha=1,edgecolor='black',dodge=False,add=False):
    #example use (overlay with violin plot):
        # xf.plot_figsize(12,10) #to set figure size
        # xf.plot_violin(data,y='recall',x='cond',linewidth=2,bandwidth=.3,inner='quartile') #to plot violin
        # xf.plot_swarm(data,y='recall',x='cond',size=16) #to overlay the dots
    import seaborn as sns
    import matplotlib.pyplot as plt
    if type(data_as_dict)==dict:
        df = dict_to_df(data_as_dict,colname_by_key=True)
    else:
        df = data_as_dict
    if hue==False:
        ax = sns.swarmplot(x=x, y=y, data=df, color='black', size=size,alpha=alpha, edgecolor=edgecolor,dodge=dodge)
    else:
        ax = sns.swarmplot(x=x, y=y, data=df, hue=hue,size=size,alpha=alpha, edgecolor=edgecolor,dodge=dodge)
    if outpath!=False:
        mkdir(outpath)
        ax.figure.savefig(outpath+'_violin'+fig_format)
        print('figure saved.')
    if add==False:
        plt.show()
    print('Swarm plot generated.')

def plot_violin(data_as_dict,y,x,outpath=False,fig_format='.png',color=False,split=False,cut=3,linewidth=None,inner='box',swarm=False,bandwidth=False,palet='muted',alpha=1,add=False):
    #data_as_dict has all the variables needed, e.g. the continuous y, the categorical condition x
    #y is the continuous variable to be ploted
    #x is the condition to be split by
    #outpath is the path and figure name to save the figure as, e.g. './figname'
    #color is another condition for y to be further split by (main contrast)
    #split is to have the compared colors take half of the violin 
    #inner: inner{“box”, “quartile”, “point”, “stick”, None}, optional
    #bandwidth is between 0-1 to indicate how much to bundle (smooth)
    #palet is the palette of the violin plot: can also be 'Set2', 'Set3', 'Accent', etc.
    import seaborn as sns
    import matplotlib.pyplot as plt
    #sns.set_theme(style="whitegrid")
    if type(data_as_dict)==dict:
        df = dict_to_df(data_as_dict,colname_by_key=True)
    else:
        df = data_as_dict
    if bandwidth ==False:
        bdwd = .5
    else: 
        bdwd = bandwidth
    if color!=False:
        ax = sns.violinplot(x=x, y=y, hue=color, data=df, cut=cut,linewidth=linewidth,palette=palet,inner=inner,split=split, bw=bdwd,alpha=alpha)
    else:
        ax = sns.violinplot(x=x, y=y, data=df, cut=cut,linewidth=linewidth,palette=palet,split=split,inner=inner, bw=bdwd, alpha=alpha) 
    if swarm==True:
        plot_swarm(data_as_dict,y,x)  
    if outpath!=False:
        mkdir(outpath)
        if outpath[-1]!='/':
            outname = outpath+'_violin'+fig_format
        elif '.' in outpath.split('/')[-1]:
            outname = outpath
        else:
            outname = outpath+fig_format
        ax.figure.savefig(outname)
        print('figure saved.')
    if add==False:
        plt.show()
    print('Violin plot generated.')

def plot_autocorr_halfmax(lag_corr,halfmax_point,fontsize=20,outpath=False,bg_style='seaborn-whitegrid',title='Auto-correlation_half-maximum',dpi=100,fig_format='.png',add=False):
    """
    Plots a vector of y-values against their index, a vertical line at x= halfmax point,
    and two horizontal lines at y=0, with a half-open frame style.
    
    Parameters:
    y_values (list of float): The y-values to be plotted.
    """
    y_values,x = lag_corr,halfmax_point
    x_values = np.arange(len(y_values))  # Generate x-values from 0 to 30
    plt.figure(figsize=(9, 6))  # Set figure size
    plt.style.use(bg_style)
    ax = plt.gca()  # Get current axes
    # Plot the y-values against x-values
    plt.plot(x_values, y_values, label='Auto correlation vs. Unit of shift')
    # Plot a vertical line at x=4
    if x!=None: #if the halfmax exists:
        plt.axvline(x=x, color='r', linestyle='--', label='Shift = '+str(halfmax_point))
    # Plot horizontal lines at y=0
    plt.axhline(y=0, color='black', linestyle='-')
    plt.axhline(y=0.5, color='red', linestyle='-')
    # Customize the frame to be half open
    ax.spines['top'].set_visible(False)  # Hide the top spine
    ax.spines['right'].set_visible(False)  # Hide the right spine
    ax.spines['left'].set_color('black')  # Set the left spine color to black
    ax.spines['bottom'].set_color('black')  # Set the bottom spine color to black
    # Add legend
    plt.legend(fontsize=fontsize*3/4)
    # Add titles and labels
    plt.title(title,fontsize=fontsize)
    plt.xlabel('Unit of shift',fontsize=fontsize)
    plt.ylabel('Lag correlation values',fontsize=fontsize)
    # Set the x-axis range from 0 to the length of lag correlation values
    plt.xlim(0, len(lag_corr))
    # Output figure
    if outpath!=False:
        plt.savefig(outpath+title+fig_format,dpi=dpi)
    # Show the plot
    if add==False:
        plt.show()



def calc_onset_from_dur(inlist,start_time=0):
    #for example: list1 = [1,2,1,4,3] each being a duration in seconds, 
    #the function returns [0,1,3,4,8,11], i.e. the onset of each item assuming start time to be 0s
    import numpy as np
    if start_time=='none':
        out = []
        start_time = 0
    else:
        out = [start_time]
    for i in range(1,len(inlist)+1):
        curr = inlist[:i]
        out.append(start_time+np.sum(curr))
    return out

def aud_dur(aud_name,unit='s'):
    #this is to calculate the duration/length of the audio wav file 
    #unit='s' returns duration in the unit of seconds, if 'ms' returns milliseconds
    import wave
    import contextlib
    with contextlib.closing(wave.open(aud_name,'r')) as f:
        frames = f.getnframes()
        rate = f.getframerate()
        duration = frames / float(rate)
        if unit=='ms':
            duration =round(duration*1000)
    return duration

def aud_sound_time(audfile,decimal=0.01,thr=0.1):
    #get the accurate (to 0.01s, i.e. 10ms) time points that has actual sound (sound signal>=thr) 
    time = []
    vec = aud_env(audfile,tr = decimal)
    counter = 0
    for i in vec:
        counter+=1
        if i>=thr:
            time.append(counter*decimal)
    return time

def aud_start_end(audfile,decimal=0.001,thr=0.1):
    #get the accurate onset and offset of an audio file
    # accurate to ms (0.001) 
    # if >=0.1, there's meaningful signal
    time = aud_sound_time(audfile,decimal=decimal,thr=thr)
    start,end = time[0],time[-1]
    return start, end

def aud_dur_ms(audfile,decimal=0.001,thr=0.1,unit='ms'):
    #this is to accurately calculate the duration/length of the sound signal in audio wav file 
    time = aud_sound_time(audfile,decimal=decimal,thr=thr)
    aud_dur = time[-1]-time[0]
    if unit=='ms':
        aud_dur = round(aud_dur*1000)
    return aud_dur

def aud_env(audfile,tr,zscore=False,outfile=False):
    #obtain the audio envelope, similar to resample(x,p,q) where p=1 in matlab
    import scipy.io.wavfile as wav
    import numpy as np
    import scipy.signal
    #read in .wav as frequency and signal (similar to audioread in matlab)
    fs,signal=wav.read(audfile)
    if np.ndim(signal)==2: #avg left/right chanel
        signal = np.mean(signal,1)
        print(signal)
    matlab_signal = signal/32767
    #transform signal
    stim_mod = abs(scipy.signal.hilbert(matlab_signal))
    #downsample data by frequency and tr
    q = int(fs*tr)
    audenv = scipy.signal.decimate(stim_mod, q, n=20*q, ftype='fir')
    if zscore==True:
        audenv = zscore(audenv)
    if outfile!=False:
        write_array(audenv,outfile)
    return audenv

def aud_rms(audpath,audname,unit='s',TR=1.5,plot=True):
    #this returns the rms of the audio file 
    #by the unit of seconds 's', or milliseconds 'ms': this is just to note
    #TR=1.5 for the interruption project 
    import os
    from pliers.extractors import RMSExtractor
    audio = os.path.join(audpath,audname)
    ext = RMSExtractor()
    df_aud = ext.transform(audio).to_df()
    df_aud['onset_tr'] = df_aud['onset']/TR
    if plot==True:
        df_aud.plot('onset', 'rms')
    return df_aud

def aud_onset_offset(audpath,audname,unit='s',thr=0.01):
    #this is to find the onset and offset of audio wav file based on rms>.01
    #unit = either 's' or 'ms'
    df_aud = aud_rms(audpath,audname,unit='s',plot=False)
    time = df_aud[df_aud['rms']>.01]['onset']
    onset = min(time)
    offset = max(time)
    if unit=='ms':
        onset = round(onset*1000)
        offset = round(offset*1000)
    return onset,offset

def downsample_rms(df_aud_rms,plot=True):
    #takes in the dataframe of audio rms info
    #outputs the downsampled vector of rms
    import matplotlib.pyplot as plt
    from scipy.interpolate import interp1d
    import numpy as np
    df = df_aud_rms
    # Linear interpolation
    linear_downsample = interp1d(df['onset_tr'][:],df['rms'][:],kind='slinear')#, kind='nlinear'
    # Specify the #of time points
    TR_time = np.arange(np.min(df['onset_tr'][:]), np.max(df['onset_tr'][:]))
    # downsample by #of time points
    downsampled_vector = linear_downsample(TR_time)
    #plot the original rms and the downsampled rms:
    if plot==True:
        plt.figure(figsize=(25,4))
        plt.axhline(0, c='k')
        plt.vlines(df['onset_tr'][:], [0], df['rms'][:], color='#4f5bd5', label='rms')
        plt.plot(TR_time, downsampled_vector, 'o-',color='#d62976', label='Downsampled signal')
        plt.legend(loc='upper left', fontsize=16)
        plt.xlabel('TR# (1.5)', fontsize=18)
        plt.ylabel('audio rms', fontsize=18)
        plt.grid()
    return downsampled_vector

def create_bp(aud_ls,onset_ls,offset_ls,order='default',outfile=False):
    #aim: the blueprint dataframe
        #aud_ls: list of audio names
        #onset_ls: list of onsets (start of crop) which are numbers (unit in ms-milliseconds, so if 2s, input 2000)
        #offset_ls: list of offsets (end of crop) which are numbers (unit in ms-milliseconds, so if 2s, input 2000)
        #default order to be 0,1,2... if not otherwise specified
    import numpy as np
    bp = init_dict(['name','start','end','length','order','old_name'])
    if len(aud_ls)==len(onset_ls)==len(offset_ls):
        for i in aud_ls:
            if i[-4:]!='.wav':
                i = i+'.wav'
            bp['name'].append(i)
        for i in onset_ls:
            bp['start'].append(i)
        for i in offset_ls:
            bp['end'].append(i)
        bp['length'] =  (np.array(bp['end'])-np.array(bp['start'])).tolist()
        bp['order'] = [i for i in range(len(aud_ls))]
        bp['old_name']=bp['name']
    else:
        print('The 3 lists inputted need to be of the same length!')
    bp_df = dict_to_df(bp,colname_by_key=True)
    bp_dict = bp
    if outfile!=False:
        import pandas as pd
        bp_df.to_csv('/'.join(outfile.split('/')[0:-1]) + '/bpOut_'+outfile.split('/')[-1]+'.csv')
        print('The created blueprint file saved.')
    return bp_df,bp_dict

def export_bp(bp_df,outfile='none',comb_gap='gap.wav',export=True):
    import numpy as np
    import pandas as pd
    if export==True:
        #get indices for gap and non-gap chunks
        inds = bp_df['name']!=comb_gap
        inds_gap = bp_df['name']==comb_gap
        new =bp_df[inds].reset_index(drop=True)
        #calc the onset, offset, length
        end = bp_df[inds]['end']
        dur = bp_df[inds_gap]['length']
        new['end']=np.array(end)+np.array(dur)
        new['length']= new['end']-new['start']
        new['onset']=calc_onset_from_dur(np.array(new['length']))[:-1]
        new['offset']= calc_onset_from_dur(np.array(new['length']))[1:]
        new['old_name'] = new['name']
        outname= outfile.split('/')[-1]
        new['name']=[outname]*len(new['old_name'])
        out= pd.DataFrame(new,columns=['order','name','old_name','onset','offset','length'])
        out['onset']=out['onset']/1000
        out['offset']=out['offset']/1000
        out['length']=out['length']/1000
        if outfile!='none':
            mkdir(outfile)
            bp_df.to_csv('/'.join(outfile.split('/')[0:-1]) + '/bpOut_'+outfile.split('/')[-1]+'.csv')
            out.to_csv('/'.join(outfile.split('/')[0:-1]) + '/export_'+outfile.split('/')[-1]+'.csv')
            print('The original and processed blueprint files saved.')
        return out
    else:
        if outfile!='none':
            mkdir(outfile)
            bp_df.to_csv('/'.join(outfile.split('/')[0:-1]) + '/bpOut_'+outfile.split('/')[-1]+'.csv')
            print('The original blueprint file saved.')
        return bp_df

def export_audio(bp_df,outfile,unit='ms',fmat='wav'):
    #this function outputs an audio at the specified blueprint file
        #bp_df is the dataframe for the blueprint
        #outfile needs to be outpath+filename (no filetype suffix), e.g. /Users/savefile
    from scrampy import splice 
    from pydub.audio_segment import AudioSegment
    aud_files = {}
    for i in bp_df['name']:
        aud_files[i] = AudioSegment.from_mp3(i)  
    audio=splice.aud_from_log(bp_df, **aud_files)
    audout = outfile+'.'+fmat
    audio.export(audout, format=fmat)
    print('Audio file saved.')

def scramble_audio(bp, aud_stories, output_name):
    """
    example use:
        ntf_bp = '/Users/xianl/Desktop/interleaved/audio/NotTheFall/current_try/try/blueprint_intact.csv'
        ntf = {'NotTheFall_sent_intact.wav':'NotTheFall_sent_intact.wav'}
        ntf_output="NotTheFall_sent_scram"
        xf.scramble_audio(ntf_bp,ntf,ntf_output)
    """
    from scrampy import splice
    import pandas as pd
    from pydub.audio_segment import AudioSegment
    import numpy as np
    print("----Loading Data----")
    df = pd.read_csv(bp, index_col=0)
    print("----Loading Audio----")
    aud_files = {}
    for key, fname in aud_stories.items():
        aud_files[key] = AudioSegment.from_mp3(fname)  
    print("----Scrambling Data & Audio----")
    df_nomusic=df.tail(len(df)-1)
    df_shuffled=df_nomusic.iloc[np.random.permutation(len(df_nomusic))]
    df_shuffled_reindexed=df.head(1).append(df_shuffled).reset_index(drop=True)
    audout=output_name + ".wav"
    audio=splice.aud_from_log(df_shuffled_reindexed, **aud_files)
    audio.export(audout, format=audout.split('.')[-1])
    print("----Exporting Blueprint----")
    # df_shuffled_reindexed['old_name'] = df_shuffled_reindexed['name']
    df_shuffled_reindexed['name']=audout
    print(df_shuffled_reindexed)
    bpout=output_name + ".csv"
    df_shuffled_reindexed.to_csv(bpout)

def calc_gap_length(aud, mindist=None, TRdur=1500):
    """Return number of ms padding before next clip such that..
    (1) next clip starts on multiple of 1.5s
    (2) clips are seperated by at least mindist sec
    """
    if type(aud) == int: t_now = aud
    else: t_now = len(aud)
    currentTR = int(t_now / TRdur)
    t_next = (currentTR + 1) * TRdur
    while t_next - t_now < mindist:
        #currentTR += 1
        #t_next = (currentTR + 1) * 1500
        t_next += TRdur
    out = t_next - t_now
    if (out/TRdur).is_integer() and mindist==0:
        out = 0
    return out

def gap_dict(row, mingap, TRdur):
    """Returns dictionary with fields for gap row. Assumes all rows will have
    gaps ending on a TR
    """
    length = calc_gap_length(row['length'], mingap, TRdur)
    return {'name': 'gap.wav',
            'length': length,
            'start': 0,
            'end': length,
            'order':-1}

def insert_gaps(bp, mingap=1000, TRdur=1500, outname=None):
    """Insert gap row after each entry.
    Parameters:
        bp: blueprint filename or dataframe
        mingap: minimum time (in ms) between segment ending and next TR
        TRdur:  length of each TR
    Note: 
        this procedure assumes that the audio begins on a TR, so that only the
        length of each segment may be used to calculate the gap length.
    """
    from pandas import DataFrame
    bp = pd.read_csv(bp) if type(bp) == str else bp
    new_rows = []
    for ii, row in bp.iterrows():
        new_rows.append(row.to_dict())
        new_rows.append(gap_dict(row, mingap, TRdur))
    new = fix_col_order(DataFrame(new_rows))
    if outname:
        new.to_csv(outname, index=None)
    return new

def fix_col_order(df):
    OUTCOLS = ['name', 'start', 'end', 'length', 'order']
    other = [col for col in df.columns if col not in OUTCOLS]
    return df[OUTCOLS + other]

def scram_list(inlist):
    # Make scramble_list_more_random available at the top level
    import random
    n = len(inlist)
    if n < 2:
        return inlist
    # Copy list to maintain original
    scrambled = inlist[:]
    for i in range(n-1, 0, -1):
        # Choose a random index from 0 to i-1 so that the item at i doesn't stay in the same position
        j = random.randrange(0, i)
        scrambled[i], scrambled[j] = scrambled[j], scrambled[i]
    # Check for the last swap to ensure no element remains in its original position
    if scrambled[0] == inlist[0]:
        # Swap the first element with any other element chosen randomly except the second element
        j = random.randrange(2, n)  # Choosing from the rest of the list excluding the first two elements
        scrambled[0], scrambled[j] = scrambled[j], scrambled[0]
    return scrambled

def scram_worker(item):
    # Define the worker function at the top level
    original_list, _ = item
    return scram_list(original_list)

def scramble_lists_in_parallel(inlist, n_permutations):
    # Example usage
        #n_items = 17
        #n_permutations = 21
        #scrambled_lists = scramble_lists_in_parallel(n_items, n_permutations)
    from concurrent.futures import ProcessPoolExecutor
    n_items=inlist
    if type(n_items)==int:
        original_list = list(range(n_items))
    else:
         original_list=inlist
    original_list = list(range(n_items))
    # Prepare the data to be processed in parallel, including a dummy value to make the input iterable
    tasks = [(original_list, i) for i in range(n_permutations)]
    with ProcessPoolExecutor() as executor:
        scrambled_lists = list(executor.map(scram_worker, tasks))
    return scrambled_lists

def count_combo(unique_items, repeated_items):
    #example use:
        # counts_dict, num_combinations = count_combinations_with_details(unique_items, repeated_items)
        # free-route has 18 different stories, yoked has 53 of these 18 stories, this is to select 18 from yoked that matches the stories in free
    # Create a dictionary to store the count of each unique item in the repeated items list
    counts_dict = {item: repeated_items.count(item) for item in unique_items}
    # Calculate the number of combinations
    num_combinations = 1
    for count in counts_dict.values():
        if count > 0:
            num_combinations *= count
    return counts_dict, num_combinations

def select_from_repeated_items(unique_items, repeated_items, repeated_couple):
    #example use:
        #selected_yoked_subid =select_from_repeated_items(stv, dat_y['yoked_to'],dat_y['ID'])
        # free-route has 18 different stories, yoked has 53 of these 18 stories, this is to select 18 from yoked that matches the stories in free
    # Create a dictionary to store indices of each unique item
    indices_dict = {item: [] for item in unique_items}
    # Populate the dictionary with indices from the repeated_items list
    for idx, item in enumerate(repeated_items):
        if item in indices_dict:
            indices_dict[item].append(idx)
    # Precompute the third list mapping
    third_list = repeated_couple
    def select_item(item, indices):
        return 'to' + item + '_sub' + third_list[random.choice(indices)]
    # Randomly select one index for each unique item and get the corresponding item from the third list
    selected_items = [select_item(item, indices) for item, indices in indices_dict.items() if indices]
    return selected_items

def compute_trial_select_from_repeated_items(args):
    #this is what we want to compute for a single trial in the next function
    unique_items, repeated_items, repeated_couple, data = args
    selected = select_from_repeated_items(unique_items, repeated_items, repeated_couple)
    isc = avg(isc_cent([data[i] for i in selected])[0], mute=True)
    return isc

def select_and_compute(unique_items, repeated_items, repeated_couple, data, n=1000, batch_size=100):
    from concurrent.futures import ProcessPoolExecutor
    #this is to do select_from_repeated_items() and compute_trial_() n times 
    args = (unique_items, repeated_items, repeated_couple, data)
    with ProcessPoolExecutor() as executor:
        futures = []
        for _ in range(n // batch_size):
            futures.extend(executor.submit(compute_trial_select_from_repeated_items, args) for _ in range(batch_size))
        isc_trials = [future.result() for future in futures]
    return isc_trials

def pval_from_dist(value, value_distribution):
    # Example usage
        # proportion = ecdf(isc_f_avg, isc_y) 
        # ecdf (empirical cumulative distribution function): gets the p-val from a distribution
    """
    Calculate the proportion of values in the distribution that are less than or equal to the given value.

    Parameters:
    - value (float): The value to compare against the distribution.
    - value_distribution (list of float): The list of values representing the distribution.

    Returns:
    - proportion (float): The proportion of values in the distribution that are less than or equal to the given value.
    """
    # Count how many values in the distribution are less than or equal to the given value
    count = sum(1 for x in value_distribution if x <= value)
    # Calculate the proportion
    p_val = count / len(value_distribution)
    return p_val



def quiry_yroute(yoked_to='all',yoked_sum='/Users/xianl/Desktop/CYOA/FixR_data/yoked/try/2_data/1_summary.xlsx'):
    #this is to check the story route from the yoked_to subject id
    check_yroute = {}
    dat,row = read_rows(read_table(yoked_sum),[7,9])
    for i in range(len(dat['yoked_to'])):
        sub = 'sub'+str(int(float(dat['yoked_to'][i])))
        yroute = clean_psiturk_route(dat['story_route'][i].split(','))
        if sub not in check_yroute:
            check_yroute[sub]=yroute
    if yoked_to =='all':
        return check_yroute
    else:
        return check_yroute[yoked_to]

def quiry_pid(debugid,pidpath='/Users/xianl/Desktop/CYOA/PassR_data/collect_data/data/first28/questiondata.xlsx'):
    #this is to get the psiturk id (in the questiondata.xlsx) from the debug id (in trialdata)
    dat,row = read_rows(read_table(pidpath),[1,2,3],skip_row1=False,key=['debugid','field_of_quiry','answer'])
    ans_list=[]
    for r in row:
        readid = row[r]['debugid'][0]
        fdofqy = row[r]['field_of_quiry'][0]
        answer = row[r]['answer'][0]
        if type(debugid)==str:
            if readid ==debugid and fdofqy=='prolific_ID1':
                return answer
        elif type(debugid)==list:
            for i in debugid:
                if readid ==debugid and fdofqy=='prolific_ID1':
                    ans_list.append(answer)
            return ans_list

def cyoa_get_eventinfo(mapfile,eventinfo,cols=[1,4,7],mapf_sheet=1):
    #free route and yoked route: get scene_lab corresponded event_lab and text 
    # cols are the column in the mapfile that corresponds to scene_lab, event_lab, and scene_text
    maptable = read_table(mapfile,sheet=mapf_sheet)
    nrows = maptable.nrows  
    counter = 0
    for row in range(nrows):
        counter+=1
        v = maptable.row_values(row)    
        col1,col2,col3= v[cols[0]-1],v[cols[1]-1],v[cols[2]-1]
        #skip the title row 
        if counter==1: 
            continue
        for key in eventinfo:  
            if str(col1).find('_') == -1:                        
                if str(int(col1))== key: #scenelabel
                    eventinfo[key].append(str(int(col2))) #eventnum
                    eventinfo[key].append(col3) #scenetxt     
            else:
                if str(col1)== key:
                    eventinfo[key].append(str(int(col2)))
                    eventinfo[key].append(col3)  
    return eventinfo

def cyoa_get_merge(mapfile,quiry_events,col_num=[1,2,10],merge='Y',pop_choice=False,choice_marker='_'):
    #mapfile: this is the story map in excel file that has the info of event_lab, scene_lab, story_text, converge or not, etc.
    #col_num: 'event_lab','scene_lab' and 'Converge' column (e.g. monthy_map)
    #quiry_events: the list of event_lab (string!) that you want to know the converge (Y/N) info for.
    #merge: 'Y'/'N' to indicate whether an event is in a merged/converged section or not
    #pop_choice: whether to pop out (exclude) the choice option events
    #choice_marker: choice option is currently detected as having '_' in the scene label
    #returns the indices of 'Y' for converge section events.
    if type(quiry_events[0])!=str:
        quiry_events = [str(i) for i in quiry_events]
    data = read_data(mapfile,col_num)
    evnt_lab,scne_lab,merge_not = [i for i in data.keys()]
    data[evnt_lab]=[str(int(float(i))) for i in data[evnt_lab]]
    dat_targ = dict_select_inds(data,evnt_lab,quiry_events)
    dat_new = init_dict([evnt_lab,scne_lab,merge_not])
    for i in range(len(dat_targ[evnt_lab])):
        event_lab = dat_targ[evnt_lab][i]
        scene_lab = dat_targ[scne_lab][i]
        if event_lab not in dat_new[evnt_lab]:
            dat_new[evnt_lab].append(event_lab)
            dat_new[scne_lab].append(scene_lab)
            dat_new[merge_not].append(dat_targ[merge_not][i])
    merge_inds=get_indices(dat_new[merge_not],merge)     
    if pop_choice==True:
        choice_inds = get_indices(dat_new[scne_lab],choice_marker,find=True)
        merge_inds = [i for i in merge_inds if i not in choice_inds]
    return merge_inds

def cyoa_get_choice_vec(story_route,to_int=True):
    #takes in story_route as a ','concatenated string 
    #outputs scene label, choice label, and choice vector, separately
    if type(story_route)==list:
        route = [i.split(',') for i in story_route]
        scenes,cholabs,chovecs = [],[],[]
        for route_sub in route:
            scene_lab = [i for i in route_sub if i.find('_')==-1]
            choice_lab = [i for i in route_sub if i.find('_')!=-1]
            if to_int!=True:
                choice_vec = [i.split('_')[-1] for i in choice_lab]
            else:
                choice_vec = [int(i.split('_')[-1]) for i in choice_lab]
            scenes.append(scene_lab)
            cholabs.append(choice_lab)
            chovecs.append(choice_vec)
    elif type(story_route)==str:
        route_sub = story_route.split(',')
        scenes = [i for i in route_sub if i.find('_')==-1]
        cholabs = [i for i in route_sub if i.find('_')!=-1]
        if to_int!=True:
            chovecs = [i.split('_')[-1] for i in choice_lab]
        else:
            chovecs = [int(i.split('_')[-1]) for i in choice_lab]
    return scenes,cholabs,chovecs

def cyoa_get_choice(story_route,choice_marker='_'):
    #this is to get the inds of the choice events and the scene label of the choices
    choice_inds = get_indices(story_route,choice_marker,find=True)
    scene_lab = apply_inds(choice_inds,story_route)
    return choice_inds,scene_lab

def cyoa_fake_choice(mapfile,col_num=1):
    #this is to detect which choice points had only one option (BeforeAlice story)
    #col_num is the column# of the scene_lab (to determine whether the current scene is a choice point)
    data,rows = read_rows(read_table(mapfile),[col_num])
    allchoices = {}
    for scene in data[[key for key in data.keys()][0]]:
        if '_' in scene:
            front = scene.split('_')[0] #scene_label = choice-point
            option= scene.split('_')[1] #choice_label = options
            if front not in allchoices:
                allchoices[front]=[]
            allchoices[front].append(option)
    fake_choice = []
    for key in allchoices:
        if len(allchoices[key])==1:
            fake_choice.append(key)
    return fake_choice

def cyoa_get_choiceinfo(mapfile,eventinfo,choices,col_num=1):
    #yoked route: get choice info (actual_choice, want/not) on top of eventinfo of scene/event/text
    #eventinfo is the scene/event/text got from cyoa_get_eventinfo(mapfile,eventinfo)
    #choices is the vector of options that the subject had selected along the way
    #col_num is the column# of the scene_lab (to determine whether the current scene is a choice point)
    fake = cyoa_fake_choice(mapfile,col_num)
    maptable = read_table(mapfile)    
    nrows = maptable.nrows  
    choice_count = 0
    for key in eventinfo:
        wanted = ' '
        counter = 0
        for row in range(nrows):
            counter+=1
            v = maptable.row_values(row)   
            scene_lab = v[col_num-1]
            #skip the storymap title row 
            if counter==1: 
                continue
            if str(scene_lab).find('_') == -1:  #not a choice point                      
                if str(int(scene_lab))== key:
                    choice = ' ' 
                    eventinfo[key].append(choice) #actual choice
                    eventinfo[key].append(wanted) #wanted or not
            elif str(scene_lab).split('_')[0] not in fake: #a real choice point        
                if str(scene_lab)== key:  
                    ind = scene_lab.find('_')
                    yoked_choice = scene_lab[ind+1:] #yoked to choice
                    choice = choices[choice_count] #actual choice
                    eventinfo[key].append(choice)
                    #wanted or not
                    wanted = '0'
                    if int(float(yoked_choice))==int(float(choice)):
                        wanted='1'
                    eventinfo[key].append(wanted)
                    choice_count+=1
            else: #a fake choice point  
                 if str(scene_lab)== key:  
                    choice = choices[choice_count] #actual choice
                    eventinfo[key].append(choice)
                    eventinfo[key].append('na')
                    choice_count+=1
    return eventinfo


def cyoa_subj_paths(mapfile,sumfile,subj_ind,map_cols=[1,4,7],sum_cols=[1,2,6],keynames=['scene_label','event_label','scene_text'],mapf_sheet=1,sumf_sheet=1):
    #free route: get map and subject data 
    alldat,allrow = read_rows(read_table(sumfile,sheet=sumf_sheet),sum_cols)  
    rowlist = [row for row in allrow]    
    subjrow = allrow[rowlist[subj_ind]]    
    #get the story_routes
    subjnum = int(float(subjrow['sub'][0]))
    subjid = int(float(subjrow['ID'][0]))
    scene_list =subjrow['story_route'][0].split(',')
    scenes = clean_psiturk_route(clean_list(scene_list))
    #get the eventinfo dict
    eventinfo = {}
    for scene_lab in scenes:
        eventinfo[str(scene_lab)]=[]
    eventinfo = cyoa_get_eventinfo(mapfile,eventinfo,cols=map_cols,mapf_sheet=mapf_sheet)
    subjinfo = dict_by_key(eventinfo,keynames)
    subj = 'sub'+str(subjnum)+'_'+str(subjid)
    return subj,subjinfo

def cyoa_subj_pass(mapfile,sumfile,subj_ind,map_cols=[1,4,7],sum_cols=[2,5,6],keynames=['scene_label','event_label','scene_text']):
    #free route: get map and subject data 
    alldat,allrow = read_rows(read_table(sumfile),sum_cols)  
    rowlist = [row for row in allrow]    
    subjrow = allrow[rowlist[subj_ind]]    
    #get the story_routes
    sub_id = int(float(subjrow['ID'][0]))
    yok_id = subjrow['yoked_to'][0].split('sub')[-1]
    scene_list =subjrow['story_route'][0].split(',')
    scenes = clean_psiturk_route(clean_list(scene_list))
    #get the eventinfo dict
    eventinfo = {}
    for scene_lab in scenes:
        eventinfo[str(scene_lab)]=[]
    eventinfo = cyoa_get_eventinfo(mapfile,eventinfo,map_cols)
    subjinfo = dict_by_key(eventinfo,keynames)
    subj = 'to'+str(yok_id)+'_'+'sub'+str(sub_id)
    return subj,subjinfo

def cyoa_subj_yoked(mapfile,sumfile,subj_ind,map_cols=[1,4,7],sum_cols=[2,7,9,10],keynames=['scene_label','event_label','scene_text','actual_choice','want_not']):
    #map_cols: 1-scene_lab, 4-event_lab, 7-scene_text
    #sum_cols: 2-ID, 7-yoked-to, 9-story_route, 10-actual_choice

    #yoked route: get map and subject data 
    alldat,allrow = read_rows(read_table(sumfile),sum_cols)  
    rowlist = [row for row in allrow]    
    subjrow = allrow[rowlist[subj_ind]] #'ID','yoked_to','story_route','actual_choices'
    #get the story_routes and choices
    sub_id = str(int(float(subjrow['ID'][0]))) 
    yok_id = str(int(float(subjrow['yoked_to'][0])))
    scenes = clean_psiturk_route(clean_list(subjrow['story_route'][0].split(',')))
    choices = clean_list(subjrow['actual_choices'][0].split(','))
    #get the eventinfo dict
    eventinfo = {}
    for scene_lab in scenes:
        eventinfo[str(scene_lab)]=[]
    eventinfo = cyoa_get_eventinfo(mapfile,eventinfo,map_cols)
    #get choice info
    eventinfo = cyoa_get_choiceinfo(mapfile,eventinfo,choices,col_num=map_cols[0])
    subjinfo = dict_by_key(eventinfo,keynames)
    subj = 'to'+str(yok_id)+'_'+'sub'+str(sub_id)
    return subj,subjinfo

def cyoa_get_survey_heroku_qdata(ques,resp,search_q='uncertainty_'):
    ius_inds = get_indices(ques,search_q,find=True)
    ius_q = apply_inds(ius_inds,ques)
    ius_a = [int(float(i)) for i in apply_inds(ius_inds,resp)]
    qlab = [int(i.split('_Q')[-1]) for i in ius_q]
    ius = sort_nlists([qlab,ius_q,ius_a])
    return ius

def cyoa_cut_events(data):
    #this is to segment the subj story info ('scene_lab','event_lab','scene_text') by events
        #e.g. subj,data = xf.cyoa_subj_paths(mapfile,sumfile,subj_ind=ind,map_cols=[2,1,3],sum_cols=[1,2,7],sumf_sheet=2)
    cutted,cut_by = group_cut(data['scene_text'],data['scene_label'],'_')
    group_by_new = group_by_both_list(data['event_label'],cut_by)
    data['event_num'] = group_by_new
    grouped_data,n_groups = group_dict(data,'event_num')
    simp,full = merge_ioflist_in_data(grouped_data)
    return simp

##############################################################################################################################################################
############################################################## Semantic and Causal Centrality:################################################################
##############################################################################################################################################################


def get_embeds(all_lines,title=False,outpath=False):
    # This is to get USE vectors (i.e. sentence embeddings)
    # all_lines is a list of strings; each element will be regarded as a row/unit
    # title: this is a string for the current inputed texts (i.e. should indicate which set of texts, e.g. sub2)
    # outpath: if wants to save the use vector per event, specify output path
    # example use:
        # get_embeds(['Alice went after the rabbit','Oh, dear!',...])
        # get_embeds(['Alice went after the rabbit','Oh, dear!',...],'sub12','./data/')
    import tensorflow.compat.v1 as tf
    import tensorflow_hub as hub
    import numpy as np
    tf.compat.v1.disable_eager_execution()
    module_url = "https://tfhub.dev/google/universal-sentence-encoder-large/4"
    embed = hub.load(module_url)
    #get embeddings
    with tf.Session() as session:
        session.run([tf.global_variables_initializer(), tf.tables_initializer()])
        embeddings = session.run(embed(all_lines))
    # save as csv
    if outpath!=False:
        mkdir(outpath)
        if title!=False:
            ftitle = title+'_SentenceEmbeddings_'+ str(len(all_lines))+'nodes.csv'
        else:
            ftitle = 'SentenceEmbeddings_'+ str(len(all_lines))+'nodes.csv'
        np.savetxt(outpath+ftitle,embeddings['outputs'],delimiter=",")
    return embeddings['outputs']

def get_embeds2(all_lines,title=False,outpath=False):
    import tensorflow as tf
    import tensorflow_hub as hub
    # Enable eager execution
    tf.compat.v1.enable_eager_execution()
    # Sentences to encode
    sentences = all_lines
    # Load the model from TensorFlow Hub
    embed = hub.load("https://tfhub.dev/google/universal-sentence-encoder/4")
    # Generate embeddings
    embeddings = embed(sentences)
    # Convert the embeddings to NumPy arrays
    embeddings_array = embeddings.numpy()
    return embeddings_array

def get_embeds_glove(glove_text="glove.6B.300d.txt"):
    #this is assuming the glove library text file is in the same path 
    import numpy as np
    embeddings_dict = {}
    with open(glove_text, 'r') as f:
        for line in f:
            values = line.split()
            word = values[0]
            vector = np.asarray(values[1:], "float32")
            embeddings_dict[word] = vector
    return embeddings_dict

def find_closest_embeddings(embed_dict,word,top=10):
    from scipy import spatial
    embedding = embed_dict[word]
    out = sorted(embeddings_dict.keys(), key=lambda word: spatial.distance.euclidean(embeddings_dict[word], embedding))
    if top!=False:
        out= out[:top]
    return out

def theme_sim(rate_word,theme_word,embed_dict,measure='max'):
    #rate_word=['word1','word2',...]
    #theme_word=['river','murder',...]
    #embed_dict = xf.get_embed_glove('glove.6B.300d.txt')
    rate_vec = []
    theme_vec = []
    for word in rate_word:
        rate_vec.append(embed_dict[word])
    for word in theme_word:
        theme_vec.append(embed_dict[word])
    sim_mat = cosine_sim(rate_vec,theme_vec)
    if measure=='max':
        out = maxpool(sim_mat)
    elif measure=='mean':
        out = avg(sim_mat,dim=1)
    else:
        out=sim_mat
    return out

def semantic_max_match(reference,test,title=False,outpath=False):
    #reference: an ndarray (e.g. story_ndarray) serving as the template to see how the test vector is maximally similar to it. 
    #test: an ndarray (e.g. recall_ndarray) that we are curious how maximally (semantically) similar it is to the template ndarray
    #outpath: if wants the reference x test similarity matrix to be saved, needs to specify the outpath to save to
    #title: if specified the outpath, can (optional) also specify the title of the matrix csv file to be saved
    #example use:
        #max_avg = semantic_max_match(story_ndarray,recall_ndarray,title,outpath='./autoScore/sim_mat/')
    import numpy as np
    import sklearn    
    from sklearn import metrics
    sem_mat = sklearn.metrics.pairwise.cosine_similarity(reference,test,dense_output=True)
    print(sem_mat.shape)
    # save matrix data for template vs. test vectors
    if outpath!=False:
        mkdir(outpath)
        if title!=False:
            csv_title = title
        else:
            csv_title = 'Reference_x_Test_matrix_'+ str(sem_mat.shape[0])+'_x_'+str(sem_mat.shape[1])+'nodes'
        np.savetxt(outpath+csv_title+'.csv', sem_mat, delimiter=",")
    #get the max match (greatest corr) across the test nodes per reference node 
    max_list = []
    for row in sem_mat:
        row_max = np.max(row)
        max_list.append(row_max)
    max_avg = np.mean(max_list)
    return max_avg,max_list


def semantic_cent(all_lines,title=False,outpath=False,labname=False,decimal=6):
    # This is to get semantic centrality per node 
    # all_lines: a list of strings/texts; each element will be regarded as a row/unit
    # title: this is for the current inputed texts (i.e. should indicate which set of texts, e.g. sub2)
    # outpath: if wants to save the matrix graph and data, specify output path
    import numpy as np
    import sklearn    
    from sklearn import metrics
    #get USE vectors (i.e. sentence embeddings)
    vectors = get_embeds(all_lines,title=title,outpath=outpath)
    # compute cosine sim mat for all events against each other
    sem_mat = sklearn.metrics.pairwise.cosine_similarity(vectors, Y=None, dense_output=True)
    print(sem_mat.shape)
    # save matrix graph and data for vector vs. vector
    if title!=False:
        graph_title = title+'_SemanticMat_'+ str(len(all_lines))+'nodes'
    else:
        graph_title = 'SemanticMat_'+ str(len(all_lines))+'nodes'
    plot_matrix(sem_mat,graph_title,outpath,labname=labname,diagmask=True)
    # calculate semantic centrality score for each event
    sem_cent=[]
    for row in sem_mat:
        sem_cent.append(round(np.mean(row),decimal))
    mesg = 'This is the semantic centrality scores:'+str(sem_cent)
    print(mesg)  
    return sem_cent, sem_mat


def modify_conn_weight(choice_events,caus,efct,rew_caus_efct,w1_w2):
    #this is to modify the connectivity strength of the identified causal pairs
        #it returns the modified connectivity strength (originally default as 1) w1 for the qualified, and w2 for the rest
    #rew_caus_efct=(1,'nan') means that only if caus event is in the choice will the conn be re-weighted to w1
    if caus*rew_caus_efct[0] in choice_events or efct*rew_caus_efct[1] in choice_events:
        conn = w1_w2[0]
    else:
        conn = w1_w2[1]
    return conn


def causal_cent(table,col1,col2,row1,pair,title=False,outpath=False,mute=True,re_weigh=False,rew_events=False,rew_caus_efct=(1,1),decimal=6):
    #table: the file that raters rated the causal relationship between events/nodes
    #col1: the first column should have the cause event
    #col2: the second column should have the effect event, 
    #row1: the row that the cause-effect number pairs starts (usually 2nd row)
    #pair: the dictionary that stores pairing relationship between the natural sequence of 1,2,3...n and the event/node labels; 
         # or an integer that specifies the total number of events (if these events follow natural sequence)
    #title: this is for the current inputed texts (i.e. should indicate which set of texts, e.g. sub2)
    #outpath: if wants to save the matrix graph and data, specify output path
    #re_weigh: re-weight the matrix to (w1,w2), with w1 being the weight assigned for the reweight_events, and w2 for the rest
    #rew_events: only needed if reweight is not False; needs to be a list of integers of the event number that's reweighted
    #rew_caus_efct: if (1,'nan') then only cause in the rew_events will be re-weighted; if (1,1), both cause and effect will be re-weighted
    import os
    import numpy as np
    #check if output path (if not False) folder exists:
    if outpath!=False and os.path.isdir(outpath) ==False:
        parentdir =  '/'+'/'.join(outpath.split('/')[1:-2])
        if os.path.isdir(parentdir) ==False:
            os.mkdir(parentdir)
        os.mkdir(outpath)
    #contruct matrix
    if type(pair)==dict:
        n_nodes = len(pair)
        caus_mat = np.zeros((n_nodes,n_nodes))
    elif type(pair)==int:
        n_nodes = pair
        caus_mat = np.zeros((n_nodes,n_nodes))
    else:
        print('ERROR: pair type needs to be either int (number of nodes) or dict (pairwise encoding with keys being the non-natural sequence node label)')
    caus_vec = []
    counter = 0  
    nrows = table.nrows      
    for row in range(nrows):
        counter+=1
        v = table.row_values(row)
        #start with row1 
        if counter<row1:
            continue   
        #only process the rows at the specified cols that are with numbers
        if type(v[col1-1])==float or type(v[col1-1])==int:
            #get cause and effect node numbers 
            caus = int(v[col1-1])
            efct = int(v[col2-1])
            #selectively enhance the connection strength of choice events
            if re_weigh !=False:
                conn = modify_conn_weight(rew_events,caus,efct,rew_caus_efct,re_weigh)
            else:
                conn = 1
            #fill in with 1 at specified places
            if type(pair)==dict:
                ind1 = pair[str(caus)]
                ind2 = pair[str(efct)]
            else:
                ind1 = caus
                ind2 = efct
            caus_mat[ind1-1,ind2-1] =conn
            caus_mat[ind2-1,ind1-1] =conn
    print(caus_mat.shape)
    #save figure for causality matrix
    if title!=False:
        graph_title = title+'_CausalityMat_'+ str(n_nodes)+'Events'
    else:
        graph_title = 'CausalityMat_'+ str(n_nodes)+'Events'
    plot_matrix(caus_mat,graph_title,outpath,labname='',diagmask=True)
    #calculate causality centrality score for each event
    caus_cent=[]
    for row in caus_mat:
        caus_cent.append(round(np.mean(row),decimal)) 
    if mute!=True:
        mesg = 'This is the causality centrality scores:'+str(caus_cent)
        print(mesg)
    return caus_cent,caus_mat


def contextual_cent(table,col1,col2,row1,pair,title=False,outpath=False,mute=True,re_weigh=False,rew_events=False,rew_curr_cntx=(1,1),decimal=6):
    #table: the file that raters rated the contextual relationship between events/nodes
    #col1: the first column should have the current event
    #col2: the second column should have the context event (the ones people are reminded of when reading current event to aid interpretation of the current event), 
    #row1: the row that the current-context number pairs starts (usually 2nd row)
    #pair: the dictionary that stores pairing relationship between the natural sequence of 1,2,3...n and the event/node labels; 
         # or an integer that specifies the total number of events (if these events follow natural sequence)
    #title: this is for the current inputed texts (i.e. should indicate which set of texts, e.g. sub2)
    #outpath: if wants to save the matrix graph and data, specify output path
    #re_weigh: re-weight the matrix to (w1,w2), with w1 being the weight assigned for the reweight_events, and w2 for the rest
    #rew_events: only needed if reweight is not False; needs to be a list of integers of the event number that's reweighted
    #rew_curr_cntx: if (1,'nan') then only current event in the rew_events will be re-weighted; if (1,1), both current and context event will be re-weighted
    import os
    import numpy as np
    #check if output path (if not False) folder exists:
    if outpath!=False and os.path.isdir(outpath) ==False:
        parentdir =  '/'+'/'.join(outpath.split('/')[1:-2])
        if os.path.isdir(parentdir) ==False:
            os.mkdir(parentdir)
        os.mkdir(outpath)
    #contruct matrix
    if type(pair)==dict:
        n_nodes = len(pair)
        cntx_mat = np.zeros((n_nodes,n_nodes))
    elif type(pair)==int:
        n_nodes = pair
        cntx_mat = np.zeros((n_nodes,n_nodes))
    else:
        print('ERROR: pair type needs to be either int (number of nodes) or dict (pairwise encoding with keys being the non-natural sequence node label)')
    counter = 0  
    nrows = table.nrows      
    for row in range(nrows):
        counter+=1
        v = table.row_values(row)
        #start with row1 
        if counter<row1:
            continue   
        #only process the rows at the specified cols that are with numbers
        if type(v[col1-1])==float or type(v[col1-1])==int:
            #get current and context node numbers 
            curr = int(v[col1-1])
            cntx = int(v[col2-1])
            #selectively enhance the connection strength of choice events
            if re_weigh !=False:
                conn = modify_conn_weight(rew_events,curr,cntx,rew_curr_cntx,re_weigh)
            else:
                conn = 1
            #fill in with 1 at specified places
            if type(pair)==dict:
                ind1 = pair[str(curr)]
                ind2 = pair[str(cntx)]
            else:
                ind1 = curr
                ind2 = cntx
            cntx_mat[ind1-1,ind2-1] =conn
            cntx_mat[ind2-1,ind1-1] =conn
    print(cntx_mat.shape)
    #save figure for contextuality matrix
    if title!=False:
        graph_title = title+'_ContextualityMat_'+ str(n_nodes)+'Events'
    else:
        graph_title = 'ContextualityMat_'+ str(n_nodes)+'Events'
    plot_matrix(cntx_mat,graph_title,outpath,labname='',diagmask=True)
    #calculate contextuality centrality score for each event
    curr_cent=[]
    for row in cntx_mat:
        curr_cent.append(round(np.mean(row),decimal)) 
    if mute!=True:
        mesg = 'This is the contextual centrality scores:'+str(curr_cent)
        print(mesg)
    return curr_cent,cntx_mat




def cent_bymat(mat,decimal=6):
    #this is to calculate the centrality as the average connection of a node to the rest of the nodes
    import numpy as np
    cent=[]
    for row in mat:
        cent.append(round(np.mean(row),decimal)) 
    return cent

def isc_cent(allsub_cents,method='pairwise',mute=True,skip_nan_min=30):
    #allsub_cents: a matrix of multiple subjects--each item is a list of centrality scores for one subject
    #method: 'pairwise' is each pair of subjects, 'average' is each subject vs. avg others
    rval,pval = [],[]
    if method=='pairwise':
        rmat,pmat=get_paircorr(allsub_cents,allsub_cents,thresh_n=skip_nan_min)
        rval = get_tril_val(rmat)
        pval = get_tril_val(pmat)
    elif method=='average':    
        for i in range(len(allsub_cents)):
            curr_cent = allsub_cents[i]
            others = avg([allsub_cents[j] for j in range(len(allsub_cents)) if j!=i],mute=mute)
            r,p= get_corr(curr_cent,others,skip_nan_min=skip_nan_min)
            rval.append(r)
            pval.append(p)
    return rval,pval

def isc_cent_cross_group(group1_cents,group2_cents,methods='all',legend=['g1','g2']):
    #methods: 'pairwise', 'average' or 'all'
    data = {}
    keys_all = [legend[0]+'-'+legend[1]+'_isc-pw',legend[0]+'avg-'+legend[1]+'_isc-avg',legend[0]+'-'+legend[1]+'avg_isc-avg']
    #add keys
    if methods =='all':
        data = dict_add_keys(data,keys_all)[0]
    elif methods=='pairwise':
        data = dict_add_keys(data,keys_all[0:1])[0]
    elif methods=='average':
        data = dict_add_keys(data,keys_all[1:])[0]
    else:
        print('Error: methods can only take on pairwise, average, or all.')
    #add vals
    if methods=='pairwise' or methods=='all':
        sim_mat= get_pairwise_pearsonr(group1_cents,group2_cents)
        allr = flatten(sim_mat)[0]
        data[keys_all[0]]=allr 
    if methods=='average' or methods=='all':
        sim_mat= get_pairwise_pearsonr([avg(group1_cents)],group2_cents)
        allr = flatten(sim_mat)[0]
        data[keys_all[1]]=allr 
        sim_mat= get_pairwise_pearsonr(group2_cents,[avg(group1_cents)])
        allr = flatten(sim_mat)[0]
        data[keys_all[2]]=allr 
    return data

def align_rcl_to_ev(rcl_ev,rcl_embed,n_ev,mute=True):
    rcl_embed_prep = init_mat((n_ev,512),fill='nan')
    for ev in range(1,n_ev+1):
        inds = get_inds_lsofls(rcl_ev,ev)
        if inds!=[]:
            temp = []
            for i in inds:
                recall = rcl_embed[i]
                temp.append(recall)
            temp = avg(temp,mute=mute)
            rcl_embed_prep[ev-1] = temp
    return rcl_embed_prep

def cosine_similarity(v1, v2):
    """Calculate cosine similarity between two vectors."""
    dot_product = np.dot(v1, v2)
    magnitude = np.linalg.norm(v1) * np.linalg.norm(v2)
    if magnitude == 0:
        return 0
    return dot_product / magnitude

def rcl_vs_ev_sim(ev_embed,rcl_embed):
    array1,array2 = ev_embed,rcl_embed
    # Ensure both arrays have the same shape
    assert array1.shape == array2.shape
    # Compute cosine similarity for each pair of rows
    cosine_similarities = np.array([
        np.nan if np.isnan(row1).any() or np.isnan(row2).any() else cosine_similarity(row1, row2)
        for row1, row2 in zip(array1, array2) ])
    return cosine_similarities


def sample_one_val_per_key(data_dict,nsample=1):
    #this is to select a random val for each of the dict[key]
    import random    
    out_dict = init_dict(['s'+str(i+1) for i in range(nsample)])
    for sample in out_dict:
        curr_sample={}
        for key in data_dict:
            curr_sample[key] = random.choice(data_dict[key])        
        out_dict[sample] = curr_sample
    return out_dict

def compute_corr_per_sample(sample_dict,cols_to_corr=[1,2]):
    #this is to compute the correlation of the two columns (val1 and val2) in each sample_dict
    allr,allp=[],[]
    for s in sample_dict:
        a = convert_list_array(dict_to_lists(sample_dict[s])[0])#get the two lists of values in the current sample
        vec1,vec2 = a[:,cols_to_corr[0]-1], a[:,cols_to_corr[1]-1]#col1 and col2
        r,p = get_corr(vec1,vec2, skip_nan_min=10)
        allr.append(r)
        allp.append(p)
    return allr,allp
        


def recall_temp_mat(rcl,tot):
    #rcl: this is subjects' recalled event number in temporal sequence, e.g. rcl = ['1.0', '2.0', '1,2,4', '18,19,21,22', '10,14', '14.0', '18,19,20,21,22,30,31', '38,39,40', '41.0', '45,46,47,48', '49.0', '52.0', '57,58', '54,55,56,57', '66.0', '66.0', '72.0', '85,86', '82.0', '59,60', '61,64', '61,64', '87,88', '96.0', '97,98,99', '100,101', '98,99,100', '109.0', '119,120,121', '129,130']
    #tot: this is the number of events in the story the subject had read, e.g. tot= 131 
    mat = init_mat((len(rcl),tot))
    for i in range(len(rcl)):
        item = rcl[i]
        if item.find(',')!=-1:
            y_ls = [int(i) for i in item.split(',')]
            for j in y_ls:
                mat[i][j-1]=1
        else:
            mat[i][int(float(item))-1] = 1
    return mat

def process_recall_seg(rcl_seg):
    # for a segment of recall being either a string of float '2.0' or a ','joined string '11,12', return the processed recall in a list 
    item = rcl_seg
    if item.find(',')!=-1:
        out = [int(i) for i in item.split(',')]
    else:
        out = [int(float(item))]
    return out

def recall_seg_split_num(rcl,thres=6):
    #rcl: this is subjects' recalled event number in temporal sequence, e.g. rcl = ['1.0', '2.0', '1,2,4', '18,19,21,22', '10,14', '14.0', '18,19,20,21,22,30,31', '38,39,40', '41.0', '45,46,47,48', '49.0', '52.0', '57,58', '54,55,56,57', '66.0', '66.0', '72.0', '85,86', '82.0', '59,60', '61,64', '61,64', '87,88', '96.0', '97,98,99', '100,101', '98,99,100', '109.0', '119,120,121', '129,130']
    #this function detects potential sticky wrong number e.g. '116117.0' in the list, and cut it into '116,117'
    ind = 0
    rcl_out = rcl.copy()
    for i in rcl_out:
        if i.find(',')==-1 and len(i)>=thres:
            i2 = i[:-2] #drop the .0 at the end of the number
            cutted = split_num(i2,thres=thres-2)
            cutted_store = []
            while len(cutted[-1])>=thres-2:
                cutted_store.append(cutted[:-1])
                cutted = split_num(cutted[-1],thres=thres-2)
            cutted_store.append(cutted)
            new = ','.join(merge_lists_in_list(cutted_store))
            rcl_out[ind]=new
        ind+=1
    return rcl_out

def recall_temp_vio(rcl):
    #rcl: this is subjects' recalled event number in temporal sequence, e.g. rcl = ['1.0', '2.0', '1,2,4', '18,19,21,22', '10,14', '14.0', '18,19,20,21,22,30,31', '38,39,40', '41.0', '45,46,47,48', '49.0', '52.0', '57,58', '54,55,56,57', '66.0', '66.0', '72.0', '85,86', '82.0', '59,60', '61,64', '61,64', '87,88', '96.0', '97,98,99', '100,101', '98,99,100', '109.0', '119,120,121', '129,130']
    #returns tv (the number of recall segments that violated the temporal sequence of event occurance), and tv_rate (the number of tv over the number of recall segments in total)
    #this is to calculate based on the subjects' segmented recall: how many of these segments violated the temporal sequence (i.e. min[current event num]<max[previous event num])
    tv=0
    for i in range(len(rcl)):
        if i==0:
            prev = process_recall_seg(rcl[i])
        else:
            prev = process_recall_seg(rcl[i-1])
            curr = process_recall_seg(rcl[i])

            if min(curr)<max(prev):
                tv+=1
    tv_rate = tv/len(rcl)
    return tv,tv_rate

def recall_temp_vio2(rcl):
    #rcl: this is subjects' recalled event number in temporal sequence, e.g. rcl = ['1.0', '2.0', '1,2,4', '18,19,21,22', '10,14', '14.0', '18,19,20,21,22,30,31', '38,39,40', '41.0', '45,46,47,48', '49.0', '52.0', '57,58', '54,55,56,57', '66.0', '66.0', '72.0', '85,86', '82.0', '59,60', '61,64', '61,64', '87,88', '96.0', '97,98,99', '100,101', '98,99,100', '109.0', '119,120,121', '129,130']
    #returns tv (the number of recall segments that violated the temporal sequence of event occurance), and tv_rate (the number of tv over the number of recall segments in total)
    #this is to calculate based on the subjects' segmented recall: how many of these segments violated the temporal sequence (i.e. min[current event num]<min[previous event num])
    tv=0
    for i in range(len(rcl)):
        if i==0:
            prev = process_recall_seg(rcl[i])
        else:
            prev = process_recall_seg(rcl[i-1])
            curr = process_recall_seg(rcl[i])

            if min(curr)<min(prev):
                tv+=1
    tv_rate = tv/len(rcl)
    return tv,tv_rate



def mesg(query):
    #given the question string, put it into the format of Q&A
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": query},
    ]
    return messages


def ask_model(query,model='gpt-4'):
    #create an API call
    if type(query)==str:
        conv = mesg(query)
    else:
        conv = query
    #ask for model response
    response = openai.ChatCompletion.create(
        model=model,
        messages=conv
    )
    response_message = response["choices"][0]["message"]['content']
    return response_message


def correct_quote_in_key(inlist,key_ls=['event_A_number','event_B_number','rating','reasoning']):
    for key in key_ls:
        inlist = replace_str(inlist,'’'+key+'’','"'+key+'"')
    return inlist


def correct_quote_in_keydata(inlist):
    import ast
    data_strings=inlist
    parsed_dicts = []
    for data_string in data_strings:
        sanitized_string = data_string.replace('’', "'").replace('‘', "'")
        try:
            item = ast.literal_eval(sanitized_string)
            parsed_dicts.append(item)
        except ValueError as e:
            print(f"Failed to parse: {data_string}\nError: {e}\n")
    return parsed_dicts
    

# def process_resp(resp,outpath_f=False):
#     #assumes the format of the resp being a string of multiple dictionaries separate by \n
#     keyls = ['event_A','event_B','rate','reason']
#     data = xf.init_dict(keyls)
#     a = resp.split('\n')
#     b = clean_quotes(a)
#     c = xf.clean_textls(b,clean_ls=['',' ','"','“','”',']','}'])
#     d = correct_quote_in_key(c)
#     if d[-1][-1]=='}':
#         e = correct_quote_in_keydata(d)
#         for i in range(len(e)):  
#             if type(e[i])==tuple:
#                 item = e[i][0]
#             else:
#                 item = e[i]
#             ev_A_ = item['event_A_number']
#             ev_B_ = item['event_B_number']
#             if str(ev_A_).find('event') !=-1:
#                 ev_A,ev_B = int(ev_A_.split('event')[-1].split('’')[0]),int(ev_B_.split('event')[-1].split('’')[0])
#             else:
#                 ev_A,ev_B = int(ev_A_),int(ev_B_)
#             rate = int(item['rating'])
#             reason = xf.clean_text(item['reasoning'])
#             xf.dict_add_vals(data,keys=keyls,vals=[ev_A,ev_B,rate,reason])
#         if outpath_f!=False:
#             xf.write_dict_anylen(data,outpath_f)
#         return data
#     else:
#         print('Error: returned response is incomplete, ending abruptly.')


def process_resp(resp,outpath_f=False):
    #assumes the format of the resp being a string of multiple dictionaries separate by \n
    keyls = ['event_A','event_B','rate','reason']
    data = init_dict(keyls)
    a = resp.split('\n')
    b = clean_quotes(a)
    c = clean_textls(b,clean_ls=['',' ','"','“','”',']','}'])
    d = correct_quote_in_key(c)
    if d[-1][-1]=='}':
        for i in range(len(d)):  
            item = d[i]
            ev_A_ = item.split('":')[1].split(',')[0]
            ev_B_ = item.split('":')[2].split(',')[0]
            if 'event' in ev_A_ and '’' in ev_A_:
                ev_A,ev_B = int(ev_A_.split('event')[-1].split('’')[0]),int(ev_B_.split('event')[-1].split('’')[0])
            elif 'event' in ev_A_:
                ev_A,ev_B = int(ev_A_.split('event')[-1].split('"')[0]),int(ev_B_.split('event')[-1].split('"')[0])
            else:
                ev_A,ev_B = int(ev_A_),int(ev_B_)
            rate = int(item.split('":')[3].split(',')[0])
            reason = clean_text(item.split('":')[4],['},','}'])
            dict_add_vals(data,keys=keyls,vals=[ev_A,ev_B,rate,reason])
        if outpath_f!=False:
            write_dict_anylen(data,outpath_f)
        return data
    else:
        print('Error: returned response is incomplete, ending abruptly.')


def clean_dot(ls):
    #this will drop the .0 for the integer strings in the list
    #example: ls = ['1.0','2.0','1,2,5','12,13','33.0'], output= ['1','2','1,2,5','12,13','33']
    outls = []
    for i in ls:
        if i.find('.')!=-1:
            j = i.split('.')[0]
        else:
            j = i
        outls.append(j)
    return outls

def pred_acc(overlap):
    if overlap==[]:
        return 0
    else:
        return 1


def neighbor_pred(rcl,ignore_nan=True,thres=0):
    #given a list of recall score, for each item, predict its recall based on previous and next item's recall (average)
    #predicted recall score for an item need to be >= thres to be kept as is, otherwise will be noted as 0
    rcl_pred = []
    for i in range(len(rcl)):
        if i>0 and i< len(rcl)-1:
            rcl_pre = rcl[i-1]
            rcl_pos = rcl[i+1]
        elif i == 0:
            rcl_pre = np.nan
            rcl_pos = rcl[i+1]
        elif i ==len(rcl)-1:
            rcl_pre = rcl[i-1]
            rcl_pos = np.nan
        rcl_cur_pred = avg([rcl_pre,rcl_pos],mute=True,ignore_nan=True)
        rcl_pred.append(rcl_cur_pred)
    rcl_pred = [i if i >= thres else 0 for i in rcl_pred]
    r,p = get_corr(rcl_pred,rcl)
    return r,p,rcl_pred

def neighbor_forward_pred(rcl,ignore_nan=True,thres=0):
    #given a list of recall score, for each item, predict its recall based on next item's recall; drop the last item (cannot be predicted)
    rcl_fwd = rcl[1:]
    r,p = get_corr(rcl_fwd,rcl[:-1])
    return r,p

def neighbor_clusters(rcl):
    binary_list=rcl
    count = 0
    in_cluster = False
    for number in binary_list:
        if number == 1:
            if not in_cluster:
                count += 1
                in_cluster = True
        else:
            in_cluster = False
    return count

def neighbor_cluster_length(rcl):
    #returns the avg, std, and the vector of length for a clump of recalled events 
    binary_vector = rcl
    lengths = []  # List to store lengths of consecutive 1s
    current_length = 0  # Current length of consecutive 1s

    for bit in binary_vector:
        if bit == 1:
            current_length += 1  # Increment length for consecutive 1s
        else:
            if current_length > 0:
                lengths.append(current_length)  # Add length to list if it's more than 0
                current_length = 0  # Reset current length
    # Check if the last element in vector was 1
    if current_length > 0:
        lengths.append(current_length)
    return avg(lengths),std(lengths),lengths
    

##############################################################################################################################################################
############################################################## Models of Mind and Brain:######################################################################
##############################################################################################################################################################

#example use:
    # find point location from distance matrix
        # orig_points = xf.rand_points(6,2)
        # targ_distance = xf.get_distance(orig_points,plot=True)
        # rand_points = xf.rand_points(6,2,plot=True)
        # xf.shepard_stress(orig_points,targ_distance,6,2)
        # xf.shepard_stress(rand_points,targ_distance,6,2)
        # loc = xf.find_locations(targ_distance,6,2)
    # matrix transform:
        # orig_points = xf.points_circle(npoints=38)
        # trans_points = xf.mat_transform(orig_points,scale=2,rot=0.2)
        # xf.plot_transform(orig_points,trans_points,38)
        # xf.nlinear_trans_demo(metric='relu')
    #hopfield net:
        # allArray = xf.get_img_data('letters_small/A.png',plot=True)  #get memory (store the states)
        # for myletter in ['B', 'C', 'D', 'E']:
        #     imageArray = xf.get_img_data(str('letters_small/'+ myletter + '.png'))
        #     allArray = np.concatenate((allArray, imageArray), axis=1)  # concatenate by column (axis = 1)
        # (numNeurons, numStates) = allArray.shape
        # squareSize = int(numNeurons**0.5) 
        # xf.plot_states(allArray,numStates,numNeurons);
        # rand = sci.random.randint(0,2, (numNeurons, 1)) #generate input data: random integer of 0/1 (in range(0,2)) in the shape of (25,1) 
        # dynamics, weight_mat = xf.hopfield_net(rand,allArray,num_neurons=numNeurons,num_states=numStates,timesteps=150,state_labs=['A','B', 'C', 'D', 'E'])




def rand_vec(vec_size,rng=[0,1],integer=True):
    #vec_size can take on any dim, e.g. (10,2) will return a 10rows x 2col mat
    #rng specifies the range of the number randomly generated for the mat
    #if integer==False, will generate float
    import numpy as np        
    if integer==True:
        random_vec = np.random.randint(low=rng[0], high=rng[1]+1, size=vec_size)
    else:
        if len(vec_size)==1:
            random_vec = np.random.rand(vec_size)
        else:
            random_vec = np.random.rand(*vec_size)
    return random_vec



def plot_points(points,labels='dot'):
    import numpy as np
    from matplotlib import pyplot as plt
    locations = points
    npoints = np.shape(locations)[0]
    plt.plot(locations[:,0], locations[:,1], '.')  
    xoffset, yoffset = +0.0, -0.08 #offset for the label (to be next to the point)
    for ipoint in range(npoints):
        current_label = labels+str(ipoint+1)
        plt.text(locations[ipoint,0]+xoffset, locations[ipoint,1]+yoffset, current_label, {'fontsize':14, 'color':'r'})
    plt.show()


def plot_distance_mat(points,labels='dot'):
    import numpy as np
    from matplotlib import pyplot as plt
    import sklearn.metrics.pairwise as metrics
    locations = points
    distance = metrics.pairwise_distances(locations)
    npoints = np.shape(points)[0]
    mylabels = []
    for i in range(npoints):
        current_label = labels+ str(i+1)
        mylabels.append(current_label)
    plt.matshow(distance)
    plt.colorbar()
    plt.title('Pairwise Distances', {'fontsize':18, 'color':'blue'})  
    plt.xticks(np.arange(npoints), mylabels);  
    plt.yticks(np.arange(npoints), mylabels);
    plt.show()


def rand_points(npoints,ndims,labels='dot',plot=False):
    #generate npoints of random points of ndim dimension(s) following normal distribution
    import numpy as np
    locations = np.random.randn(npoints, ndims) 
    #can plot if ndim is 2d and plot is not False
    if plot !=False:    
        plot_points(points=locations,labels=labels) 
        plot_distance_mat(points=locations, labels=labels)       
    return locations


def get_distance(points,plot=False):
    import sklearn.metrics.pairwise as metrics
    locations = points
    distances = metrics.pairwise_distances(locations)
    if plot!=False:
        plot_points(points=locations) 
        plot_distance_mat(points=locations)   
    return distances


def shepard_stress(locations_input, targetdistances, npoints, ndims, plot=False):
    #Shepard's stress metric: Stress = \frac{\sum_{i,j}(d_{ij}-\delta_{ij})^2}{\sum_{ij}d_{ij}^2} $$
    import numpy as np  
    import sklearn.metrics.pairwise as metrics
    from matplotlib import pyplot as plt
    mylocations = np.reshape(locations_input, (npoints, ndims))  #convert the flat list of location parameters into a 2-D location matrix with shape (Npoints, Ndims)
    mydistances = metrics.pairwise_distances(mylocations) #compute pairwise distances using the 2-D location matrix    
    if plot !=False:
        plot_points(mylocations)
        plot_distance_mat(mylocations)
        plt.matshow(targetdistances)
    indices = np.triu_indices(npoints, 1)  #this handy helper functions tells us which indices of the distance matrix correspond to the "upper triangle" of the matrix
    dd = mydistances[indices]   #extract the upper triangle of distances based on the 2-D location matrix
    tt = targetdistances[indices]   #extract the upper triangle of target distances    
    stress = (np.dot(dd-tt,dd-tt))/np.dot(dd,dd)  
    return stress


def find_locations(targetdistances,npoints,ndims,plot=False):
    import numpy as np
    from scipy.optimize import minimize
    startinglocations = np.random.randn(npoints, ndims)   #we start our search from random locations
    optimization_result = minimize(shepard_stress, np.ravel(startinglocations), (targetdistances, npoints, ndims))  #run the optimizer
    optimized_locations = optimization_result.x   #extract the output from the optimizer
    optimized_locations = np.reshape(optimized_locations, (npoints, ndims))  # re-shape the optimized locations from flat vector into matrix shape
    if plot!=False:
        plot_points(optimized_locations)
    return optimized_locations


def points_circle(npoints=16,radius=1,phaseshift=0):
    #Matrix multiplication is a linear transformation
        #radius = 1  #radius of circle of test points
        #phaseshift = 0  #rotational shift of points on circumference
    import numpy as np
    from matplotlib import pyplot as plt
    anglestep = np.arange(0, 2*np.pi, 2*np.pi/npoints)+phaseshift  #angle between points
    # we convert from polar coordinates (radius and angle) to cartesian coordinates (x and y)
    circle_xcoords = radius*np.cos(anglestep)
    circle_ycoords = radius*np.sin(anglestep)
    plt.plot(circle_xcoords, circle_ycoords, 'o');
    plt.axis('equal');
    points = np.array([circle_xcoords, circle_ycoords])
    return points


def plot_transform(orig_points,trans_points,npoints):
    import numpy as np
    from matplotlib import pyplot as plt
    #create some handy color RGB values for later plotting
    myred = np.sin(np.linspace(0.2,1,npoints))
    mygreen = np.sin(np.linspace(0.0,1,npoints))
    myblue = np.cos(np.linspace(0.8,0.6,npoints))
    mycolors = np.vstack((myred, mygreen, myblue)).T
    #difference 
    difference = trans_points - orig_points
    #plot
    plt.rcParams["figure.figsize"] = (8,6)
    axis_scale = 3
    plt.axis([-axis_scale, axis_scale, -axis_scale, axis_scale]);
    xcoords = orig_points[0]
    ycoords = orig_points[1]
    for istep in range(npoints):
        xorig = xcoords[istep]
        yorig = ycoords[istep]
        origlines, = plt.plot(orig_points[0,istep], orig_points[1,istep], 'o', color=mycolors[istep,:])
        translines, = plt.plot(trans_points[0,istep], trans_points[1,istep], 's', color=mycolors[istep,:])
        plt.arrow(xorig, yorig, difference[0,istep], difference[1,istep], color=mycolors[istep,:], width=0.04, head_width=0.1, linestyle='dashed')
        plt.legend((origlines, translines), ('Original Points', 'Transformed Points'), loc='lower right')
        plt.xlabel('x coordinate')
        plt.ylabel('y coordinate')
        plt.axis('equal')
    plt.show()
    

def mat_transform(points,scale=1,trans_mat='none',rot='none'):
    # Returns a transformed version of the inputpoints
    # Input: a 2-by-N Numpy array of input coordinates
    # Output: a 2-by-N Numpy array of transformed coordinates
    import numpy as np
    if trans_mat != 'none':
        mymatrix = trans_mat
    elif rot != 'none':
        beta = rot
        alpha = scale
        mymatrix = alpha* np.array([
            [np.cos(beta),  -np.sin(beta)],
            [np.sin(beta) ,  np.cos(beta)]
        ])
    else:
        mymatrix = np.array([
            [scale ,  0],
            [0 ,  scale]
        ])
    out_points = np.dot(mymatrix, points)
    return out_points


def nlinear_trans_demo(metric,x='default',overlay=True):
    import numpy as np
    from matplotlib import pyplot as plt
    if x =='default':   
        x = np.arange(-200, 200) / 40
    if metric == 'tan':
        y = np.tanh(x)
    elif metric == 'sin':
        y = np.sin(x)
    elif metric == 'cos':
        y = np.cos(x)
    elif metric == 'relu':
        y = np.copy(x)
        y[y < 0] = 0
    plt.plot(x,y);
    if overlay !=True:
        plt.legend([metric]);
        plt.show()


def toLong(A):
    #A: np.ndarray
    #this function returns a "long" 1-D unraveled array from a "square" 2-D array input
    import numpy as np
    x = A.size
    return np.reshape(A, (x,1))


def toSquare(longvector):
    #this function returns a "square" 2-D array of size N by N when given an N^2 by 1 array input
    import numpy as np
    x = longvector.size
    edgeLen = int(x**0.5)
    if float(edgeLen) != x**0.5:
        print("Error: Matrix not a perfect square")
    else:
        return np.reshape(longvector, (edgeLen,edgeLen))


def heb_getWeights(x):
    #initialize the weight matrix, w, with zeros
    #the value w[i,j] is the link from node j to node i
    import numpy as np
    w = np.zeros([len(x),len(x)])
    # loop over all edges in the network, by considering all node pairs (i,j)
    for i in range(len(x)): #loop over target nodes
        for j in range(i,len(x)):  #loop over source nodes
            if i == j:   #check if the network edge is a self-connection from a node it itself
                w[i,j] = 0
            else:
                # note: the weight formula is slightly different from Hopfield,
                # because Hopfield used binary (0/1) neurons, and here we use -1/+1 neurons
                # the equations produce equivalent dynamics, but are adjusted
                # to account for using -1/+1 activation states vs 0/1 activation states
                w[i,j] = x[i]*x[j] #+1 if nodes have the same state; -1 if they have different states
                w[j,i] = w[i,j]   # ensure that weights are symmetric (i to j == j to i)
    return w


def hop_update(w,x):
    #this function picks a random node in the network and updates it
    import numpy as np
    import scipy as sci
    m = len(list(x))
    i = sci.random.randint(0,m)   #choose random index i
    u = np.dot(w[i,:], x/2.0+0.5)   #input to node i [note that we scale variable x from +1/-1 range to +1/0 here before computing dot product]
    if u > 0:
        x[i] = 1    #the "on" state of the node is +1
    elif u < 0:
        x[i] = -1   #the "off" state of the node is -1
    return x


def hop_trainPatterns(M):
    # this function accepts a 2-D array, M, where each column is a "stable state" or "memory" to be stored in the connection weights
    # it returns a 2-D matrix, w, which stores the stable states according to Hopfield's equations
    import numpy as np
    w = np.zeros([M.shape[0],M.shape[0]])
    for column in M.T:
        theseWeights = heb_getWeights(column)
        w = w + theseWeights
    return w


def get_img_data(img,convert='1/-1',plot=False):
    #example use: array = get_img_data('letters_small/A.png')
    from PIL import Image
    import numpy as np
    from matplotlib import pyplot as plt
    imageIn = Image.open(img) #this is the n pixels x n pixels, with 255-white, 0-black 
    imageArray = np.array(imageIn.getdata(0)) == 0 #convert 255/0 to False/True 
    if convert!='1/-1':
        imageArray = toLong(1*imageArray) #convert False/True to 0/1
    else:
        imageArray = toLong(2*imageArray-1) #convert False/True to -1/+1
    if plot !=False:
        plt.imshow(toSquare(imageArray));
        plt.show()
    return imageArray


def plot_states(states,num_states,num_neurons):
    #states is the num_neurons (row) x num_states (col) array
    from PIL import Image
    from matplotlib import pyplot as plt
    allPlot = plt.imshow(states, extent=[0, num_states, 0, num_neurons], aspect=num_states/num_neurons);
    allPlot.set_cmap('hot')
    plt.xlabel('Stable State to be Stored');
    plt.ylabel('Neuron #');
    plt.show()


def plot_hop_tc(dynamics,states,num_neurons,num_states,timesteps):
    from PIL import Image
    from matplotlib import pyplot as plt
    #create 3 axes for plotting
    f, (ax1, ax2, ax3) = plt.subplots(1, 3, sharey=True, figsize=(8,8))
    # image map on the left will show how the network states (columns) change over time (rows)
    dyplot = ax1.imshow(dynamics, extent=[0, timesteps, 0, num_neurons], aspect = 'auto', interpolation='none')
    plt.sca(ax1)
    plt.title('Network State over Time')
    plt.xlabel('Time Step')
    plt.ylabel('Neuron #')
    plt.colorbar(dyplot)
    # to create some spacing, we delete this middle axis -- we don't need it
    plt.delaxes(ax2)
    #this graph will show the target patterns, for comparison against the network state
    VSplot = ax3.imshow(states, extent=[0, num_states, 0, num_neurons], aspect='auto');
    plt.sca(ax3)
    plt.title('Target Patterns')
    plt.xlabel('Target Pattern #')
    plt.ylabel('Neuron #')
    plt.set_cmap('viridis')
    plt.show()


def plot_hop_match_tc(dynamics,states,state_labs=False):
    # e.g. labs = ['A','B','C','D','E']
    # how well does the network state match each of the target states?
    # let's compute the dot product ('match') between the target patterns ("allArray") and the netowrk state at each timepoint ("dynamics")
    import numpy as np
    from matplotlib import pyplot as plt
    tcs = np.dot(states.T, dynamics)
    plt.plot(tcs.T);   #plot the timecourse of the "match"
    plt.xlabel('Time Steps');
    plt.ylabel('Match Against Target');
    if state_labs!=False:
        plt.legend(state_labs);
    plt.show()


def plot_hop_startend(dynamics,num_neurons,timesteps):
    from PIL import Image
    from matplotlib import pyplot as plt
    squareSize= int(num_neurons**0.5)   #this will not work if the image is not a square
    plt.subplot(1,2,1)
    imageOut = Image.new('1', (squareSize, squareSize))
    imageOut.putdata(dynamics[:,0])
    plt.imshow(imageOut)
    plt.title('Starting State (t=0)')
    plt.subplot(1,2,2)
    imageOut = Image.new('1', (squareSize, squareSize))
    imageOut.putdata(dynamics[:,timesteps-1])
    plt.imshow(imageOut)
    plt.title('Ending State (t = ' + str(timesteps)+')');
    plt.show()


def hopfield_net(input,states,num_neurons,num_states,timesteps,plot_startend=True,plot_timecourse=True,state_labs=False,plot_wmat=True):
    import numpy as np
    import scipy as sci
    from PIL import Image
    from matplotlib import pyplot as plt
    input = np.array(input)
    input.reshape((num_neurons,1))
    dynamics = np.zeros((num_neurons,timesteps)) #this array will hold the state of all neurons as a function of time
    initial = 2*np.squeeze(input) - 1  #convert 0/1 to -1/1; this serves as the initial input
    dynamics[:,0] = initial
    Wtrained = hop_trainPatterns(states)
    for istep in range(0,timesteps-1):          #loop over time timesteps
        dynamics[:,istep+1] = hop_update(Wtrained, dynamics[:,istep])     #update network activation state at each timestep
    if plot_startend==True:
        plot_hop_startend(dynamics,num_neurons,timesteps);
    if plot_timecourse ==True:
        plot_hop_tc(dynamics,states,num_neurons,num_states,timesteps);
        if state_labs!=False:
            plot_hop_match_tc(dynamics,states,state_labs=state_labs);
        else:
            plot_hop_match_tc(dynamics,states);
    if plot_wmat ==True: #the connection matrix (i.e. weight matrix)
        plt.imshow(Wtrained);
        plt.xlabel('To Neuron #')
        plt.ylabel('From Neuron #')
        plt.colorbar();
        plt.show()
    return dynamics, Wtrained
        

##############################################################################################################################################################
############################################################ Naturalistic fMRI data processing:###############################################################
##############################################################################################################################################################


def get_nifti_data(file):
    #load fmri data
    import nibabel as nib
    import numpy as np
    data = nib.load(file).get_data()
    shape = np.shape(data)
    return data, shape

def get_nifti_info(file):
    data,shape = get_nifti_data(file)
    info, info_title=mat_info(data)
    return info

def nifti_select_time(brain_file,time_lists):
    #takes in nii file and the time (TR#) to be kept, assuming the 4th dim is time for the 4d brain img.
    #time_lists = [[start1,end1],[start2,end2],...] where start and end TRs will both be kepted/included
    #returns: sliced brain_img (nii-like), brain data (ndarray), shape
    import nibabel as nib
    import numpy as np
    #load in brain image
    fdata,shape = get_nifti_data(brain_img)
    count = 0
    for tlist in time_lists:
        count+=1
        start,end = tlist[0],tlist[1]+1        
        if count==1:
            curr_fdata = fdata[:,:,:,start:end]
        else:
            curr_fdata = np.concatenate((curr_fdata,fdata[:,:,:,start:end]),3)
    #put ndarray back to nibabel.nifti1.Nifti1Image
    targ_img = nib.Nifti1Image(curr_fdata, img.affine)
    shape = np.shape(curr_fdata)  
    return targ_img,curr_fdata,shape

def get_mask(maskn='50',plot_mask=True,mask_file=False,netn=7):
    #maskn: get the whole brain roi seg/parc mask with options of:
        #'50' for k50_2mm, 
        #'268' for n268_2mm, 
        #int: 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 for Schaefer2018, 
        #http address for seg/parc mask nii file
    #mask_file: give a local mask.nii file and output a single ROI mask
        # e.g. mask,shape,size = get_mask(mask_file='./a1_rev.nii')
    from nltools.data import Brain_Data
    from nltools.mask import expand_mask
    if mask_file==False:
        if maskn=='50':
            mask = Brain_Data('http://neurovault.org/media/images/2099/Neurosynth%20Parcellation_0.nii.gz')
        elif maskn=='268':
            mask = Brain_Data('https://neurovault.org/media/images/8423/shen_2mm_268_parcellation.nii.gz')
        elif type(maskn)==int: #schaefer2018 parcellation of 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 '
            import nilearn
            from nilearn import plotting
            parc = nilearn.datasets.fetch_atlas_schaefer_2018(n_rois=maskn, yeo_networks=netn, resolution_mm=1, data_dir=None, base_url=None, resume=True, verbose=1)
            atlas = parc.maps
            mask = Brain_Data(atlas)    
        else:#provide the http address for the mask nii file
            mask = Brain_Data(mask)
        mask_x = expand_mask(mask)
        mask_nodes = str(len(mask_x))
        if plot_mask!=False:
            mask.plot()
        return mask,mask_x,mask_nodes
    else:
        import nibabel as nib
        import numpy as np
        mask = nib.load(mask_file).get_data()
        mask_shape = np.shape(mask)
        mask_size = np.nonzero(mask)[0].shape
        return mask,mask_shape,mask_size

def plot_nifti(nii_file,mode='stats',outpath=None):
    #mode: 'stats', 'glass','blank'
    import nilearn
    import nilearn.plotting
    if mode=='stats':
        nilearn.plotting.plot_stat_map(nii_file,output_file=outpath)
    elif mode=='glass':
        nilearn.plotting.plot_glass_brain(nii_file,output_file=outpath)
    elif mode=='blank':
        nilearn.plotting.plot_img(nii_file,output_file=outpath)
    print('plotted.')

def plot_anat(nii_f,minmax=(0,1000),outpath=None):
    from nilearn.plotting import plot_anat
    plot_anat(nii_f, colorbar=True, vmin=minmax[0], vmax=minmax[1],cbar_tick_format="%i",output_file=outpath)
    print('figure plotted')

def nifti_split_hemi(mask_f,x_mid):
    import os
    #this splits the brain by left/right hemisphere
    #x_mid: the center slice for x-axis (<x_mid is left, > is right)
    mask_l = '/'.join(mask_f.split('/')[:-1])+'/'+ mask_f.split('/')[-1].split('.nii')[0] +'_left'
    mask_r = '/'.join(mask_f.split('/')[:-1])+'/'+ mask_f.split('/')[-1].split('.nii')[0] +'_right'
    os.system("fslmaths "+mask_f+" -bin -roi "+str(x_mid)+ " -1 -1 -1 -1 -1 -1 -1 "+mask_r)
    os.system("fslmaths "+mask_f+" -sub "+mask_r+" "+mask_l)
    print('left and right separately saved.')

def nifti_math(nii_f1,nii_f2=False,math='sub',outpath='default'):
    #math: sub, add, mul, div, bin, etc. from fslmath
    #if math=='sub': returns the difference  map of: nii_f1 - nii_f2 (math can be div, add, etc.)
    import os
    if nii_f2!=False:
        if outpath=='default':
            n1=nii_f1.split('/')[-1].split('.nii')[0]
            n2=nii_f2.split('/')[-1].split('.nii')[0]
            n = n1+'_'+math+'_'+n2
            out = '/'.join(nii_f1.split('/')[0:-1])+'/'+n
        else:
            out=outpath
        os.system("fslmaths "+nii_f1+" -"+math+" "+nii_f2+" "+out)
    else:
        os.system("fslmaths "+nii_f1+" -"+math+" "+out)
    print('map stored at input file1 location.')

def export_nifti(mat,outpath):
    #this exports matrix (arrays) to .nii file. But the output needs to be visualized NOT in fsleyes
    import nibabel as nib
    import numpy as np
    convert_array = np.array(mat, dtype=np.float32) 
    affine = np.eye(4)
    nif = nib.Nifti1Image(convert_array, affine)
    nib.save(nif, outpath)
    print('Nifti file saved.')

def wipe_mat(mat,retain,wipe_out=0,binary=False,dtype='int',outpath=False):
    #retain is the list of items to retain in the mat
    #wipe_out = 0 is to set all other element (items not in the retain list) to 0
    #binary = False is to set the retained item as they are, but if = 1, set them as 1 in the mat
    #dtype = 'int' is to take the int(element) for the elements in the mat (when checked against retain list)
    import numpy as np
    shape = np.shape(mat)
    vec_shape = np.prod([i for i in shape])
    vec_mat = np.reshape(mat,vec_shape)
    if dtype=='int':
        vec_mat = [int(i) for i in vec_mat]
    for i in range(len(vec_mat)):
        if vec_mat[i] not in retain:
            vec_mat[i]=wipe_out
        else:
            if binary !=False:
                vec_mat[i]=binary
    out = np.reshape(vec_mat,shape)
    if outpath !=False:
        export_nifti(out,outpath)
    return out

def make_mask(atlas_f,retain,outpath='default',resample=False,interp='nearest'):
#     example use: 
#         inds = [3,4,19,20]
#         out = '../masks/'+'amyg.nii'
#         atlas_f = '../masks/Tian_Subcortex_S2_3T_2009cAsym.nii' #this has better be the original resolution atlas nii
#         make_mask(atlas_f,retain=inds,outpath=out)
#     if resample is not false, needs to be a nifti file of the target dimension
    import os
    files = []
    #for all the indices to be retained as 1
    for i in retain:
        temp = 'parcel_'+str(i)+'.nii'
        files.append(temp)
        os.system("fslmaths "+atlas_f+" -thr "+str(i)+" -uthr "+str(i)+" -bin "+temp)
    #specify the outputs
    if outpath=='default':
        out = 'parcel_'+'_'.join([str(i) for i in retain])
    else:
        out=outpath
    #add all the parcels
    sub_par = ' '.join(files)
    os.system("fsladd "+out+" "+sub_par)
    os.system("gzip -d "+out)
    #delete the intermediate files
    for f in files:
        os.system("rm "+f+'.gz')
    os.system("rm "+out+'.gz')
    if resample!=False:
        resample_map(out,resample,outpath=out,interp=interp)
    print('done.')

def apply_mask(data_file,mask_file,keep=1,pop=0):
    #this function returns the masked 4d time-series and the shape of it 
    #e.g. (607,616) where 607 is the #of kept voxels in the mask, 616 is the #of volumns
        #data_file = './fmri.nii'
        #mask_file = './a1_rev.nii'
        #keep = 1 this is to keep the voxels in mask file that are 1
        #if keep ==False, set pop to specify values unwant from the mask, e.g. 0 or Nan etc.
    if type(data_file)==np.memmap or type(data_file)==np.ndarray:
        data = data_file
    else:
        data,data_shape = get_nifti_data(data_file)
    mask,mask_shape,mask_size= get_mask(mask_file=mask_file)
    if keep!=False:
        masked_data = data[mask==keep]
    else:
        masked_data = data[mask!=pop]
    shape = np.shape(masked_data)
    return masked_data,shape

def lagcorr(swav,a1tc,unit=30,yrng=(-25,25),shift=-3,outpath=False,show_fig=True,limit=[-6,6],title='a1tc-audio'):
    #this is to check the a1tc match with audio signal
    #swav: a list/array of numbers reflecting sound wave signal (either rms or audio envelope)
    #a1tc: primary auditory cortex activity (a1 timecourse)
    #unit: shift the a1tc or swav by how many TRs
    #returns the peak location (-3 means shift a1tc forward in time by 3TRs to match with audio: delay/HRF)
    all_r = []
    all_p = []
    aud_env = a1tc
    #neg: shift window on a1tc right by 0-60 TR(s) to match with the start of the swav
    for ind in range(-unit,0):
        if ind ==0:
            continue
        else:
            array1 = swav[:ind]
            array2 = aud_env[-ind:]
        r,p=get_corr(array1,array2)
        all_r.append(float(r))
        all_p.append(float(p))
    #pos: shift window on swav right by 0-60 TR(s) to match with the start of the a1tc
    for ind in range(unit+1):
        if ind ==0:
            array1 = swav
            array2 = aud_env
        else:
            array1 = swav[ind:]
            array2 = aud_env[:-ind]
        r,p=get_corr(array1,array2)
        all_r.append(float(r))
        all_p.append(float(p))
    if limit==None or limit==False:
        loc,peak = find_peak_latency(all_r,time=range(-unit,unit+1),window=[-unit,unit])
    else:
        start = abs(-unit-(limit[0]))
        lim_len = len([i for i in range(limit[0],limit[1]+1)])
        loc,peak = find_peak_latency(all_r[start:start+lim_len],time=range(limit[0],limit[1]+1),window=[limit[0],limit[1]])
        # loc,peak = find_peak_latency(all_r,time=range(-unit,unit+1),window=[limit[0],limit[1]])
    plot_line(all_r,x=[i for i in range(-unit,unit+1)],title=title+' corr: '+str(peak)+' at '+str(loc),x_lab='TR shift',add=True)
    plt.plot([loc]*(yrng[1]-yrng[0]),[i/100 for i in range(yrng[0],yrng[1])])
    if outpath !=False:
        plt.savefig(outpath+'.png')
    if show_fig==True:
        plt.show()
    return loc,peak

def auto_correlation(tc,unit=30):
    #this is to check the tr of shift it takes for a time-series's auto-correlation fall <= half-max (r=0.5)
    #tc: timecourse to check autocorrelation of
    #unit: shift the a1tc or swav by how many TRs
    a1tc = tc
    all_r = []
    all_p = []
    #pos: shift window on swav right by 0-60 TR(s) to match with the start of the a1tc
    for ind in range(unit+1):
        if ind ==0:
            array1 = a1tc
            array2 = a1tc
        else:
            array1 = a1tc[ind:]
            array2 = a1tc[:-ind]
        r,p=get_corr(array1,array2)
        all_r.append(float(r))
        all_p.append(float(p))
    return all_r

def auto_correlation_halfmax(tc,unit=30,half_max=0.5,interpolate=False,show_fig=True,bg_style='seaborn-whitegrid',outpath=False,title='Auto-correlation_half-maximum'):
    #given a timeseries, compute the auto-correlation with 0 to 30 tr(s) shift, and find the first tr that falls to or below the half-max (0.5) 
    allr = auto_correlation(tc,unit=unit)
    halfmax_tr = find_halfmax_point(allr,half_max=half_max,interpolate=interpolate)
    if show_fig==True:
        plot_autocorr_halfmax(allr,halfmax_tr,outpath=outpath,title=title,bg_style=bg_style)        
    return halfmax_tr,allr

def extract_mvp_raw(data_file,mask_file,trs='all',outpath=False,transpose=False):
    #this is to return the mvp of that masked region for each given TRs
        #data_file='/preproced_scan.nii' the timeseries that TRs corresponds to.
        #mask_file='roi_mask.nii' 
        #trs=[23,24,67,68,69] these are volume# to extract the MVP from 
    #returns: vox# x tr#
    masked_data, shape = apply_mask(data_file,mask_file)
    mvp_list = masked_data
    if trs!='all': #otherwise, trs need to be a list of integer
        trs_inds = (np.array(trs)-1).tolist()
        mvp_list = apply_inds(trs_inds,masked_data,ind_pos=1)
    if outpath!=False:
        fname = data_file.split('/')[-1].split('.')[0] + '_' + mask_file.split('/')[-1].split('.')[0]+'_mvp.csv' 
        write_csv(mvp_list,outpath+fname,transpose=transpose)
    return mvp_list    
    
def extract_mvp(data_file,mask_file,trs,mute=True):
    #this is to return the mvp of that masked region averaged across TRs
        #data_file='/preproced_scan.nii' the timeseries that TRs corresponds to.
        #trs=[23,24,67,68,69] these are volume# to extract the MVP from and average across
        #mask_file='roi_mask.nii' 
    #notes: 
        #masked_data is supposed to be voxel# x TR#; 
        #ind_pos=1 instead of 0 because applying to TR
        #trs_masked_data is supposed to be selected TR# x voxel#
        #avg across dim=0 because averaging across TR
    masked_data, shape = apply_mask(data_file,mask_file)
    trs_inds = (np.array(trs)-1).tolist()
    trs_masked_data = apply_inds(trs_inds,masked_data,ind_pos=1)
    mvp = avg(trs_masked_data,0,mute=mute)
    return mvp 

def extract_group_mvp(data_file,mask_file,trs,mute=True):
    #this is to return the mvp of that masked region averaged across subjects and their respective TRs
    allsub_mvp = []
    for i in range(len(data_file)):
        mvp = extract_mvp(data_file[i],mask_file,trs[i])
        allsub_mvp.append(mvp)
    out = avg(allsub_mvp,0,mute=mute)
    return out

def extract_voi_masked(data_file,mask_file,mute=False):
    #returns the time-series at masked region (averaged across voxels within the mask at each TR)
    mvp_ls = extract_mvp_raw(data_file,mask_file)
    return avg(mvp_ls,mute=mute)

def extract_voi(data_file,maskn=200,netn=17,outfile=False,to_csv=True):
    #data_file: the nifti file
    #maskn: the number of ROIs the whole brain is desected to
    #outfile: the path and filename to save the output data: nVOI x nTR
    #to_csv: if True, save output as .csv, else save as .xslx
    import pandas as pd
    from nltools.data import Brain_Data
    mask, mask_x, mask_nodes = get_mask(maskn=maskn,netn=netn)
    voi = Brain_Data(data_file).extract_roi(mask)
    if to_csv==False: #save to .xslx
        if outfile==True:
            write_ndarray(voi,'./voi'+str(maskn)+'.xlsx')
        elif outfile!=False and outfile!=True:
            write_ndarray(voi,outfile)
    else: #save to .csv
        if outfile==True:
            pd.DataFrame(voi.T).to_csv('./voi_n'+str(maskn)+'.csv',index=False)
        elif outfile!=False and outfile!=True:
            pd.DataFrame(voi.T).to_csv(outfile,index=False)
    return voi

def load_voi(voi_file,ftype='.csv',exclude=[]):
    #load single subject's func session's voi data
    #panda auto reads in csv with the first row being the header
    #exclude needs to be row#-1 for indexing, e.g. to exclude the first two actual rows (skip header), exclude=[0,1]
    #returns a dataframe of nVOI x nTR 
    if ftype=='.csv':
        import pandas as pd
        func_data = pd.read_csv(voi_file)
        func_data=func_data.drop(exclude)
    return func_data

def load_group_voi(voi_folder,ftype='.csv',exclude=[],file_list=False):
    #voi_folder is the path to all the voi files to be loaded
    #file_list is the specified list of files to load; if False, will load all files in voi_folder
    #load all subjects' func session's voi data
    #returns sub_timeseries: a dictionary with subjects as keys, and each dic[sub] is a dataframe
    sub_timeseries = {}
    if file_list==False:
        flist = grab_all(voi_folder,ftype)
    else:
        flist = file_list
    for file in flist:
        voi_file = voi_folder+file
        sub_data = load_voi(voi_file,exclude=exclude)
        sub = file.split('_')[0]
        sub_timeseries[sub]=sub_data
    return sub_timeseries





################################################################# fMRI GLM-related ###########################################################

def coord_mni2vox(x,y,z,to_int=False):
    #Only valid for 3mm nifti files of (61,73,61) dimension 
    #this is to to transform the 3d MNI coordinate (negative to postiive) to voxel coordinate (all positive indexing)
    #returns (x,y,z) of the transformed coordinate
    x = x*(1/3)+30
    y = y*(1/3)+42
    z = z*(1/3)+24
    if to_int!=False:
        x,y,z = int(round(x)),int(round(y)),int(round(z))
    return (x,y,z)

def coord_vox2mni(x,y,z,to_int=False):
    #Only valid for 3mm nifti files of (61,73,61) dimension 
    #this is to transform the 3d voxel coordinate (all positive indexing) to MNI coordinate (negative to postiive).
    #returns (x,y,z) of the transformed coordinate
    x = x*3-90
    y = y*3-126
    z = z*3-72
    if to_int!=False:
        x,y,z = int(round(x)),int(round(y)),int(round(z))
    return (x,y,z)

def resample_map(source_img,targ_img,outpath=False,interp='nearest'):
    #this is to align the shape of the different nifti files (they should already be in the same space, but just need to be resampled)
    #source_img and targ_img are both nii file or nii-like object loaded from nibabel; function returns the resampled nii-like object
    # e.g. source_img is a subcortical atlas of (91, 109, 91) shape, targ_img is a 3mm MNI template of (61, 73, 61)
        #this function will then return the subcortical atlas with the shape of (61, 73, 61) 
    import nilearn
    import nibabel as nib
    if type(source_img)!=nib.nifti1.Nifti1Image:
        source_img = nib.load(source_img)
    if type(targ_img)!=nib.nifti1.Nifti1Image:
        targ_img = nib.load(targ_img)
    resampled_img = nilearn.image.resample_to_img(source_img,targ_img,interpolation=interp)
    if outpath!=False:
        resampled_img.to_filename(outpath)
    return resampled_img

def xyz_label_name(cort_num, subc_num, cort_labtxt,subc_labtxt):
    #cort_num: the label number in the cortical atlas where 0 is out of bound
    #subc_num: the label number in the subcortical atlas where 0 is out of bound
    #cort_labtxt = '../masks/use_labels/tian_subcortex32_label.txt'
    #subc_labtxt =  '../masks/use_labels/schaefer_cortex200_net7_label.txt'
    #returns the label name of the corresponding label number for cortical and subcortical maps
    if type(cort_labtxt)!=list:
        cort_labtxt = read_text(cort_labtxt,sep='\n')
        subc_labtxt = read_text(subc_labtxt,sep='\n')
    if cort_num==0:
        cort_lab ='NA'
    else:
        cort_lab = cort_labtxt[cort_num-1]
    if subc_num==0:
        subc_lab = 'NA'
    else:
        subc_lab = subc_labtxt[subc_num-1]
    return cort_lab,subc_lab    

def xyz_label_num(xyz, atlas_cort_img, atlas_subc_img,xyz_to_vox=True):
    # xyz = (x,y,z) where x y z should be numbers
    # atlas_cort_img = '../masks/atlas-cortex_schaefer400net17.nii'
    # atlas_sub_img = '../masks/atlas-subcortex_tian32s2_nearest'
    # xyz_to_vox: if true, transform the coordinate xyz from MNI space to voxel space
    # returns the atlas label number for coordinate xyz
    cort_map,_,_ = get_mask(mask_file=atlas_cort_img)
    subc_map,_,_ = get_mask(mask_file=atlas_subc_img)
    X,Y,Z = xyz[0],xyz[1],xyz[2]
    vox_xyz = [int(i) for i in coord_mni2vox(X,Y,Z)]
    cort_num = int(cort_map[vox_xyz[0],vox_xyz[1],vox_xyz[2]])
    subc_num = int(subc_map[vox_xyz[0],vox_xyz[1],vox_xyz[2]])
    return cort_num,subc_num
    
def xyz_label_info(xyz, atlas_cort_img, atlas_subc_img,cort_labtxt,subc_labtxt, xyz_to_vox=True):
    cort_num,subc_num = xyz_label_num(xyz, atlas_cort_img, atlas_subc_img,xyz_to_vox=xyz_to_vox)
    cort_lab,subc_lab = xyz_label_name(cort_num, subc_num, cort_labtxt, subc_labtxt)
    return cort_num,cort_lab, subc_num,subc_lab
    
def add_conditions(design_mat,factor_list='default',show=False,outpath=False,add_cons=False):
    #this is for fmri glm, to add in the conditions array 
    #outputs a dict of conditions with key = cond or contrast, value = corresponding one-hot vec array
    from numpy import array
    key_list = [i for i in design_mat.keys()]
    if factor_list=='default':
        factor_list = [i for i in key_list if i.find('drift')==-1 and i.find('constant')==-1]
    conditions = init_dict(factor_list)
    count=0
    for f in factor_list:
        init_zero = init_mat((1,len(key_list)))
        init_zero[0,count]=1
        conditions[f]=init_zero
        count+=1
    cond_list = factor_list
    if show!=False:
        plot_conds(conditions,cond_list,design_mat,outpath=outpath)
    if add_cons!=False:
        conditions,cond_list = add_contrasts(design_mat,conditions,factor_list=factor_list,show=show,outpath=outpath)
    return conditions,cond_list

def add_contrasts(design_mat,conditions,factor_list='default',show=False,outpath=False,fig_format='.png'):
    key_list = [i for i in design_mat.keys()]
    if factor_list=='default':
        factor_list = [i for i in key_list if i.find('drift')==-1 and i.find('constant')==-1]
    pairs_ind_up,pairs_name_up = get_triu_coords(factor_list,lists_name=factor_list,diag_offset=1)
    pairs_ind_low,pairs_name_low = get_tril_coords(factor_list,lists_name=factor_list,diag_offset=1)
    pairs_ind = merge_lists_in_list([pairs_ind_up,pairs_ind_low])
    pairs_name = merge_lists_in_list([pairs_name_up,pairs_name_low])    
    cond_list = factor_list
    count=0
    for i in pairs_ind:
        pair_name = pairs_name[count]
        factor1 = factor_list[i[0]]          
        factor2 = factor_list[i[1]]
        conditions[pair_name] = conditions[factor1]-conditions[factor2]
        cond_list.append(pair_name)
        count+=1
        if show!=False:
            import numpy as np
            import matplotlib.pyplot as plt
            from nilearn.plotting import plot_contrast_matrix
            con = np.vstack((conditions[factor1],conditions[factor2])) 
            plot_contrast_matrix(con, design_mat)
            if outpath!=False:
                plt.savefig(outpath+i+fig_format)
            plt.show()
    return conditions,cond_list

def plot_conds(conditions,cond_list,design_mat,outpath=False,fig_format='.png'):
    #e.g. cond_list = ['Pre','Post','Post-Pre']
    import matplotlib.pyplot as plt
    from nilearn.plotting import plot_contrast_matrix
    for i in cond_list:
        cond = conditions[i]
        plot_contrast_matrix(cond, design_mat)
        if outpath!=False:
            plt.savefig(outpath+i+fig_format)
        plt.show()
    return 'figure shown.'

def plot_design_mat(design_mat,outpath=False):
    from nilearn.plotting import plot_design_matrix
    import matplotlib.pyplot as plt
    if outpath!=False:
        mkdir(outpath)
        plot_design_matrix(design_mat,output_file=outpath)
        mesg='design matrix plotted and saved.'
    else:
        plot_design_matrix(design_mat)
        mesg='design matrix plotted.'
    plt.show()
    return mesg
    
def model_glm(brain_img,events,tr,title=False,hrf='spm',std=False,hp=.01,fwhm=None,roi_mask=None,regressors=None,drift='cosine',outpath=False):
    #takes in brain image, event file (onset/duration/trial_type/modulation)
        #uses nilearn.glm.first_level.FirstLevelModel to specify the GLM model and uses the .fit to fit brain data to model
        #e.g. tr is the sampling rate, e.g. tr=1.5, or tr=0.8 etc.
        #returns fmri_glm and design_matrix
    #alternative to the current function which : 
        #use nilearn.glm.first_level.make_first_level_design_matrix and .run_glm
        #make_first_level_design_matrix: lets you make the design matrix with add_reg and add_reg_name (more flexible)
        #run_glm: takes in Y and X where Y is an ndarray of (n_voxels,n_time) and X is the design matrix with all the parameters
    import nibabel as nib
    from nilearn.glm.first_level import FirstLevelModel    
    import matplotlib.pyplot as plt
    #load in the brain image
    if type(brain_img)!=nib.nifti1.Nifti1Image:
        img = nib.load(brain_img)
    else:
        img = brain_img
    if roi_mask!=None:
        mask_img = nib.load(roi_mask)
    #model glm
    fmri_glm = FirstLevelModel(t_r=tr, hrf_model=hrf, standardize=std, drift_model=drift, high_pass=hp, smoothing_fwhm=fwhm,mask_img=roi_mask)
    fmri_glm = fmri_glm.fit(img, events,confounds=regressors)
    design_matrix = fmri_glm.design_matrices_[0]
    #plot design matrix
    if outpath!=False:
        if title!=False:
            save_target = outpath+title+ '_design_matrix.png'
        else:
            save_target = outpath+'design_matrix.png'
        plot_design_mat(design_matrix,outpath=save_target)
    else:
        plot_design_mat(design_matrix)
    #plot predicted bold response for each factor
    keys = [i for i in design_matrix.keys()]
    factors = [i for i in design_matrix.keys() if i.find('drift')==-1 and i.find('constant')==-1]
    for f in factors:
        plt.plot(design_matrix[f])
        plt.xlabel('Time')
        plt.title(f)
        if outpath!=False:
            if title!=False:
                save_target = outpath+title+ '_modeled_response_'+f+'.png'
            else:
                save_target = outpath+ 'modeled_response_'+f+'.png'
            plt.savefig(save_target)
        plt.show()
    return fmri_glm,design_matrix

def thresh_map(stat_map,clust_thresh=20,thresh=1.,alpha=.05,mcc='fdr',outpath=False,show=False):
    #stat_map needs to be a nii file or nii-like object
    import nibabel as nib
    from nilearn.glm import threshold_stats_img
    if type(stat_map)!=nib.nifti1.Nifti1Image:
        stat_map = nib.load(stat_map)    
    clean_map, threshold = threshold_stats_img(stat_map, alpha=alpha, height_control=mcc, cluster_threshold=clust_thresh)
    if show!=False:
        if thresh=='sig':
            thr = threshold
        else:
            thr = thresh            
        view_brain(clean_map,view_list=['lateral','medial'],thresh=thr,outpath=outpath)
    return clean_map,threshold

def show_glm_results(z_map,alpha=.05,mcc='fdr',clust_thresh=20,title=None,outpath=False,fig_format='.png',view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior'],thresh=1.):
    # this is to show global activation given the z_map (a nii-like object)
    import nibabel as nib
    from nilearn.glm import threshold_stats_img
    from nilearn.plotting import plot_stat_map, plot_anat, plot_img
    import matplotlib.pyplot as plt
    #read in z_map
    if type(z_map)!=nib.nifti1.Nifti1Image:
        z_map = nib.load(z_map)
    clean_map, threshold = threshold_stats_img(z_map, alpha=alpha, height_control=mcc, cluster_threshold=clust_thresh)
    if title==None:
        fig_title = 'fdr'+str(alpha)+'_clust'+str(clust_thresh)+'vox'
    else:
        fig_title = title+'_fdr'+str(alpha)+'_clust'+str(clust_thresh)+'vox'
    plot_stat_map(clean_map, bg_img=z_map, threshold=threshold,display_mode='z', cut_coords=3, black_bg=True,title=fig_title)
    if outpath!=False:
        mkdir(outpath)
        plt.savefig(outpath+fig_title+fig_format)
    plt.show()
    if thresh=='sig':
        thr=threshold
    else:
        thr=thresh
    view_brain(clean_map,outpath=outpath+'globalview_'+fig_title+'_',fig_format=fig_format,thresh=thr,view_list=view_list)
    return threshold

def save_glm_results(conditions,fmri_glm,outpath,show=True,title=None,alpha=.05,mcc='fdr',clust_thresh=20,thresh=1.,view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior']):
    #this is to save the glm results (global activation per condition and contrast) from model_glm and add_conditions function
    from nilearn.reporting import get_clusters_table
    for cond in conditions:
        if title==None:
            fname = cond
        else:
            fname = title+'_'+cond
        mkdir(outpath)
        eff_map = fmri_glm.compute_contrast(conditions[cond],output_type='effect_size')
        z_map = fmri_glm.compute_contrast(conditions[cond],output_type='z_score')
        z_map.to_filename(outpath+'zmap_'+fname+'.nii.gz')
        eff_map.to_filename(outpath+'betamap_'+fname+'.nii.gz') 
        if show!=False:
            threshold = show_glm_results(z_map,alpha=alpha,mcc=mcc,clust_thresh=clust_thresh,title=fname,outpath=outpath,view_list=view_list,thresh=thresh)
            table = get_clusters_table(z_map, stat_threshold=threshold,cluster_threshold=20)
            table.to_csv(outpath+'table_'+fname+'.csv')
    return 'saved.'







################################################################# fMRI processing and ISC-related ###########################################################

def make_motion_covariates(mc, tr):
    #this is only to be called from post_fmriprep
    z_mc = zscore(mc)
    all_mc = pd.concat([z_mc, z_mc**2, z_mc.diff(), z_mc.diff()**2], axis=1)
    all_mc.fillna(value=0, inplace=True)
    return Design_Matrix(all_mc, sampling_freq=1/tr)

def post_fmriprep(base_dir,sub_list,tr=2,fwhm=6,outlier_cutoff=3,hdf5=False,do_roi=False):
    #This function:
        #smooth and denoise/despike (will skip if targte files already exist)
        #optional: make a hdf5 version file of the *smooth_denoised*.nii file
        #optional: extract voi from the nii file and save to csv 
    #tr: change tr to suit the study if were to do smooth and denoise
    #fwhm: smooth kernel
    #outlier_cutoff: for finding spikes
    #hdf5: for making hdf5 (fast loading for nltools) from nii
    #do_roi: for extracting voi, with options of:
        #'50' for k50_2mm, 
        #'268' for n268_2mm, 
        #int: 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 for Schaefer2018, 
    import os
    import glob
    import numpy as np
    import pandas as pd
    import matplotlib.pyplot as plt
    import seaborn as sns
    from nltools.stats import regress, zscore
    from nltools.data import Brain_Data, Design_Matrix
    from nltools.stats import find_spikes 
    from nltools.mask import expand_mask
    #loop through all subjects in list
    for sub_num in sub_list:
        sub= 'sub-'+sub_num  
        mesg = 'Started with '+sub+'...'
        print(mesg)
        path = os.path.join(base_dir,sub,'func/*preproc*')  
        sub_files = grab_all(path,'gz',exclude='denoise',mute=True)  
        output_list = insert_str(sub_files,'_task-',str('_denoise_smooth'+str(fwhm)+'mm'))
        print(output_list)
        for file_name in output_list:
            file = os.path.join(base_dir,sub,'func',file_name) 
            if os.path.exists(file):
                print(str('Smoothed and denoised file already exists for '+file_name))
                continue
            else:
                file_name = delete_str([file_name],str('_denoise_smooth'+str(fwhm)+'mm'))[0]
                print(str('Starting to smooth and denoise for: '+file_name))
                file = os.path.join(base_dir,sub,'func',file_name) 
                task = file[file.find('_task-')+6:file.find('_space-')]
                data = Brain_Data(file)
                smoothed = data.smooth(fwhm=fwhm)
                spikes = smoothed.find_spikes(global_spike_cutoff=outlier_cutoff, diff_spike_cutoff=outlier_cutoff)
                cov_tsv = base_dir+sub+'/func/'+sub+'_task-'+task+'_desc-confounds_regressors.tsv'  
                covariates = pd.read_csv(cov_tsv, sep='\t')
                mc = covariates[['trans_x','trans_y','trans_z','rot_x', 'rot_y', 'rot_z']]
                mc_cov = make_motion_covariates(mc, tr)
                csf = covariates['csf'] 
                dm = Design_Matrix(pd.concat([csf, mc_cov, spikes.drop(labels='TR', axis=1)], axis=1), sampling_freq=1/tr)
                dm = dm.add_poly(order=2, include_lower=True) # Add Intercept, Linear and Quadratic Trends
                smoothed.X = dm
                stats = smoothed.regress()
                stats['residual'].data = np.float32(stats['residual'].data) # cast as float32 to reduce storage space
                stats['residual'].write(os.path.join(base_dir, sub, 'func', f'{sub}_denoise_smooth{fwhm}mm_task-{task}_space-MNI152NLin2009cAsym_desc-preproc_bold.nii.gz'))
                print(str('Smooth and denoise completed for '+ file_name))
        print(output_list)
        for file_name in output_list:
            task = file[file.find('_task-')+6:file.find('_space-')]
            file = os.path.join(base_dir,sub,'func',file_name) 
            if hdf5==True:
                print(str('Saving as hdf5 file for : '+file_name))
                Brain_Data(file).write(os.path.join(base_dir, sub, 'func', f"{file_name.split('.nii.gz')[0]}.hdf5"))
                print('hdf5 file saved for nltools')
            if do_roi!=False:
                mask, mask_x, mask_nodes = get_mask(maskn=do_roi)
                roi = Brain_Data(file).extract_roi(mask)
                pd.DataFrame(roi.T).to_csv(os.path.join(base_dir,sub,'func',f"{sub}_{task}_Average_ROI_n{mask_nodes}.csv" ), index=False)
                print('ROI file saved')
    print('Post-fmriprep processing completed.') 


def make_hdf5(nii_file,out='default'):
    #out='default' is to save at the same place as the input file
    from nltools.data import Brain_Data
    out= nii_file.replace('.nii.gz','.hdf5')
    if file_exist(out)==False:
        Brain_Data(nii_file).write(out)
        print('hdf5 file saved for nltools')
    else:
        print('hdf5 file already exist.')


def hdf5_apply_mask(fmriprep_dir,roi_ind,subj='*',name_in_file='*hdf5',maskn='50'):
    #example use: 
        #all_data = hdf5_apply_mask(base_dir,4,name_in_file=name)
        #base_dir= './Sherlock/fmriprep/'
        #name= '*crop*Part2*hdf5' (defaults to grab all hdf5 files in the directory)
        #maskn: 
            #'50' for k50_2mm, 
            #'268' for n268_2mm, 
            #int: 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 for Schaefer2018, 
    import os
    import glob
    from nltools.mask import create_sphere, expand_mask
    from nltools.data import Brain_Data   
    import warnings
    warnings.simplefilter('ignore')
    #get mask
    mask,mask_x,mask_nodes = get_mask(maskn=maskn)
    #check ROI
    roi_mask = mask_x[roi_ind]
    roi_mask.plot()
    #apply mask to hdf5 files
    file_list = glob.glob(os.path.join(fmriprep_dir, subj, 'func', f'{name_in_file}'))
    all_data = []
    file_list.sort()
    for f in file_list:
        sub = os.path.basename(f).split('_')[0]
        print(sub)
        data = Brain_Data(f)
        all_data.append(data.apply_mask(roi_mask))
    return all_data


def plot_hyper_isc(data,hyperalign=False,transformed_list=False,voxel_ind=0,tr_ind=0,nsubj=5,lab_ftsize=16,title_ftsize=18,wid=15,hgt=5,map_color='RdBu_r'):
    #exmaple use:
        #plot_hyper_align(all_data[:15],hyperalign)
        #all_data[:15] = hdf5_apply_mask(base_dir,4,name_in_file=name)[:15] (i.e. the first 15 subjects)
        #hyperalign = align(all_data[:15], method='procrustes')
    import pandas as pd
    import numpy as np
    import seaborn as sns
    import matplotlib.pyplot as plt
    from nltools.data import Adjacency
    import warnings
    warnings.simplefilter('ignore')    
    if hyperalign!=False:
        transformed = hyperalign['transformed']
    if transformed_list!=False:
        transformed = transformed_list
    if hyperalign==False and transformed_list==False:
        print('ERROR: should input either the hyperalign dictionary or the transformed list')
    else:
        voxel_unaligned = pd.DataFrame([x.data[:, voxel_ind] for x in data]).T
        voxel_aligned = pd.DataFrame([x.data[:, voxel_ind] for x in transformed]).T
        #see the ISC for a specified voxel
        plot1 = plt.figure(1)
        print(str('ISC for Voxel Index-'+str(voxel_ind)+':'))
        f, a = plt.subplots(nrows=2, figsize=(wid, hgt), sharex=True)
        a[0].plot(voxel_unaligned, linestyle='-', alpha=.2)
        a[0].plot(np.mean(voxel_unaligned, axis=1), linestyle='-', color='navy')
        a[0].set_ylabel('Unaligned Voxel', fontsize=lab_ftsize)
        a[0].yaxis.set_ticks([])
        a[1].plot(voxel_aligned, linestyle='-', alpha=.2)
        a[1].plot(np.mean(voxel_aligned, axis=1), linestyle='-', color='navy')
        a[1].set_ylabel('Aligned Voxel', fontsize=lab_ftsize)
        a[1].yaxis.set_ticks([])
        plt.xlabel(str('Voxel Index-'+str(voxel_ind)+' Time Course (TRs)'), fontsize=lab_ftsize)
        a[0].set_title(f"Unaligned Voxel ISC: r={Adjacency(voxel_unaligned.corr(), matrix_type='similarity').mean():.02}", fontsize=title_ftsize)
        a[1].set_title(f"Aligned Voxel ISC: r={Adjacency(voxel_aligned.corr(), matrix_type='similarity').mean():.02}", fontsize=title_ftsize)
        plt.show()
        #see the impact of hyperalignment on the spatial topography
        plot2 = plt.figure(3)
        print(str('Spatial Topography for TR Index-'+str(tr_ind)+':'))
        f,a = plt.subplots(ncols=nsubj, nrows=2, figsize=(wid,hgt), sharex=True, sharey=True)
        for i in range(nsubj):
            sns.heatmap(np.rot90(data[i][tr_ind].to_nifti().dataobj[30:60, 10:28, 37]), cbar=False, cmap='RdBu_r', ax=a[0,i])
            a[0,i].set_title(f'Subject: {i+1}', fontsize=18)
            a[0,i].axes.get_xaxis().set_visible(False)
            a[0,i].yaxis.set_ticks([])
            sns.heatmap(np.rot90(transformed[i][tr_ind].to_nifti().dataobj[30:60, 10:28, 37]), cbar=False, cmap=map_color, ax=a[1,i])
            a[1,i].axes.get_xaxis().set_visible(False)
            a[1,i].yaxis.set_ticks([])
        a[0,0].set_ylabel('Unaligned Voxels', fontsize=lab_ftsize)
        a[1,0].set_ylabel('Aligned Voxels', fontsize=lab_ftsize)
        plt.tight_layout()
        plt.show()
        if hyperalign!=False:
            #see the ISC after hyperalignment for all voxels
            fig_num = 2+(nsubj*2)+1
            plot3 = plt.figure(fig_num)
            print(str('ISC after hyperalignment for all voxels:'))
            plt.hist(hyperalign['isc'].values())
            plt.axvline(x=np.mean(list(hyperalign['isc'].values())), linestyle='--', color='red', linewidth=2)
            plt.ylabel('Frequency', fontsize=lab_ftsize)
            plt.xlabel('Voxel ISC Values', fontsize=lab_ftsize)
            plt.title('Hyperalignment ISC', fontsize=title_ftsize)
            print(f"Mean ISC: {np.mean(list(hyperalign['isc'].values())):.2}")
            plt.show()


def plot_srm_isc(srm=False,transformed_list=False,component_ind = 0,nsubj=4,wid=15,hgt=8,lab_ftsize=16,title_ftsize=18,map_color='RdBu_r'):
    #srm is the output dict from the align functiona, e.g. srm = align(all_data_z, method='deterministic_srm', n_features=100)
    #component_ind is the index number of the component you want to plot, e.g. 0,2,...99 (note that component1 is of index 0)
    #nsubj is the number of subjects you want to show in spatial plot
    #wid is the width of the spatial plot
    #hgt is the height of the spatial plot
    #lab_ftsize is the label's fontsize of the x- and y-axis
    #title_ftsize is the title fontsize
    import pandas as pd
    import numpy as np
    import seaborn as sns
    import matplotlib.pyplot as plt
    import matplotlib.gridspec as gridspec
    from nltools.data import Adjacency
    import warnings
    warnings.simplefilter('ignore')
    if srm!=False:
        component_aligned = pd.DataFrame([x[:, component_ind] for x in srm['transformed']]).T
    if transformed_list!=False:
        component_aligned = pd.DataFrame([x.data[:, component_ind] for x in transformed_list]).T
    if srm==False and transformed_list==False:
        print('ERROR: should input either the srm dictionary or the transformed list')
    else:
        #Plot Aligned Latent Component Time Course
        plot1 = plt.figure(1)
        print(str('ISC for Component Index-'+str(component_ind)+':'))        
        f, a = plt.subplots(nrows=1, figsize=(15, 5), sharex=True)
        a.plot(component_aligned, linestyle='-', alpha=.2)
        a.plot(np.mean(component_aligned, axis=1), linestyle='-', color='navy')
        a.set_ylabel('Aligned Component', fontsize=lab_ftsize)
        a.yaxis.set_ticks([])
        plt.xlabel('Component Time Course (TRs)', fontsize=lab_ftsize)
        a.set_title(f"Aligned Component ISC: r={Adjacency(component_aligned.corr(), matrix_type='similarity').mean():.02}", fontsize=title_ftsize)
        plt.show()
        if srm!=False:
            #Plot Subject Spatial Patterns associated with each latent component
            plot2 = plt.figure(3)
            print(str('Spatial Topography for Component Index-'+str(component_ind)+':'))
            f = plt.figure(constrained_layout=True, figsize=(wid,hgt))
            spec = gridspec.GridSpec(ncols=nsubj, nrows=nsubj, figure=f)
            for i in range(nsubj):
                a0 = f.add_subplot(spec[i, 0])
                a0.imshow(np.rot90(srm['transformation_matrix'][i][component_ind].to_nifti().dataobj[30:60, 10:28, 37]),cmap=map_color)
                a0.set_ylabel(f'Subject {i+1}', fontsize=lab_ftsize)
                a0.yaxis.set_ticks([])
                a0.xaxis.set_visible(False)    
                a1 = f.add_subplot(spec[i, 1:])
                a1.plot(srm['transformed'][i][:,component_ind])
                a1.xaxis.set_visible(False)
                a1.yaxis.set_visible(False)
                if i < 1:
                    a0.set_title('Spatial Pattern', fontsize=lab_ftsize)
                    a1.set_title('Latent Timecourse', fontsize=lab_ftsize)
            plt.show()
            #overall distribution of ISC across all components
            fig_num = 2+(nsubj*2)+1
            plot3 = plt.figure(fig_num)
            print('ISC for all Components:')
            plt.hist(srm['isc'].values())
            plt.axvline(x=np.mean(list(srm['isc'].values())), linestyle='--', color='red', linewidth=2)
            plt.ylabel('Frequency', fontsize=lab_ftsize)
            plt.xlabel('Voxel ISC Values', fontsize=lab_ftsize)
            plt.title('Shared Response Model ISC', fontsize=title_ftsize)
            plt.show()    

def view_surf(isc_r,isc_p,maskn='50',netn=7,thresh=.005,outpath=False,minmax=(None,None),symmetric_cmap=True):
    #this is to visualize isc_r and isc_p on brain surface
    #e.g. outpath = '/Users/xianl/Desktop/ntf_isc_int008_thr005.nii.gz' saves the whole brain isc_r (thresholded by isc_p) to a nifti file
    import pandas as pd
    from nltools.stats import threshold
    from nltools.mask import expand_mask, roi_to_brain
    from nilearn.plotting import view_img_on_surf
    mask,mask_x,mask_nodes = get_mask(maskn=maskn,netn=netn,plot_mask=False)
    isc_r_brain, isc_p_brain = roi_to_brain(pd.Series(isc_r), mask_x), roi_to_brain(pd.Series(isc_p), mask_x)
    wb_thr_isc = threshold(isc_r_brain, isc_p_brain, thr=thresh).to_nifti()
    if minmax[0]!=minmax[1]:
        symmetric_cmap = False
    surf = view_img_on_surf(wb_thr_isc,vmin=minmax[0],vmax=minmax[1],symmetric_cmap=symmetric_cmap)
    if outpath!=False:
        import nibabel as nib
        nib.save(wb_thr_isc,outpath)
    return surf

def view_brain(brain_img,surf_type='smooth',view_list=['lateral', 'medial', 'dorsal', 'ventral', 'anterior', 'posterior'],hemi_list=['left','right'],title=False,colorbar=True,outpath=False,thresh=.00,view=True,fig_format='.png',vmax=None):
    #assuming that img is in standard MNI space, this function takes in a nifti file or nii-like img for visualization
    import nibabel as nib
    from nilearn import datasets
    from nilearn import surface
    from nilearn import plotting
    #default thresh=.00 means no thresholding, and just showing the stats as
    #load in the brain image
    if type(brain_img)!=nib.nifti1.Nifti1Image:
        stat_img = nib.load(brain_img)
    else:
        stat_img = brain_img
    #load in the surface map
    if surf_type=='smooth':
        fsaverage = datasets.fetch_surf_fsaverage('fsaverage')
    elif surf_type=='rough':
        fsaverage = datasets.fetch_surf_fsaverage()
    else:
        print('surf flag can only be smooth or rough')
    #loop through the viewpoints and hemispheres
    for hemi in hemi_list:
        if hemi=='right':
            texture = surface.vol_to_surf(stat_img, fsaverage.pial_right)
            surf,surf_bg = fsaverage.infl_right,fsaverage.sulc_right
        elif hemi=='left':
            texture = surface.vol_to_surf(stat_img, fsaverage.pial_left)
            surf,surf_bg = fsaverage.infl_left,fsaverage.sulc_left
        for view in view_list:
            if title ==False or title==None:
                fname = '_'.join([hemi+'-hemi',view])
            else:
                fname = '_'.join([title,hemi+'-hemi',view])
            if outpath==False:
                out_target=None
            else:
                out_target = outpath+fname+fig_format
            #here's the plot of the figure
            if title==None:
                fig = plotting.plot_surf_stat_map(surf, texture, hemi=hemi, view = view, output_file=out_target,threshold=thresh, bg_map=surf_bg,vmax=vmax, colorbar=colorbar)
            else:
                fig = plotting.plot_surf_stat_map(surf, texture, hemi=hemi, view = view, title=fname, output_file=out_target,threshold=thresh, bg_map=surf_bg,vmax=vmax, colorbar=colorbar)
            if view==True:
                fig.show()
    return fig
    
def combine_figures_w_title(fpath, ftype='.png',outpath=False,show=False,columns='default', dpi=300,title=False,select=['lateral','medial']):
    '''combine multiple figures one'''
    #fpath is the path to all the figures to be combined
    from PIL import Image
    import math
    output_path=outpath
    if columns=='default':
        columns = len(select)
    inds = get_indices(grab_all(fpath,ftype),select,True)
    figure_paths = [fpath+i for i in apply_inds(inds,grab_all(fpath,ftype))]
    # Load images
    images = [Image.open(fp) for fp in figure_paths]
    # Determine the number of rows and columns
    rows = math.ceil(len(images) / columns)
    # Determine max width and height of the images
    max_width = max(image.size[0] for image in images)
    max_height = max(image.size[1] for image in images)
    # Create a new image with a white background
    combined_image = Image.new('RGB', (max_width * columns, max_height * rows), (255, 255, 255))
    # Paste images into the combined image
    for index, image in enumerate(images):
        row = index // columns
        col = index % columns
        combined_image.paste(image, (max_width * col, max_height * row))
    # Use matplotlib to add a title and save with specified DPI
    fig, ax = plt.subplots(figsize=(max_width * columns / dpi, max_height * rows / dpi), dpi=dpi)
    ax.imshow(combined_image)
    if title!=False:
        ax.set_title(title, fontsize=20)
    ax.axis('off')  # Hide axes
    # Save the figure
    plt.savefig(output_path, dpi=dpi, bbox_inches='tight')
    if show==False:
        plt.show()
    else:
        plt.close()
    print('complete.')

def crop_image2(img, crop_tuple):
    """Crop the image based on the crop_tuple (left, top, right, bottom)."""
    left, top, right, bottom = crop_tuple
    width, height = img.size
    new_left = left
    new_upper = top
    new_right = width - right
    new_lower = height - bottom
    return img.crop((new_left, new_upper, new_right, new_lower))

def combine_figures_w_title2(fpath, ftype='.png', outpath=False, show=False, columns='default', dpi=300, title=False, select=['lateral', 'medial'], crop_tuple=(0, 0, 0, 0)):
    '''combine multiple figures with optional cropping and title'''
    #example use:
        # root = '../results/3_global_isc/cutint_8tr/'
        # ls =xf.grab_all(root,'ntf*')
        # outp = root+'sum/'
        # xf.mkdir(outp)
        # for a in ls:
        #     fpath = root+a+'/n200_net7_shift_0tr/'
        #     xf.combine_figures_w_title2(fpath, columns=4, ftype='.png', outpath=outp+a+'.png', dpi=300, title=a, select=['lateral', 'medial'], crop_tuple=(100, 50, 200, 100)) #(left, top, right, bottom)
    from PIL import Image
    import matplotlib.pyplot as plt
    import math
    import os
    if not outpath:
        outpath = 'combined_figure.png'  # Default output path if not specified
    inds = get_indices(grab_all(fpath, ftype), select,find=True)    
    figure_paths = [os.path.join(fpath, i) for i in apply_inds(inds, grab_all(fpath, ftype))]
    print(len(figure_paths))
    # Load and optionally crop images
    images = [crop_image2(Image.open(fp), crop_tuple) for fp in figure_paths if os.path.exists(fp)]
    if not images:
        print("No images found matching the criteria. Exiting.")
        return
    if columns == 'default' or isinstance(columns, str):
        columns = len(select)
    rows = math.ceil(len(images) / columns)
    # Avoid calling max() on an empty list
    if images:
        max_width = max(image.size[0] for image in images)
        max_height = max(image.size[1] for image in images)
    else:
        print("Error: No images were loaded. Please check the file paths and formats.")
        return
    combined_image = Image.new('RGB', (max_width * columns, max_height * rows), (255, 255, 255))
    for index, image in enumerate(images):
        row = index // columns
        col = index % columns
        combined_image.paste(image, (max_width * col, max_height * row))
    # Using matplotlib for adding title and saving the image
    plt.figure(figsize=(max_width * columns / dpi, max_height * rows / dpi), dpi=dpi)
    plt.imshow(combined_image)
    plt.axis('off')  # Hide axes
    if title:
        plt.title(title, fontsize=20)
    plt.savefig(outpath, dpi=dpi, bbox_inches='tight')
    if not show:
        plt.close()
    else:
        plt.show()
    print('Complete.')

def combine_figures(fpath, ftype='.png', outpath='combined_image.png', columns=2, select=['lateral', 'medial']):
    '''Combine multiple figures into one.'''
    from PIL import Image
    import os
    # Get paths to the figures to be combined
    figure_paths = [os.path.join(fpath, fname) for fname in os.listdir(fpath) if fname.endswith(ftype) and any(s in fname for s in select)]
    # Load images
    images = [Image.open(fp) for fp in figure_paths]
    # Determine the number of rows and columns
    rows = -(-len(images) // columns)  # Ceiling division
    # Determine max width and height of the images
    max_width = max(image.size[0] for image in images)
    max_height = max(image.size[1] for image in images)
    # Create a new image with a white background
    combined_image = Image.new('RGB', (max_width * columns, max_height * rows), 'white')
    # Paste images into the combined image
    for index, image in enumerate(images):
        row = index // columns
        col = index % columns
        combined_image.paste(image, (max_width * col, max_height * row))
    # Save the combined image directly using PIL, preserving the original resolution
    combined_image.save(outpath)
    print('Complete.')

def crop_image(img, left, right, top, bottom):
    """Crop the image by removing 'left' pixels from the left, 'right' pixels from the right,
       'top' pixels from the top, and 'bottom' pixels from the bottom."""
    width, height = img.size
    # New dimensions
    crop_left = left
    crop_top = top
    crop_right = width - right
    crop_bottom = height - bottom
    return img.crop((crop_left, crop_top, crop_right, crop_bottom))

def combine_and_crop_images(fpath, crop_sides, ftype='.png', outpath='combined_cropped_image.png', columns=2, select=['lateral', 'medial'], title_space=100):
    """Combine and crop images with specified cropping from each side and a title."""
    # Example usage
        # combine_and_crop_images(fpath, (100, 200, 50, 10),columns=4,outpath=outp+a+'.png')
    from PIL import Image, ImageDraw, ImageFont
    import os
    figure_paths = [os.path.join(fpath, fname) for fname in os.listdir(fpath) if fname.endswith(ftype) and any(s in fname for s in select)]
    # Crop sides should be a tuple or list of four integers: (left, right, top, bottom)
    left, right, top, bottom = crop_sides    
    # Load and crop images
    images = [crop_image(Image.open(fp), left, right, top, bottom) for fp in figure_paths]
    # Calculate rows and new dimensions based on cropped images
    rows = -(-len(images) // columns)  # Ceiling division to determine number of rows
    if images:  # Check if images list is not empty to avoid errors
        max_width = max(image.size[0] for image in images)
        max_height = max(image.size[1] for image in images) + (title_space // len(images))
    else:
        max_width, max_height = 0, 0  # Default values in case there are no images
    combined_image = Image.new('RGB', (max_width * columns, (max_height * rows) + title_space), 'white')
    for index, image in enumerate(images):
        row = index // columns
        col = index % columns
        x_offset = max_width * col
        y_offset = max_height * row + title_space
        combined_image.paste(image, (x_offset, y_offset))
    combined_image.save(outpath)
    print(f'Combined and cropped image saved to {outpath}.')

def srm_align(all_data,n_features=100,component_ind=1):
    from nltools.stats import align
    import warnings
    warnings.simplefilter('ignore')
    all_data_z = [x.standardize(method='zscore') for x in all_data]
    srm = align(all_data, method='deterministic_srm', n_features=n_features)
    plot_srm_isc(srm=srm,component_ind=component_ind)
    return srm


def hyper_align(all_data,voxel_ind=0,tr_ind=0,nsubj=5):
    from nltools.stats import align
    import warnings
    warnings.simplefilter('ignore')
    hyperalign = align(all_data, method='procrustes')    
    plot_hyper_isc(all_data,hyperalign=hyperalign,voxel_ind=voxel_ind,tr_ind=tr_ind,nsubj=nsubj)
    return hyperalign


def funcalign_new_data(new_data,model,tr_ind=0,mthd='procrustes'):
    #example use:
        #funcalign_new_data(new_data,model)
        #new_data = all_data[-1]
        #model = align(all_data[:15], method='procrustes')
        #or, model_dict = srm = align(all_data_z, method='deterministic_srm', n_features=100), and mthd='srm'
    import matplotlib.pyplot as plt
    from nltools.stats import align
    import warnings
    warnings.simplefilter('ignore')
    #use the z-score of the new data if method is SRM, because srm is trained on z-scores
    if mthd =='srm':
        mthd= 'deterministic_srm'
        new_data= new_data.standardize(method='zscore')
    print(str('Plot the unaligned spatial topography at TR Index-'+str(tr_ind)))
    new_data[tr_ind].plot()
    aligned_new_sub = new_data.align(model['common_model'], method=mthd)
    print(str('Plot the aligned spatial topography at TR Index-'+str(tr_ind)))
    if mthd=='deterministic_srm':
        aligned_new_sub['transformation_matrix'][tr_ind].plot()
    else:
        aligned_new_sub['transformed'][tr_ind].plot()
    return aligned_new_sub


def funcalign_cross_validate(model,test_data,mthd='procrustes',voxel_ind=0,component_ind = 0,tr_ind=0):
    #example_use:
        #dp1_srm_trans=funcalign_cross_validate(srm,data_part1[:15],mthd='srm')
        #srm = srm_align(all_data[:15], n_features=100) where all_data= hdf5_apply_mask(base_dir,roi_ind=4,name_in_file='*crop*Part2*hdf5') 
        #data_part1 = hdf5_apply_mask(base_dir,roi_ind=4,name_in_file='*crop*Part1*hdf5') 
    import numpy as np
    import warnings
    warnings.simplefilter('ignore')
    if type(test_data)!=list:
        x = test_data
        new_x = test_data.copy()
        if mthd=='srm':
            new_x.data = np.dot(x.standardize(method='zscore').data, model['transformation_matrix'].data.T)
        else:
            new_x.data = np.dot(x.data, model['transformation_matrix'].data.T)
        transformed = new_x
    else:
        transformed = []
        for i,x in enumerate(test_data):
            new_x = x.copy()
            if mthd=='srm':
                new_x.data = np.dot(x.standardize(method='zscore').data, model['transformation_matrix'][i].data.T)
            else:
                new_x.data = np.dot(x.data, model['transformation_matrix'][i].data.T)
            transformed.append(new_x)
    return transformed


def hdf5_extract_voi_fsl(fsl_dir,roi_ind='all',subj='*',name_in_file='*hdf5',maskn='50',netn=17):
    #example use:
        #hdf5_extract_voi(base_dir,name_in_file)
        #base_dir = data_dir+'/fmriprep/'
        #name_in_file = ['*crop*Part1*','*crop*Part2*']
    #fmriprep_dir: e.g. '/Users/xli239/Desktop/naturalistic_fMRI_processing/3-datasets/Sherlock/fmriprep/'
    #roi_ind: default to write all roi's timeseries (i.e. VOI) to csv; otherwise, specify ROI index to be written
    #subj: default to loop through all existing subjects 'sub-01','sub-02',...
    #name_in_file: has to be either a string with #hdf5 or a list of such strings (to combine these files)
    #maskn: 
        #'50' for k50_2mm (parcellation based on patterns of coactivation from the Neurosynth database (de la Vega et al., 2016)), 
        #'268' for n268_2mm, 
        #int: 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 for Schaefer2018,
    #netn: number of net, either 7 or 17 for schaefer2018 atlas 
    import os
    import glob
    import pandas as pd
    from nltools.mask import create_sphere, expand_mask
    from nltools.data import Brain_Data
    import warnings
    #get the brain mask segment/parcel for ROI:
    mask, mask_x, mask_nodes = get_mask(maskn=maskn,netn=netn)
    #get the list of hdf5 files and write the VOIs to csv file:
    if type(name_in_file)==list:
        for scan in name_in_file:
            file_list = glob.glob(os.path.join(fsl_dir, subj, f'{scan}*hdf5'))
            file_list.sort()
            scan_name = scan.replace('*','')
            for f in file_list:
                sub = os.path.basename(f).split('_')[0]
                print(sub)
                data = Brain_Data(f)
                if roi_ind !='all':
                    roi_mask = mask_x[roi_ind]
                    roi = data.extract_roi(roi_mask)
                else:
                    roi = data.extract_roi(mask)
                pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{scan_name}_Average_ROI_n{mask_nodes}.csv" ), index=False)        
    elif name_in_file != '*hdf5':
        file_list = glob.glob(os.path.join(fsl_dir, subj, f'{name_in_file}*hdf5'))
        a = os.path.join(fsl_dir, subj, f'{name_in_file}*hdf5')
        print(a)
        file_list.sort()
        scan_name = name_in_file.replace('*','')
        for f in file_list:
            sub = os.path.basename(f).split('_')[0]
            print(sub)
            data = Brain_Data(f)
            if roi_ind !='all':
                roi_mask = mask_x[roi_ind]
                roi = data.extract_roi(roi_mask)
            else:
                roi = data.extract_roi(mask)
            roi = data.extract_roi(mask)
            pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{scan_name}_Average_ROI_n{mask_nodes}.csv" ), index=False)        
    else:
        file_list = glob.glob(os.path.join(fsl_dir, subj, f'{name_in_file}'))
        for f in file_list:
            file_name_parts = os.path.basename(f).split('_')
            file_list.sort()
            sub = file_name_parts[0]
            for part in file_name_parts:
                if parts.find('task')!=-1:
                    task = part
                else:
                    task = 'task'
            print(sub)
            data = Brain_Data(f)
            if roi_ind !='all':
                roi_mask = mask_x[roi_ind]
                roi = data.extract_roi(roi_mask)
            else:
                roi = data.extract_roi(mask)
            roi = data.extract_roi(mask)
            pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{task}_Average_ROI_{mask_name}.csv" ), index=False)        
    print('The VOIs have been written into csv file')



def hdf5_extract_voi(fmriprep_dir,roi_ind='all',subj='*',name_in_file='*hdf5',maskn='50'):
    #example use:
        #hdf5_extract_voi(base_dir,name_in_file)
        #base_dir = data_dir+'/fmriprep/'
        #name_in_file = ['*crop*Part1*','*crop*Part2*']
    #fmriprep_dir: e.g. '/Users/xli239/Desktop/naturalistic_fMRI_processing/3-datasets/Sherlock/fmriprep/'
    #roi_ind: default to write all roi's timeseries (i.e. VOI) to csv; otherwise, specify ROI index to be written
    #subj: default to loop through all existing subjects 'sub-01','sub-02',...
    #name_in_file: has to be either a string with #hdf5 or a list of such strings (to combine these files)
    #maskn: 
        #'50' for k50_2mm (parcellation based on patterns of coactivation from the Neurosynth database (de la Vega et al., 2016)), 
        #'268' for n268_2mm, 
        #int: 100, 200, 300, 400 (default), 500, 600, 700, 800, 900, 1000 for Schaefer2018, 
    import os
    import glob
    import pandas as pd
    from nltools.mask import create_sphere, expand_mask
    from nltools.data import Brain_Data
    import warnings
    #get the brain mask segment/parcel for ROI:
    mask, mask_x, mask_nodes = get_mask(maskn=maskn)
    #get the list of hdf5 files and write the VOIs to csv file:
    if type(name_in_file)==list:
        for scan in name_in_file:
            file_list = glob.glob(os.path.join(fmriprep_dir, subj, 'func', f'{scan}*hdf5'))
            file_list.sort()
            scan_name = scan.replace('*','')
            for f in file_list:
                sub = os.path.basename(f).split('_')[0]
                print(sub)
                data = Brain_Data(f)
                if roi_ind !='all':
                    roi_mask = mask_x[roi_ind]
                    roi = data.extract_roi(roi_mask)
                else:
                    roi = data.extract_roi(mask)
                pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{scan_name}_Average_ROI_n{mask_nodes}.csv" ), index=False)        
    elif name_in_file != '*hdf5':
        file_list = glob.glob(os.path.join(fmriprep_dir, subj, 'func', f'{name_in_file}*hdf5'))
        file_list.sort()
        scan_name = name_in_file.replace('*','')
        for f in file_list:
            sub = os.path.basename(f).split('_')[0]
            print(sub)
            data = Brain_Data(f)
            if roi_ind !='all':
                roi_mask = mask_x[roi_ind]
                roi = data.extract_roi(roi_mask)
            else:
                roi = data.extract_roi(mask)
            roi = data.extract_roi(mask)
            pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{scan_name}_Average_ROI_n{mask_nodes}.csv" ), index=False)        
    else:
        file_list = glob.glob(os.path.join(fmriprep_dir, subj, 'func', f'{name_in_file}'))
        for f in file_list:
            file_name_parts = os.path.basename(f).split('_')
            file_list.sort()
            sub = file_name_parts[0]
            for part in file_name_parts:
                if parts.find('task')!=-1:
                    task = part
                else:
                    task = 'task'
            print(sub)
            data = Brain_Data(f)
            if roi_ind !='all':
                roi_mask = mask_x[roi_ind]
                roi = data.extract_roi(roi_mask)
            else:
                roi = data.extract_roi(mask)
            roi = data.extract_roi(mask)
            pd.DataFrame(roi.T).to_csv(os.path.join(os.path.dirname(f), f"{sub}_{task}_Average_ROI_{mask_name}.csv" ), index=False)        
    print('The VOIs have been written into csv file')


def combine_voi_csv(fmriprep_dir,comb_root,comb,subj='*',stand='none',savefig='none'):
    #example use: dict_sub1 = combine_voi_csv(base_dir,comb_root,comb,subj='sub-01')
        # comb_root = 'Part*_Average_ROI_n50'
        # comb = ['Part1_Average_ROI_n50','Part2_Average_ROI_n50'] 
        # base_dir = data_dir+'/fmriprep/'
        # subj: can specify individual subject e.g. 'sub-01', or default as all subj
        # stand: standardize/z-score each timeseries before/after concatenation
        # save: save the 2d matrix per subject to csv if input an output path, e.g. '/Users/xli239/Desktop/' or '/Users/xli239/Desktop/n200_' 
    import os
    import glob   
    import pandas as pd
    sub_list = [os.path.basename(x).split('_')[0] for x in glob.glob(os.path.join(fmriprep_dir, '*', 'func', f'*{comb[0]}*csv'))]
    sub_list.sort()
    sub_timeseries = {}
    for sub in sub_list:
        print(sub)
        part1 = pd.read_csv(os.path.join(fmriprep_dir, sub, 'func', f'{sub}_{comb[0]}.csv'))
        for ind in range(len(comb)-1):
            ind+=1
            current_part = pd.read_csv(os.path.join(fmriprep_dir, sub, 'func', f'{sub}_{comb[ind]}.csv'))
            if stand !='none':
                from sklearn.preprocessing import StandardScaler
                part1 = pd.DataFrame(StandardScaler().fit_transform(part1))
                current_part = pd.DataFrame(StandardScaler().fit_transform(current_part))
            sub_data = part1.append(current_part)
        sub_data.reset_index(inplace=True, drop=True)   
        sub_timeseries[sub] = sub_data 
    if savefig!='none':
        for sub in sub_timeseries:
            mat = sub_timeseries[sub]
            save_as = str(savefig+sub+'_voi.csv')
            pd.DataFrame(mat).to_csv(save_as)
    return sub_timeseries


def get_subject_roi(data, roi_ind):
    #example use: sub_rois = get_subject_roi(sub_timeseries, roi_ind)
        #roi_ind = 32
        #mask_x[roi_ind].plot() (this is to check out the roi at that index)  
    import pandas as pd
    sub_rois = {}
    for sub in data:
        sub_rois[sub] = data[sub].iloc[:, roi_ind]
    pd.DataFrame(sub_rois).head()
    return pd.DataFrame(sub_rois)


def plot_circular_shift(sub_rois, sub, sampling_freq):
    #this is to show what this subject's timeseries at this ROI look like versus it being randomly circular shifted 
    import matplotlib.pyplot as plt
    import numpy as np
    from numpy.fft import fft, ifft, fftfreq
    from nltools.stats import circle_shift
    #original timeseries
    f,a = plt.subplots(nrows=2, ncols=2, figsize=(15, 5))
    a[0,0].plot(sub_rois[sub], linewidth=2)
    a[0,0].set_ylabel('Avg Activity', fontsize=16)
    a[0,1].set_xlabel('Time (TR)', fontsize=18)
    a[0,0].set_title('Observed Data', fontsize=16)
    #fft: fast fourier transform the timeseries of this specified subject: original data power spectrum
    fft_data = fft(sub_rois[sub]) 
    freq = fftfreq(len(fft_data), 1/sampling_freq) #returns the fft sample frequencies: fftfreq(window length, time step)
    n_freq = int(np.floor(len(fft_data)/2)) # show only the positive range of the above freq range 
    a[0,1].plot(freq[:n_freq], np.abs(fft_data)[:n_freq], linewidth=2) 
    a[0,1].set_xlabel('Frequency (Hz)', fontsize=18)
    a[0,1].set_ylabel('Amplitude', fontsize=18)
    a[0,1].set_title('Power Spectrum', fontsize=18)
    #circular shift by random amount: circular shifted data
    circle_shift_data = circle_shift(sub_rois[sub]) 
    a[1,0].plot(circle_shift_data, linewidth=2, color='red')
    a[1,0].set_ylabel('Avg Activity', fontsize=16)
    a[1,0].set_xlabel('Time (TR)', fontsize=16)
    a[1,0].set_title('Circle Shifted Data', fontsize=16)
    #fft: circular shifted data power spectrum
    fft_circle = fft(circle_shift_data)
    a[1,1].plot(freq[:n_freq], np.abs(fft_circle)[:n_freq], linewidth=2, color='red')
    a[1,1].set_xlabel('Frequency (Hz)', fontsize=18)
    a[1,1].set_ylabel('Amplitude', fontsize=18)
    a[1,1].set_title('Circle Shifted Power Spectrum', fontsize=18)
    #plot out the figure
    plt.tight_layout()
    print('Plot shown')


def plot_phase_randomize(sub_rois, sub, sampling_freq):
    #this is to show what this subject's timeseries at this ROI look like versus it being randomly circular shifted 
    import matplotlib.pyplot as plt
    import numpy as np
    from numpy.fft import fft, ifft, fftfreq
    from nltools.stats import phase_randomize
    #original timeseries
    f,a = plt.subplots(nrows=2, ncols=2, figsize=(15, 5))
    a[0,0].plot(sub_rois[sub], linewidth=2)
    a[0,0].set_ylabel('Avg Activity', fontsize=16)
    a[0,1].set_xlabel('Time (TR)', fontsize=18)
    a[0,0].set_title('Observed Data', fontsize=16)
    #fft: fast fourier transform the timeseries of this specified subject: original data power spectrum
    fft_data = fft(sub_rois[sub]) 
    freq = fftfreq(len(fft_data), 1/sampling_freq) #returns the fft sample frequencies: fftfreq(window length, time step)
    n_freq = int(np.floor(len(fft_data)/2)) # show only the positive range of the above freq range 
    a[0,1].plot(freq[:n_freq], np.abs(fft_data)[:n_freq], linewidth=2) 
    a[0,1].set_xlabel('Frequency (Hz)', fontsize=18)
    a[0,1].set_ylabel('Amplitude', fontsize=18)
    a[0,1].set_title('Power Spectrum', fontsize=18)
    #circular shift by random amount: circular shifted data
    phase_random_data = phase_randomize(sub_rois[sub])
    a[1,0].plot(phase_random_data, linewidth=2, color='red')
    a[1,0].set_ylabel('Avg Activity', fontsize=16)
    a[1,0].set_xlabel('Time (TR)', fontsize=16)
    a[1,0].set_title('Phase Shifted Data', fontsize=16)
    #fft: circular shifted data power spectrum
    fft_phase = fft(phase_random_data)
    a[1,1].plot(freq[:n_freq], np.abs(fft_phase)[:n_freq], linewidth=2, color='red')
    a[1,1].set_xlabel('Frequency (Hz)', fontsize=18)
    a[1,1].set_ylabel('Amplitude', fontsize=18)
    a[1,1].set_title('Phase Shifted Power Spectrum', fontsize=18)
    #plot out the figure
    plt.tight_layout()
    print('Plot shown')


def bootstrap_subject_matrix(sim_mat=False,sub_roi_dict=False, input = 'sim_mat',random_state=None):
    '''This function shuffles subjects within a similarity matrix based on recommendation by Chen et al., 2016'''   
    import numpy as np
    import pandas as pd
    from sklearn.metrics import pairwise_distances
    from sklearn.utils import check_random_state
    if input=='sim_mat':
        similarity_matrix = sim_mat
    elif input =='sub_roi_dict':
        similarity_matrix = 1 - pairwise_distances(pd.DataFrame(sub_roi_dict).T, metric='correlation')
    else:
        print('Needs either similarity matrix or the ROI timeseries across all subs (as dictionary)')
    plot_matrix(similarity_matrix)
    random_state = check_random_state(random_state)
    n_sub = similarity_matrix.shape[0]
    bootstrap_subject = sorted(random_state.choice(np.arange(n_sub), size=n_sub, replace=True)) #this is to pick 1 subj out of the 16 subj for 16 times with replacement
    print(bootstrap_subject)
    bootstrap_matrix =  similarity_matrix[bootstrap_subject, :][:, bootstrap_subject] #compute the similarity matrix of the resulting 16 subjects (some subjects could have been repeatedly sampled) 
    plot_matrix(bootstrap_matrix)
    return bootstrap_matrix


def plot_isc_null(stats, method, add_title='none',savefig=False,outpath=False,eps=False):
    #This is to plot the null distribution formed from the permutation test on a specific ROI
    #add_title: added to the end of the plot title, e.g. 'n50_roi8'
    #save: needs to be an output path: '/Users/xli239/Desktop/figures/'
    import matplotlib.pyplot as plt
    import seaborn as sns
    plt.figure(figsize=(12,6))
    if method=='circle'or method=='circle_shift':
        #null dist, observed ISC, legend 
        sns.distplot(stats['null_distribution'], kde=True, label='Bootstrap')
        plt.axvline(stats['isc'], linestyle='-', color='red', linewidth=4)
        plt.legend(['Observed ISC', 'Circle Shift Null'], fontsize=18)
        #vertical line marking 95% CI 
        plt.axvline(stats['ci'][0], linestyle='--', color='blue')
        plt.axvline(stats['ci'][1], linestyle='--', color='blue')
    elif method=='phase' or method=='phase_randomize':
        #null dist, observed ISC, legend 
        sns.distplot(stats['null_distribution'], kde=True, label='Bootstrap')
        plt.axvline(stats['isc'], linestyle='-', color='red', linewidth=4)
        plt.legend(['Observed ISC', 'Phase Randomize Null'], fontsize=18)
        #vertical line marking 95% CI 
        plt.axvline(stats['ci'][0], linestyle='--', color='blue')
        plt.axvline(stats['ci'][1], linestyle='--', color='blue')
    elif method=='bs' or method== 'bootstrap':
        #null dist, observed ISC, legend 
        sns.distplot(stats['null_distribution'] - stats['isc'], kde=True, label='Bootstrap')
        plt.axvline(stats['isc'], linestyle='-', color='red', linewidth=4)
        plt.legend(['Observed ISC', 'Bootstrap Null'], fontsize=18)
        #vertical line marking 95% CI 
        plt.axvline(stats['ci'][0] - stats['isc'], linestyle='--', color='blue')
        plt.axvline(stats['ci'][1] - stats['isc'], linestyle='--', color='blue')
    else:
        print('Error: input method not found')
    #axis labels and title
    plt.ylabel('Frequency', fontsize=18)
    plt.xlabel('ISC Values (correlation)', fontsize=18)
    if add_title !='none':
        plt.title(str('ISC Null Distribution for: '+add_title), fontsize=20)
    else:
        plt.title('ISC Null Distribution', fontsize=20)
    mesg = 'Figure plotted'
    if savefig!=False:
        if outpath!=False:
            plt.savefig(outpath+str(add_title)+'_isc_null.png')
            if eps!=False:
                plt.savefig(outpath+str(add_title)+'_isc_null.eps')
        else:
            plt.savefig(str(add_title)+'_isc_null.png')
            if eps!=False:
                plt.savefig(str(add_title)+'_isc_null.eps')
        mesg = mesg + ' and saved at specified location'
    print(mesg)

def get_avg_others(ndarray,mute=True,skip_nan=True):
    #this is to prep for isc analysis, where to compute the average other subjects' ndarray in advance
    #ndarray's dim1 must be subjects, e.g. ndarray[0] is sub1's data
    #skip_nan: if any subjects' specific voxel is nan, skip it when computing the average mvp
    avg_others = []
    for i in range(len(ndarray)):
        avg_oth = avg([ndarray[j] for j in range(len(ndarray)) if j!=i],mute=mute, ignore_nan=skip_nan)
        avg_others.append(avg_oth)
    return avg_others

def isc_avg(ndarray,fill='nan',mute=True,method='pearson'):
    rval,pval=[],[]
    for i in range(len(ndarray)):
        ls1 = ndarray[i]
        ls2 = avg([ndarray[j] for j in range(len(ndarray)) if j!=i],mute=mute)
        if method=='pearson':
            r,p = get_corr(ls1,ls2)
        elif method=='cosine':
            r = cosine_sim(ls1,ls2)
            p = 'na'
        rval.append(r)
        pval.append(p)
    return rval,pval

def isc_test(sub_timeseries,roi_ind='all',maskn='50',netn=7,rand_method='circle_shift',plot_mask=False,plot_whole_brain=False,thresh=.001, plot_null=False, outpath=False,permute = 5000, meassure='median', return_bt=True,eps=False):
    #Note: this permutation test is for pairwise ISC only.
    #roi_ind: can also be an integer for roi index
    #plot_null: needs to be an output path, e.g. '/Users/xli239/Desktop/' 
    import pandas as pd
    from nltools.stats import isc, threshold
    from nltools.data import Brain_Data
    from nltools.mask import expand_mask, roi_to_brain
    from nilearn.plotting import view_img_on_surf, view_img
    #grab whole-brain mask for ROIs:
    if plot_mask !=False:
        mask,mask_x,mask_nodes = get_mask(maskn=maskn,netn=netn)
    else:
        mask,mask_x,mask_nodes = get_mask(maskn=maskn,netn=netn,plot_mask=False)
    #test ISC: whole-brain
    if roi_ind== 'all': 
        isc_r, isc_p = {}, {}
        for roi in range(int(maskn)):
            sub_roi = get_subject_roi(sub_timeseries, roi)
            stats = isc(sub_roi, method=rand_method, n_bootstraps=permute, metric=meassure,return_bootstraps=return_bt)
            if plot_null!=False:
                add = 'n'+str(int(maskn))+'_roi'+str(int(roi+1))
                plot_isc_null(stats, rand_method, add_title=add,savefig=plot_null,outpath=outpath,eps=eps)
            isc_r[roi], isc_p[roi] = stats['isc'], stats['p']
        if plot_whole_brain!=False:
            mask_x = expand_mask(mask)
            isc_r_brain, isc_p_brain = roi_to_brain(pd.Series(isc_r), mask_x), roi_to_brain(pd.Series(isc_p), mask_x)
            isc_r_brain.plot(cmap='RdBu_r')
            view3 = view_img(isc_r_brain.to_nifti())
            surf = view_img_on_surf(threshold(isc_r_brain, isc_p_brain, thr=thresh).to_nifti())
            return isc_r,isc_p,view3,surf
        else:
            return isc_r,isc_p
    #test isc at an roi index 
    elif type(roi_ind)==int: 
        sub_roi = get_subject_roi(sub_timeseries, roi_ind)
        stats = isc(sub_roi, method=rand_method, n_bootstraps=permute, metric=meassure, return_bootstraps=return_bt)
        print(f"ISC: {stats['isc']:.06}, p = {stats['p']:.06}") #'.06': how many digits we want to keep/show after the dot 
        if plot_null!=False:
            add = 'roi'+str(int(roi_ind+1))
            plot_isc_null(stats, rand_method, add_title=add,savefig=plot_null,eps=eps)
        return stats
    else:
        print('Error: roi_ind input not recognized')
            
def ttc(mvp_trs1,mvp_trs2='None'):
    '''time-time-correlation (for a single subject): 
        computes the pairwise mvp correlation across time (trs)'''
    #mvp_trs1 should take the shape of (n_trs, n_roi_size)
    return get_pairwise_pearsonr(mvp_trs1,mvp_trs2)

def time_time_mat(sim_mat,time_window,title,outpath=False,save_matrix=False,map_color='CMRmap',diagmask=True):
    plot_figsize(10,10)
    on_p = (time_window*2+1)/2
    plot_vline(on_p,[on_p,on_p*2],add=True,color='black',lwd=2)
    plot_hline(on_p,[0,on_p],add=True,color='black',lwd=2)
    plot_label([on_p-2],[on_p-3],labels=['interrupt'],add=True,font_size=20)
    plot_label([on_p-1],[on_p-1],labels=['onset'],add=True,font_size=20)
    plot_matrix(sim_mat,labname=[str(i)+'TR' for i in range(-time_window,time_window+1)],map_color=map_color,diagmask=diagmask,title=title,outpath=outpath,save_matrix=save_matrix)

def plot_network_degree(mat,savefig=False):
    '''Plot the degree of the thresholded Adjaceny matrix'''
    import networkx as nx
    import matplotlib.pyplot as plt
    from nltools.data import Adjacency
    if not isinstance(mat, Adjacency):
        mat = Adjacency(mat, matrix_type='similarity') #default to consider input matrix as similarity matrix
    plt.figure(figsize=(20,15))
    G = mat.to_graph()
    pos = nx.kamada_kawai_layout(G)
    node_and_degree = G.degree()
    nx.draw_networkx_edges(G, pos, width=3, alpha=.4)
    nx.draw_networkx_labels(G, pos, font_size=14, font_color='darkslategray')

    nx.draw_networkx_nodes(G, pos, nodelist=list(dict(node_and_degree).keys()),
                           node_size=[x[1]*100 for x in node_and_degree],
                           node_color=list(dict(node_and_degree).values()),
                           cmap=plt.cm.Reds_r, linewidths=2, edgecolors='darkslategray', alpha=1, with_labels=True)
    if savefig!=False:
        plt.savefig(savefig+'network_graph.png')
    print('Completed')


def get_isfc(sub_timeseries,outpath=False,plot_allsub=False,savefig=False,map_color='RdBu_r'):
    #get isfc matrix for each subject, computes group average, and plots the matrix
    #returns 1) the average matrix across all subj, 2) each subj's matrix
    import numpy as np
    import seaborn as sns
    import matplotlib.pyplot as plt
    from nltools.stats import isfc
    data = list(sub_timeseries.values())
    isfc_allsub = isfc(data)
    isfc_avg = np.array(isfc_allsub).mean(axis=0)
    if outpath!=False:
        mkdir(outpath)
        outpath_file = outpath+'isfc_'+str(len(sub_timeseries))+'sub_avg'
        plot_matrix(isfc_avg,title='ISFC Group Average',outpath = outpath_file,map_color=map_color)    
    else:
        plot_matrix(isfc_avg,title='ISFC Group Average',map_color=map_color)
    if plot_allsub!=False:
        for sub in range(len(isfc_allsub)):
            isfc_sub = isfc_allsub[sub]
            outname = 'isfc_'+ list(sub_timeseries.keys())[sub]
            if savefig!=False:
                plot_matrix(isfc_sub,title=outname,outpath=save,map_color=map_color)
            else:
                plot_matrix(isfc_sub,title=outname,map_color=map_color)
    return isfc_allsub,isfc_avg


def fc_test(isfc_allsub,fdr_thr=.0000001,plot_mat=False, plot_net=False,savefig=False,plot_brain=False,maskn='50'):
    # threshold the connectivity matrix by running a one-sample ttest on each ISFC value 
    # and correcting for multiple comparisons using FDR
    #savefig would otherwise needs to be a path to save to e.g. '/Users/xli239/Desktop/Figures/'
    import numpy as np
    from scipy.stats import ttest_1samp
    from nltools.stats import fdr, threshold
    import seaborn as sns
    import matplotlib.pyplot as plt
    #1sample-ttest against 0 and thresh pos, fdr p
    t, p = ttest_1samp(np.array([x.reshape(-1) for x in isfc_allsub]), 0)
    tmap = np.reshape(t, isfc_allsub[0].shape)
    pmap = np.reshape(p, isfc_allsub[0].shape)
    thresh = fdr(p, fdr_thr)
    thresholded_t_pos = t.copy()
    thresholded_t_pos[p > thresh] = 0
    thresholded_t_pos[thresholded_t_pos <= 0] = 0
    thresholded_t_pos[thresholded_t_pos > 0] = 1
    thresholded_t_pos = np.reshape(thresholded_t_pos, isfc_allsub[0].shape)
    #go through the plots 
    if plot_mat !=False:
        num_nodes=np.shape(isfc_allsub[0])
        xlab = str(num_nodes[0])+'Nodes'
        ylab = str(num_nodes[1])+'Nodes'
        ttl = 'Positive Edges Thresholded at '+str(fdr_thr)+' with FDR correction'
        sns.heatmap(thresholded_t_pos, square=True, xticklabels=False, yticklabels=False)
        plt.title(ttl, fontsize=20)
        plt.xlabel(xlab, fontsize=18)
        plt.ylabel(ylab, fontsize=18)
        if savefig!=False:
            add = xlab+'_x_'+ylab+'fdar_thr'
            plt.savefig(str(savefig+add+'.png'))
    if plot_net!=False:
        plot_network_degree(thresholded_t_pos,savefig = savefig)
    if plot_brain!=False:
        #project the number of connections (i.e., degree) with each node back into brain space.
        import pandas as pd
        from nltools.data import Adjacency
        from nltools.mask import roi_to_brain
        from nilearn.plotting import view_img_on_surf
        mask,mask_x,mask_nodes = get_mask(maskn=maskn)
        degree = pd.Series( dict( Adjacency(thresholded_t_pos, matrix_type='similarity').to_graph().degree() ) )
        brain_degree = roi_to_brain(degree, mask_x)
        brain_degree.plot()
        surf = view_img_on_surf(brain_degree.to_nifti())
        return thresholded_t_pos,tmap,pmap,surf
        print('Completed')
    else:
        return thresholded_t_pos,tmap,pmap
        print('Completed')
    
def apply_phase_shift_and_correlate(target_series, series):
    #permutation of the timeseries need to keep the power spectrum, so use phase shift
    # FFT of the series: Fast Fourier Transform
    fft_series = np.fft.fft(series)
    # Random phase shift
    random_phases = np.exp(2j * np.pi * np.random.rand(len(series)))
    # Apply the phase shift
    fft_series_shifted = fft_series * random_phases
    # IFFT to get the shifted series back in time domain
    shifted_series = np.fft.ifft(fft_series_shifted)
    # Compute correlation
    correlation = pearsonr(target_series, shifted_series.real)[0]  # Use .real part as IFFT returns complex numbers
    return correlation

def single_subj_isc_at_roi(data, target_subject_key, roi_index, num_permutations):
    target_series = data[target_subject_key].iloc[:, roi_index].values
    non_target_keys = [key for key in data if key != target_subject_key]
    
    # Precompute FFT of target series (not needed here but shown for consistency)
    fft_target = np.fft.fft(target_series)

    actual_correlations = [pearsonr(target_series, data[key].iloc[:, roi_index].values)[0] for key in non_target_keys]
    actual_avg_r = np.mean(actual_correlations)

    null_distribution = np.zeros(num_permutations)
    for perm in range(num_permutations):
        permuted_correlations = [apply_phase_shift_and_correlate(target_series, data[key].iloc[:, roi_index].values) for key in non_target_keys]
        null_distribution[perm] = np.mean(permuted_correlations)

    p_value = np.sum(np.abs(null_distribution) >= np.abs(actual_avg_r)) / num_permutations
    return roi_index, actual_avg_r, p_value, null_distribution

def single_subj_isc(data, target_subject_key, roi_indices, num_permutations=100):
    # Example usage
        # subjects_data_dict = sub_timeseries.copy() 
        # target_subject_key = 'sub-023'
        # rois_to_analyze = [i for i in range(10)] 
        # allr,allp,rois,results = single_subj_isc(subjects_data_dict, target_subject_key, rois_to_analyze)

    from concurrent.futures import ProcessPoolExecutor
    results = {}                      
    with ProcessPoolExecutor() as executor:
        futures = [executor.submit(single_subj_isc_at_roi, data, target_subject_key, roi, num_permutations) for roi in roi_indices]
        for future in futures:
            roi_index, actual_avg_r, p_value, null_distribution = future.result()
            results[roi_index] = (actual_avg_r, p_value, null_distribution)
    rois =[i for i in results.keys()]
    allr,allp =[results[i][0] for i in rois],[results[i][1] for i in rois]
    return allr,allp,rois,results




# ################## read_table and read_rows substitute for python 3.9 ###################

# def read_table(infile, sheet=1):
#     # infile specifies the name of the file and its path
#     # example use:
#     # table = read_table('/Users/xianl/Desktop/data/sub1.xlsx') 
#     # this returns a table that you can process the rows for, e.g.:
#     # nrows = table.max_row
#     # for row in range(1, nrows + 1):
#     #    row_counter+=1
#     #    vec = table.row_values(row)
#     import openpyxl
#     workbook = openpyxl.load_workbook(infile, data_only=True)
#     worksheet = workbook.worksheets[sheet - 1]
#     class TableWrapper:
#         def __init__(self, worksheet):
#             self.worksheet = worksheet
#             self.nrows = worksheet.max_row
#         def row_values(self, row):
#             return [cell.value for cell in self.worksheet[row]]
    
#     return TableWrapper(worksheet)

# def read_rows(table,col_num,key=True,skip_row1=True,clear_empty=True,clear=''):
#     #table: read_table('./2_data/1_summary.xlsx')    
#     #col_num: either int or list of int for the column number (not index) wanted
#     #key: defaults key to be the title row; if no title row, need to input a list of strings for key name
#     #skip_row1: if no title row, needs to input the var key, and specify skip_row1=False
#     #example use: 
#         #all_subs = read_rows(table,8)
#         #all_subs = read_rows(table,[1,7,8])
#     row_count =0
#     nrows = table.nrows
#     all_rows = {}
#     #read row:
#     for row in range(1,nrows+1):
#         row_count+=1
#         v = table.row_values(row)
#         #skip first row (if key not specified: grab title)
#         if skip_row1==True and row_count == 1:
#             if key ==True:
#                 if type(col_num)==int:
#                     key = str(v[col_num-1])
#                 elif type(col_num)==list:
#                     key = []
#                     for i in col_num:
#                         key.append(str(v[i-1]))
#                 else:
#                     print('ERROR: col_num type needs to be either int or list of int')
#             continue
#         elif skip_row1==False and row_count==1:
#             key=['col_'+str(i) for i in col_num]
#         #update progress of row reading:
#         current_row = 'row'+str(row_count)
#         #print(current_row)
#         #check out and save the col value(s)
#         if type(col_num)==int:
#             out = v[col_num-1]
#             all_rows[current_row]={}
#             if key not in all_rows[current_row]:
#                 all_rows[current_row][key]=[]
#             all_rows[current_row][key].append(out)
#             #print(out) 
#         if type(col_num)==list:
#             all_rows[current_row]={}
#             key_count=0
#             for keyname in key:
#                 key_count+=1
#                 if keyname not in all_rows[current_row]:
#                     all_rows[current_row][keyname]=[]
#                 item = str(v[col_num[key_count-1]-1])
#                 all_rows[current_row][keyname].append(item)    
#             #print(all_rows[current_row])
#     data = data_by_key(all_rows,clear_empty=clear_empty,clear=clear)
#     return data, all_rows 



# import time
# from contextlib import contextmanager
# @contextmanager

# def timer():
#     # Example usage:
#         # with measure_time():
#         #     # Your code here
#         #     for i in range(1000000):
#         #         pass
#     """
#     Context manager to measure and print the runtime of a block of code.
#     """
#     start_time = time.time()
#     yield
#     end_time = time.time()
#     elapsed_time = end_time - start_time
#     print(f"Elapsed time: {elapsed_time:.2f} seconds")




# ##################old scripts no longer in use###################

# def sort_square_mtx(mtx, vct):
#     """
#     Sorts rows/columns of a matrix according to a separate vector.
#     """
#     inds = vct.argsort()
#     mtx_sorted = mtx.copy()
#     mtx_sorted = mtx_sorted[inds, :]
#     mtx_sorted = mtx_sorted[:, inds]
#     return mtx_sorted


# def scale_mtx(mtx):
#     """
#     Scales a matrix to have values between 0 and 1.
#     """
#     return (mtx-np.min(mtx))/(np.max(mtx)-np.min(mtx))


# def set_aspect_ratio(ax):
#     """
#     helper function to make square axes with equal x and y ranges
#     """
#     xmin, xmax = ax.get_xlim()
#     ymin, ymax = ax.get_ylim()
#     both_min = np.min([xmin, ymin])
#     both_max = np.max([xmax, ymax])
#     ax.set_xlim((both_min, both_max))
#     ax.set_ylim((both_min, both_max))
#     x0,x1 = ax.get_xlim()
#     y0,y1 = ax.get_ylim()
#     ax.set_aspect(abs(x1-x0)/abs(y1-y0))



            
# def isc_avg(sub_timeseries,perm=200,plot_whole_brain=True,outpath=False):
#     #this is to compute the isc of each subject with averaged others; then average across all subjects' isc
#     all_r = []
#     all_p = []
#     zdata = zscore_multi_dfs(sub_timeseries)
#     for key in zdata.keys():
#         #average other subjects' timeserires
#         others_ts = zdata.copy()    
#         del others_ts[key]
#         avg_others = avg_df(others_ts,z=False)
#         #construct the subj pair
#         sub_pair = {}
#         sub_pair[key]=zdata[key]
#         sub_pair['avg_others']=avg_others
#         #compute individual isc stats
#         isc_r,isc_p = isc_test(sub_pair,rand_method='bootstrap',maskn=200, meassure='mean', permute=perm, thresh=.005)
#         all_r.append([isc_r[key] for key in isc_r])
#         all_p.append([isc_p[key] for key in isc_p])
#         #save individual isc stats
#         if outpath!=False:
#             subpair_data = {'isc_r':isc_r, 'isc_p':isc_p}
#             write_dict_anylen(subpair_data,outpath=outpath+str(key) +'_isc_perm'+str(perm)+'_stats.xlsx')
#     #save overall isc
#     avg_r = avg(all_r)
#     avg_p = avg(all_p)
#     if outpath!=False:
#         subpair_data = {'isc_r':isc_r, 'isc_p':isc_p}
#         write_dict_anylen(subpair_data,outpath=outpath+'nsub'+str(len(sub_timeseries.keys()))+'_isc_perm'+str(perm)+'_stats.xlsx')
#     #plot overall isc
#     if plot_whole_brain !=False:
#         surf = view_surf(avg_r,avg_p,thresh=0.01,maskn=200)
#         return avg_r, avg_p, surf
#     else:
#         return avg_r,avg_p



# def get_indices2(search_list,searched_item):
#     #if search_list = [1,2,3,5,1,81] and searcher_item = [1,81], this returns ind_list =[0,4,5]
#     ind_list=[]
#     dt = type(searched_item)
#     if dt==str or dt==int or dt==float:
#         item = searched_item
#         for i in range(len(search_list)):
#             if search_list[i]==item:
#                 if i not in ind_list:
#                     ind_list.append(i)
#     else:
#         for item in searched_item:
#             for i in range(len(search_list)):
#                 if search_list[i]==item:
#                     if i not in ind_list:
#                         ind_list.append(i)
#     return sorted(ind_list)

